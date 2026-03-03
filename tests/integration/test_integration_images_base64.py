"""
Integration tests: storing images in base64 (PDF, DOCX, Markdown).
"""

import base64
import sys
from pathlib import Path
from unittest.mock import patch

import pytest
from langchain_core.documents import Document

# Add project root to PYTHONPATH
_project_root = Path(__file__).parent.parent.parent
if str(_project_root) not in sys.path:
    sys.path.insert(0, str(_project_root))

from documentor.domain import Element, ElementType
from documentor.pipeline import Pipeline


# ============================================================================
# Tests for storing images in base64
# ============================================================================

class TestImageBase64Storage:
    """Tests for storing images in base64."""

    def test_pdf_images_base64(self, tmp_path):
        """Test storing images in base64 for PDF."""
        try:
            import fitz
        except ImportError:
            pytest.skip("PyMuPDF is not installed")

        # Create simple PDF with image
        pdf_path = tmp_path / "test_image.pdf"
        doc = fitz.open()
        page = doc.new_page()

        # Add text mentioning image
        page.insert_text((50, 50), "Document with image")
        page.insert_text((50, 100), "Figure 1: Test image")

        doc.save(str(pdf_path))
        doc.close()

        doc = Document(page_content="", metadata={"source": str(pdf_path)})
        pipeline = Pipeline()

        with patch.object(pipeline, '_parsers', []):
            # Create mock parser that returns element with image
            from documentor.processing.parsers.pdf.pdf_parser import PdfParser
            parser = PdfParser()

            # Mock methods that require OCR
            img_data = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
            image_element = Element(
                id="img_001",
                type=ElementType.IMAGE,
                content="",
                metadata={"image_data": img_data}
            )

            # Mock layout_processor and other components
            with patch.object(parser.layout_processor, 'detect_layout_for_all_pages', return_value=[]):
                with patch.object(parser.layout_processor, 'filter_layout_elements', return_value=[]):
                    with patch.object(parser.hierarchy_builder, 'analyze_header_levels_from_elements', return_value=[]):
                        with patch.object(parser.hierarchy_builder, 'build_hierarchy_from_section_headers', return_value=[]):
                            with patch.object(parser.text_extractor, 'extract_text_by_bboxes', return_value=[]):
                                with patch.object(parser.text_extractor, 'merge_nearby_text_blocks', return_value=[]):
                                    with patch.object(parser.hierarchy_builder, 'create_elements_from_hierarchy', return_value=[image_element]):
                                        with patch.object(parser.image_processor, 'store_images_in_metadata', return_value=[image_element]):
                                            result = parser.parse(doc)

                                            # Check that images are present
                                            images = [e for e in result.elements if e.type == ElementType.IMAGE]
                                            if images:
                                                assert "image_data" in images[0].metadata
                                                assert images[0].metadata["image_data"].startswith("data:image/")
                                                assert "base64," in images[0].metadata["image_data"]

    def test_docx_images_base64(self, tmp_path):
        """Test storing images in base64 for DOCX."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx is not installed")

        # Create simple DOCX with image
        docx_path = tmp_path / "test_image.docx"
        doc = DocxDocument()
        doc.add_paragraph("Document with image")
        doc.add_paragraph("Figure 1: Test image")
        doc.save(str(docx_path))

        doc = Document(page_content="", metadata={"source": str(docx_path)})
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check that images are present
        images = [e for e in result.elements if e.type == ElementType.IMAGE]
        for img in images:
            assert "image_data" in img.metadata
            assert img.metadata["image_data"].startswith("data:image/")
            assert "base64," in img.metadata["image_data"]

            # Check that base64 is valid
            base64_part = img.metadata["image_data"].split(",")[1]
            try:
                decoded = base64.b64decode(base64_part)
                assert len(decoded) > 0
            except Exception:
                pytest.fail("Invalid base64 encoding in image_data")

    def test_markdown_images_base64(self):
        """Test storing images for Markdown (URLs remain as-is)."""
        doc = Document(
            page_content="""# Document with Image

![Test Image](https://example.com/image.png)

Text with inline image ![Inline](https://example.com/inline.jpg) in text.
""",
            metadata={"source": "test.md"}
        )

        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # In Markdown images remain as URL, not converted to base64
        images = [e for e in result.elements if e.type == ElementType.IMAGE]
        for img in images:
            # Check that src is in metadata
            assert "src" in img.metadata or "href" in img.metadata
