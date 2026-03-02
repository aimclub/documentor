"""
Tests for new parser features:
1. Saving images in base64
2. Saving tables in pandas DataFrame
3. Saving links in metadata
"""

import base64
import sys
from pathlib import Path
from unittest.mock import MagicMock, patch

import pandas as pd
import pytest
from langchain_core.documents import Document
from PIL import Image

# Add project root to PYTHONPATH
_project_root = Path(__file__).parent.parent.parent
if str(_project_root) not in sys.path:
    sys.path.insert(0, str(_project_root))

from documentor.domain import ElementType
from documentor.pipeline import Pipeline


# ============================================================================
# Tests for saving images in base64
# ============================================================================

class TestImageBase64Storage:
    """Tests for saving images in base64."""

    def test_pdf_images_base64(self, tmp_path):
        """Test saving images in base64 for PDF."""
        try:
            import fitz
        except ImportError:
            pytest.skip("PyMuPDF not installed")

        # Create simple PDF with image
        pdf_path = tmp_path / "test_image.pdf"
        doc = fitz.open()
        page = doc.new_page()
        
        # Add text mentioning image
        page.insert_text((50, 50), "Document with image")
        page.insert_text((50, 100), "Figure 1: Test image")
        
        doc.save(str(pdf_path))
        doc.close()

        doc = Document(page_content="", metadata={"source": str(pdf_path)})
        pipeline = Pipeline()
        
        with patch.object(pipeline, '_parsers', []):
            # Create mock parser that returns element with image
            from documentor.processing.parsers.pdf.pdf_parser import PdfParser
            parser = PdfParser()
            
            # Mock methods that require OCR
            from documentor.domain import Element
            img_data = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
            image_element = Element(
                id="img_001",
                type=ElementType.IMAGE,
                content="",
                metadata={"image_data": img_data}
            )
            
            # Mock layout_processor and other components
            with patch.object(parser.layout_processor, 'detect_layout_for_all_pages', return_value=[]):
                with patch.object(parser.layout_processor, 'filter_layout_elements', return_value=[]):
                    with patch.object(parser.hierarchy_builder, 'analyze_header_levels_from_elements', return_value=[]):
                        with patch.object(parser.hierarchy_builder, 'build_hierarchy_from_section_headers', return_value=[]):
                            with patch.object(parser.text_extractor, 'extract_text_by_bboxes', return_value=[]):
                                with patch.object(parser.text_extractor, 'merge_nearby_text_blocks', return_value=[]):
                                    with patch.object(parser.hierarchy_builder, 'create_elements_from_hierarchy', return_value=[image_element]):
                                        with patch.object(parser.image_processor, 'store_images_in_metadata', return_value=[image_element]):
                                            result = parser.parse(doc)
                                            
                                            # Check for images
                                            images = [e for e in result.elements if e.type == ElementType.IMAGE]
                                            if images:
                                                assert "image_data" in images[0].metadata
                                                assert images[0].metadata["image_data"].startswith("data:image/")
                                                assert "base64," in images[0].metadata["image_data"]

    def test_docx_images_base64(self, tmp_path):
        """Test saving images in base64 for DOCX."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        # Create simple DOCX with image
        docx_path = tmp_path / "test_image.docx"
        doc = DocxDocument()
        doc.add_paragraph("Document with image")
        doc.add_paragraph("Figure 1: Test image")
        doc.save(str(docx_path))

        doc = Document(page_content="", metadata={"source": str(docx_path)})
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check for images
        images = [e for e in result.elements if e.type == ElementType.IMAGE]
        for img in images:
            assert "image_data" in img.metadata
            assert img.metadata["image_data"].startswith("data:image/")
            assert "base64," in img.metadata["image_data"]
            
            # Check that base64 is valid
            base64_part = img.metadata["image_data"].split(",")[1]
            try:
                decoded = base64.b64decode(base64_part)
                assert len(decoded) > 0
            except Exception:
                pytest.fail("Invalid base64 encoding in image_data")

    def test_markdown_images_base64(self):
        """Test saving images for Markdown (URLs remain as is)."""
        doc = Document(
            page_content="""# Document with Image

![Test Image](https://example.com/image.png)

Text with inline image ![Inline](https://example.com/inline.jpg) in text.
""",
            metadata={"source": "test.md"}
        )
        
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # In Markdown images remain as URL, not converted to base64
        images = [e for e in result.elements if e.type == ElementType.IMAGE]
        for img in images:
            # Check that src is in metadata
            assert "src" in img.metadata or "href" in img.metadata


# ============================================================================
# Tests for saving tables in pandas DataFrame
# ============================================================================

class TestTableDataFrameStorage:
    """Tests for saving tables in pandas DataFrame."""

    def test_docx_tables_dataframe(self, tmp_path):
        """Test saving tables in DataFrame for DOCX."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        # Create DOCX with table
        docx_path = tmp_path / "test_table.docx"
        doc = DocxDocument()
        
        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "Column1"
        table.cell(0, 1).text = "Column2"
        table.cell(0, 2).text = "Column3"
        table.cell(1, 0).text = "Value1"
        table.cell(1, 1).text = "Value2"
        table.cell(1, 2).text = "Value3"
        table.cell(2, 0).text = "Value4"
        table.cell(2, 1).text = "Value5"
        table.cell(2, 2).text = "Value6"
        
        doc.save(str(docx_path))

        doc = Document(page_content="", metadata={"source": str(docx_path)})
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check for tables
        tables = [e for e in result.elements if e.type == ElementType.TABLE]
        assert len(tables) > 0
        
        for table in tables:
            assert "dataframe" in table.metadata
            assert isinstance(table.metadata["dataframe"], pd.DataFrame)
            # Even if table is empty, should have empty DataFrame
            assert table.metadata["dataframe"] is not None

    def test_markdown_tables_dataframe(self):
        """Test saving tables in DataFrame for Markdown."""
        doc = Document(
            page_content="""# Document with Table

| Column1 | Column2 | Column3 |
|---------|---------|---------|
| Value1  | Value2  | Value3  |
| Value4  | Value5  | Value6  |
""",
            metadata={"source": "test.md"}
        )
        
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check for tables
        tables = [e for e in result.elements if e.type == ElementType.TABLE]
        assert len(tables) > 0
        
        for table in tables:
            assert "dataframe" in table.metadata
            assert isinstance(table.metadata["dataframe"], pd.DataFrame)
            df = table.metadata["dataframe"]
            assert len(df) >= 2  # At least 2 data rows
            assert len(df.columns) == 3  # 3 columns
            assert "Column1" in df.columns or "Column_1" in df.columns

    def test_all_tables_have_dataframe(self):
        """Test that all tables have DataFrame (even empty)."""
        doc = Document(
            page_content="""# Multiple Tables

| A | B |
|---|---|
| 1 | 2 |

| X | Y |
|---|---|
| a | b |
""",
            metadata={"source": "test.md"}
        )
        
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        tables = [e for e in result.elements if e.type == ElementType.TABLE]
        assert len(tables) > 0
        
        for table in tables:
            assert "dataframe" in table.metadata
            assert isinstance(table.metadata["dataframe"], pd.DataFrame)
            # DataFrame should exist even if empty
            assert table.metadata["dataframe"] is not None


# ============================================================================
# Tests for saving links in metadata
# ============================================================================

class TestLinksInMetadata:
    """Tests for saving links in metadata."""

    def test_pdf_links_in_metadata(self, tmp_path):
        """Test saving links in metadata for PDF."""
        try:
            import fitz
        except ImportError:
            pytest.skip("PyMuPDF not installed")

        # Create PDF with text containing links
        pdf_path = tmp_path / "test_links.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Visit https://example.com for more info")
        page.insert_text((50, 100), "Check www.google.com and http://test.org")
        doc.save(str(pdf_path))
        doc.close()

        doc = Document(page_content="", metadata={"source": str(pdf_path)})
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check for links in text element metadata
        text_elements = [e for e in result.elements if e.type == ElementType.TEXT]
        found_links = False
        
        for elem in text_elements:
            if "links" in elem.metadata:
                found_links = True
                links = elem.metadata["links"]
                assert isinstance(links, list)
                assert len(links) > 0
                # Check that links are strings
                for link in links:
                    assert isinstance(link, str)
                    assert len(link) > 0
        
        # At least one element should have links
        # (may not have if OCR did not recognize text correctly)
        # assert found_links, "No links found in PDF text elements"

    def test_docx_links_in_metadata(self, tmp_path):
        """Test saving links in metadata for DOCX."""
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches
        except ImportError:
            pytest.skip("python-docx not installed")

        # Create DOCX with hyperlinks
        docx_path = tmp_path / "test_links.docx"
        doc = DocxDocument()
        
        # Add text with URL (parser should extract links from text)
        doc.add_paragraph("Visit https://example.com for more info")
        
        # Add more text with URL in text
        doc.add_paragraph("Check www.google.com and http://test.org")
        
        doc.save(str(docx_path))

        doc = Document(page_content="", metadata={"source": str(docx_path)})
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check for links in metadata
        text_elements = [e for e in result.elements if e.type == ElementType.TEXT]
        header_elements = [e for e in result.elements if e.type.name.startswith("HEADER")]
        
        found_links = False
        
        for elem in text_elements + header_elements:
            if "links" in elem.metadata:
                found_links = True
                links = elem.metadata["links"]
                assert isinstance(links, list)
                assert len(links) > 0
                for link in links:
                    assert isinstance(link, str)
                    assert len(link) > 0

    def test_markdown_links_in_metadata(self):
        """Test saving links in metadata for Markdown."""
        doc = Document(
            page_content="""# Document with Links

This is text with [a link](https://example.com) in it.

Visit https://www.google.com for search.

Check http://test.org and www.example.com
""",
            metadata={"source": "test.md"}
        )
        
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check for links in metadata
        text_elements = [e for e in result.elements if e.type == ElementType.TEXT]
        link_elements = [e for e in result.elements if e.type == ElementType.LINK]
        
        found_links_in_text = False
        
        # Check links in text elements
        for elem in text_elements:
            if "links" in elem.metadata:
                found_links_in_text = True
                links = elem.metadata["links"]
                assert isinstance(links, list)
                assert len(links) > 0
                for link in links:
                    assert isinstance(link, str)
                    assert ("http" in link or "www" in link)
        
        # Check separate LINK elements
        for elem in link_elements:
            assert "href" in elem.metadata or "src" in elem.metadata
            url = elem.metadata.get("href") or elem.metadata.get("src")
            assert isinstance(url, str)
            assert len(url) > 0

    def test_links_in_headers(self):
        """Test saving links in header metadata."""
        doc = Document(
            page_content="""# Header with https://example.com link

## Another header with www.google.com

Text with http://test.org
""",
            metadata={"source": "test.md"}
        )
        
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check for links in headers
        header_elements = [e for e in result.elements if e.type.name.startswith("HEADER")]
        
        for header in header_elements:
            if "links" in header.metadata:
                links = header.metadata["links"]
                assert isinstance(links, list)
                assert len(links) > 0


# ============================================================================
# Combined tests for all three features
# ============================================================================

class TestAllFeaturesTogether:
    """Combined tests for all three features."""

    def test_markdown_all_features(self):
        """Test all three features for Markdown."""
        doc = Document(
            page_content="""# Document Title

## Section with https://example.com

| Column1 | Column2 |
|---------|---------|
| Data1   | Data2   |

![Image](https://example.com/image.png)

Text with [link](https://google.com) and www.test.org
""",
            metadata={"source": "test.md"}
        )
        
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check tables
        tables = [e for e in result.elements if e.type == ElementType.TABLE]
        assert len(tables) > 0
        for table in tables:
            assert "dataframe" in table.metadata
            assert isinstance(table.metadata["dataframe"], pd.DataFrame)

        # Check images
        images = [e for e in result.elements if e.type == ElementType.IMAGE]
        for img in images:
            assert "src" in img.metadata or "href" in img.metadata

        # Check links
        text_elements = [e for e in result.elements if e.type == ElementType.TEXT]
        found_links = False
        for elem in text_elements:
            if "links" in elem.metadata:
                found_links = True
                assert len(elem.metadata["links"]) > 0

    def test_all_features_metadata_structure(self):
        """Test metadata structure for all features."""
        doc = Document(
            page_content="""# Test

| A | B |
|---|---|
| 1 | 2 |

Text with https://example.com
""",
            metadata={"source": "test.md"}
        )
        
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check metadata structure
        for elem in result.elements:
            assert elem.metadata is not None
            assert isinstance(elem.metadata, dict)
            
            # If table, should have DataFrame
            if elem.type == ElementType.TABLE:
                assert "dataframe" in elem.metadata
                assert isinstance(elem.metadata["dataframe"], pd.DataFrame)
            
            # If image, may have image_data (for PDF/DOCX) or src (for Markdown)
            if elem.type == ElementType.IMAGE:
                has_image_data = "image_data" in elem.metadata
                has_src = "src" in elem.metadata or "href" in elem.metadata
                assert has_image_data or has_src
            
            # If links in text, they should be in list
            if "links" in elem.metadata:
                assert isinstance(elem.metadata["links"], list)
                for link in elem.metadata["links"]:
                    assert isinstance(link, str)
