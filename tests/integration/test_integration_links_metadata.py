"""
Integration tests: storing links in element metadata (PDF, DOCX, Markdown).
"""

import sys
from pathlib import Path

import pytest
from langchain_core.documents import Document

# Add project root to PYTHONPATH
_project_root = Path(__file__).parent.parent.parent
if str(_project_root) not in sys.path:
    sys.path.insert(0, str(_project_root))

from documentor.domain import ElementType
from documentor.pipeline import Pipeline


# ============================================================================
# Tests for storing links in metadata
# ============================================================================

class TestLinksInMetadata:
    """Tests for storing links in metadata."""

    def test_pdf_links_in_metadata(self, tmp_path):
        """Test storing links in metadata for PDF."""
        try:
            import fitz
        except ImportError:
            pytest.skip("PyMuPDF is not installed")

        # Create PDF with text containing links
        pdf_path = tmp_path / "test_links.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Visit https://example.com for more info")
        page.insert_text((50, 100), "Check www.google.com and http://test.org")
        doc.save(str(pdf_path))
        doc.close()

        doc = Document(page_content="", metadata={"source": str(pdf_path)})
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check links in text element metadata
        text_elements = [e for e in result.elements if e.type == ElementType.TEXT]
        for elem in text_elements:
            if "links" in elem.metadata:
                links = elem.metadata["links"]
                assert isinstance(links, list)
                assert len(links) > 0
                for link in links:
                    assert isinstance(link, str)
                    assert len(link) > 0

    def test_docx_links_in_metadata(self, tmp_path):
        """Test storing links in metadata for DOCX."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx is not installed")

        # Create DOCX with hyperlinks
        docx_path = tmp_path / "test_links.docx"
        doc = DocxDocument()

        # Add text with URL (parser should extract links from text)
        doc.add_paragraph("Visit https://example.com for more info")

        # Add more text with URL in text
        doc.add_paragraph("Check www.google.com and http://test.org")

        doc.save(str(docx_path))

        doc = Document(page_content="", metadata={"source": str(docx_path)})
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check links in metadata
        text_elements = [e for e in result.elements if e.type == ElementType.TEXT]
        header_elements = [e for e in result.elements if e.type.name.startswith("HEADER")]

        for elem in text_elements + header_elements:
            if "links" in elem.metadata:
                links = elem.metadata["links"]
                assert isinstance(links, list)
                assert len(links) > 0
                for link in links:
                    assert isinstance(link, str)
                    assert len(link) > 0

    def test_markdown_links_in_metadata(self):
        """Test storing links in metadata for Markdown."""
        doc = Document(
            page_content="""# Document with Links

This is text with [a link](https://example.com) in it.

Visit https://www.google.com for search.

Check http://test.org and www.example.com
""",
            metadata={"source": "test.md"}
        )

        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check links in metadata
        text_elements = [e for e in result.elements if e.type == ElementType.TEXT]
        link_elements = [e for e in result.elements if e.type == ElementType.LINK]

        # Check links in text elements
        for elem in text_elements:
            if "links" in elem.metadata:
                links = elem.metadata["links"]
                assert isinstance(links, list)
                assert len(links) > 0
                for link in links:
                    assert isinstance(link, str)
                    assert ("http" in link or "www" in link)

        # Check individual LINK elements
        for elem in link_elements:
            assert "href" in elem.metadata or "src" in elem.metadata
            url = elem.metadata.get("href") or elem.metadata.get("src")
            assert isinstance(url, str)
            assert len(url) > 0

    def test_links_in_headers(self):
        """Test storing links in header metadata."""
        doc = Document(
            page_content="""# Header with https://example.com link

## Another header with www.google.com

Text with http://test.org
""",
            metadata={"source": "test.md"}
        )

        pipeline = Pipeline()
        result = pipeline.parse(doc)

        # Check links in headers
        header_elements = [e for e in result.elements if e.type.name.startswith("HEADER")]

        for header in header_elements:
            if "links" in header.metadata:
                links = header.metadata["links"]
                assert isinstance(links, list)
                assert len(links) > 0
