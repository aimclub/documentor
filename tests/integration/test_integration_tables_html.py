"""
Integration tests: tables stored as HTML in element content (DOCX, Markdown).
"""

import sys
from pathlib import Path

import pytest
from langchain_core.documents import Document

# Add project root to PYTHONPATH
_project_root = Path(__file__).parent.parent.parent
if str(_project_root) not in sys.path:
    sys.path.insert(0, str(_project_root))

from documentor.domain import ElementType
from documentor.pipeline import Pipeline


# ============================================================================
# Tests for storing tables as HTML in content
# ============================================================================

class TestTableHtmlContent:
    """Tests for tables stored as HTML in element content."""

    def test_docx_tables_html(self, tmp_path):
        """Test that DOCX tables have HTML in content."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx is not installed")

        docx_path = tmp_path / "test_table.docx"
        doc = DocxDocument()
        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "Column1"
        table.cell(0, 1).text = "Column2"
        table.cell(0, 2).text = "Column3"
        table.cell(1, 0).text = "Value1"
        table.cell(1, 1).text = "Value2"
        table.cell(1, 2).text = "Value3"
        table.cell(2, 0).text = "Value4"
        table.cell(2, 1).text = "Value5"
        table.cell(2, 2).text = "Value6"
        doc.save(str(docx_path))

        doc = Document(page_content="", metadata={"source": str(docx_path)})
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        tables = [e for e in result.elements if e.type == ElementType.TABLE]
        assert len(tables) > 0
        for table in tables:
            assert isinstance(table.content, str)
            if table.content.strip():
                assert "<table>" in table.content

    def test_markdown_tables_html(self):
        """Test that Markdown tables have HTML in content."""
        doc = Document(
            page_content="""# Document with Table

| Column1 | Column2 | Column3 |
|---------|---------|---------|
| Value1  | Value2  | Value3  |
| Value4  | Value5  | Value6  |
""",
            metadata={"source": "test.md"}
        )
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        tables = [e for e in result.elements if e.type == ElementType.TABLE]
        assert len(tables) > 0
        for table in tables:
            assert isinstance(table.content, str)
            assert "<table>" in table.content
            assert "</table>" in table.content

    def test_all_tables_have_content(self):
        """Test that all table elements have content (HTML)."""
        doc = Document(
            page_content="""# Multiple Tables

| A | B |
|---|---|
| 1 | 2 |

| X | Y |
|---|---|
| a | b |
""",
            metadata={"source": "test.md"}
        )
        pipeline = Pipeline()
        result = pipeline.parse(doc)

        tables = [e for e in result.elements if e.type == ElementType.TABLE]
        assert len(tables) > 0
        for table in tables:
            assert isinstance(table.content, str)
            assert "<table>" in table.content
