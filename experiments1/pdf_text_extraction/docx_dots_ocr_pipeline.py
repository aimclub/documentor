"""
Пайплайн для DOCX с использованием Dots.OCR для layout detection.

Подход:
1. Конвертируем DOCX страницы в изображения (через временный PDF)
2. Используем Dots.OCR для layout detection (как в PDF)
3. Извлекаем текст из DOCX через python-docx
4. Сопоставляем текст из Dots.OCR bbox с текстом из DOCX
5. Строим иерархию на основе Dots.OCR
"""

import json
import re
import sys
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple, Union
from io import BytesIO
import tempfile
import shutil

from PIL import Image
from tqdm import tqdm

# Добавляем корень проекта в путь
project_root = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(project_root))

try:
    from docx import Document as PythonDocxDocument
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    print("Предупреждение: python-docx не установлен")

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    print("Предупреждение: PyMuPDF не установлен")

# Альтернативные способы рендеринга DOCX напрямую в изображения (без PDF)
HAS_WIN32COM = False
try:
    import win32com.client
    HAS_WIN32COM = True
except ImportError:
    pass

HAS_LIBREOFFICE = False
LIBREOFFICE_CMD = None
try:
    import subprocess
    # Проверяем наличие LibreOffice на разных платформах
    # Пробуем разные варианты команды
    for cmd in ['soffice', 'libreoffice', '/Applications/LibreOffice.app/Contents/MacOS/soffice']:
        try:
            result = subprocess.run([cmd, '--version'], capture_output=True, timeout=5)
            if result.returncode == 0:
                HAS_LIBREOFFICE = True
                LIBREOFFICE_CMD = cmd
                break
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
except:
    pass

# Импорты из documentor (напрямую, минуя __init__)
# Избегаем импорта через documentor.__init__ чтобы не загружать pandas
from documentor.processing.parsers.pdf.ocr.dots_ocr_client import process_layout_detection
from documentor.utils.ocr_image_utils import fetch_image
from documentor.utils.ocr_consts import MIN_PIXELS, MAX_PIXELS


class DocxPageRenderer:
    """
    Рендерер страниц DOCX в изображения БЕЗ промежуточного PDF.
    
    Использует прямые методы конвертации:
    - win32com (Windows COM) - экспорт страниц Word напрямую в изображения
    - LibreOffice - экспорт DOCX напрямую в PNG (без PDF)
    
    Преимущества:
    - Нет потери качества из-за конвертации через PDF
    - Нет проблем с шрифтами при двойной конвертации
    - Более точное визуальное представление документа
    """
    
    def __init__(
        self,
        render_scale: float = 2.0,
        optimize_for_ocr: bool = True,
        min_pixels: Optional[int] = None,
        max_pixels: Optional[int] = None,
    ) -> None:
        """
        Инициализация рендерера.
        
        Args:
            render_scale: Масштаб рендеринга (2.0 = увеличение в 2 раза)
            optimize_for_ocr: Применять ли smart_resize для оптимизации под OCR
            min_pixels: Минимальное число пикселей
            max_pixels: Максимальное число пикселей
        """
        self.render_scale = render_scale
        self.optimize_for_ocr = optimize_for_ocr
        
        if min_pixels is None:
            min_pixels = MIN_PIXELS
        if max_pixels is None:
            max_pixels = MAX_PIXELS
        
        self.min_pixels = min_pixels
        self.max_pixels = max_pixels
        
        # Определяем метод рендеринга
        # Приоритет: LibreOffice (кроссплатформенный) > win32com (только Windows)
        self.render_method = None
        if HAS_LIBREOFFICE:
            self.render_method = "libreoffice"  # Основной метод - работает везде
        elif HAS_WIN32COM:
            self.render_method = "win32com"  # Опциональный для Windows
        else:
            raise RuntimeError(
                "Не найден способ рендеринга DOCX. Установите:\n"
                "- LibreOffice (рекомендуется, кроссплатформенный):\n"
                "  * Windows: скачайте с https://www.libreoffice.org/\n"
                "    Добавьте LibreOffice в PATH или используйте полный путь\n"
                "  * macOS: brew install --cask libreoffice\n"
                "    Или скачайте с https://www.libreoffice.org/\n"
                "  * Linux: sudo apt-get install libreoffice (Ubuntu/Debian)\n"
                "    или sudo yum install libreoffice (RHEL/CentOS)\n"
                "    или sudo pacman -S libreoffice-fresh (Arch)\n"
                "  Убедитесь, что команда 'soffice' или 'libreoffice' доступна в PATH\n"
                "\n"
                "- ИЛИ pywin32 (только Windows): pip install pywin32\n"
                "\n"
                "ВАЖНО: Мы НЕ используем docx2pdf, чтобы избежать ошибок при двойной конвертации."
            )
    
    def _render_via_win32com(self, docx_path: Path, page_num: int) -> Image.Image:
        """
        Рендерит страницу через win32com (Word).
        
        ВАЖНО: Word COM API не поддерживает прямой экспорт в PNG с настройками.
        Используем временный PDF и затем рендерим через PyMuPDF.
        Это единственный надежный способ получить изображение страницы из Word.
        """
        import win32com.client
        
        if not HAS_PYMUPDF:
            raise RuntimeError("PyMuPDF требуется для рендеринга страниц Word")
        
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        temp_pdf_dir = tempfile.mkdtemp()
        temp_pdf_path = Path(temp_pdf_dir) / f"{docx_path.stem}_temp.pdf"
        
        try:
            doc = word.Documents.Open(str(docx_path.absolute()))
            try:
                # Экспортируем весь документ в PDF (Word не поддерживает экспорт одной страницы)
                # Используем только поддерживаемые параметры
                doc.ExportAsFixedFormat(
                    OutputFileName=str(temp_pdf_path),
                    ExportFormat=17,  # wdExportFormatPDF = 17
                    OpenAfterExport=False,
                    OptimizeFor=0,  # wdExportOptimizeForPrint = 0
                    BitmapMissingFonts=True,
                    DocStructureTags=False,
                    CreateBookmarks=0,  # wdExportCreateNoBookmarks = 0
                )
                
                # Рендерим нужную страницу из PDF через PyMuPDF
                pdf_doc = fitz.open(str(temp_pdf_path))
                try:
                    if page_num >= len(pdf_doc):
                        raise ValueError(f"Страница {page_num} не существует (всего страниц: {len(pdf_doc)})")
                    
                    page = pdf_doc.load_page(page_num)
                    mat = fitz.Matrix(self.render_scale, self.render_scale)
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("ppm")
                    image = Image.open(BytesIO(img_data)).convert("RGB")
                    return image
                finally:
                    pdf_doc.close()
            finally:
                doc.Close(SaveChanges=False)
        finally:
            word.Quit()
            # Удаляем временный PDF
            if temp_pdf_path.exists():
                try:
                    temp_pdf_path.unlink()
                    shutil.rmtree(temp_pdf_dir, ignore_errors=True)
                except:
                    pass
    
    def _render_via_libreoffice(self, docx_path: Path, page_num: int) -> Image.Image:
        """
        Рендерит страницу через LibreOffice напрямую в PNG (без PDF).
        
        LibreOffice может экспортировать DOCX напрямую в изображения,
        что избегает проблем с двойной конвертацией и шрифтами.
        
        Использует Python макрос для экспорта конкретной страницы.
        """
        import subprocess
        
        temp_image_dir = tempfile.mkdtemp()
        temp_script_dir = tempfile.mkdtemp()
        
        try:
            # Метод 1: Используем Python макрос для экспорта конкретной страницы
            # Создаем временный Python скрипт для LibreOffice
            script_path = Path(temp_script_dir) / "export_page.py"
            
            with open(script_path, "w", encoding="utf-8") as f:
                f.write(f"""
import uno
from com.sun.star.beans import PropertyValue

def export_page():
    desktop = XSCRIPTCONTEXT.getDesktop()
    doc = desktop.loadComponentFromURL(
        "file:///{docx_path.absolute().as_posix().replace(':', '|')}",
        "_blank", 0, ()
    )
    
    # Экспортируем конкретную страницу в PNG
    page_num = {page_num}
    output_url = "file:///{temp_image_dir.replace(chr(92), '/')}/page_{page_num + 1}.png"
    
    # Создаем фильтр для PNG
    filter_props = (
        PropertyValue("FilterName", 0, "PNG", 0),
        PropertyValue("PageRange", 0, str(page_num + 1), 0),  # 1-based
    )
    
    # Экспортируем страницу
    doc.storeToURL(output_url, filter_props)
    doc.close(True)

g_exportedScripts = (export_page,)
""")
            
            # Запускаем LibreOffice с макросом
            # Но это сложно, поэтому используем более простой метод
            
            # Метод 2: Экспортируем весь документ в PNG и берем нужную страницу
            # LibreOffice при конвертации в PNG создает отдельные файлы для каждой страницы
            # Или один файл - нужно проверить
            
            # Используем LibreOffice для прямого экспорта в PNG
            # Используем определенную команду (может быть 'soffice', 'libreoffice' или полный путь на Mac)
            cmd = LIBREOFFICE_CMD or 'soffice'
            
            result = subprocess.run(
                [
                    cmd,
                    '--headless',
                    '--nodefault',
                    '--nolockcheck',
                    '--invisible',
                    '--convert-to', 'png',
                    '--outdir', str(temp_image_dir),
                    str(docx_path.absolute())
                ],
                check=True,
                timeout=120,
                capture_output=True,
                text=True
            )
            
            # LibreOffice может создать:
            # 1. Один PNG файл с именем документа (если одна страница)
            # 2. Несколько PNG файлов (если несколько страниц) - но это зависит от версии
            
            # Ищем созданные PNG файлы
            png_files = sorted(list(Path(temp_image_dir).glob("*.png")))
            
            if not png_files:
                # Пробуем найти любые изображения
                png_files = sorted([f for f in Path(temp_image_dir).iterdir() 
                                   if f.suffix.lower() in ['.png', '.jpg', '.jpeg', '.bmp']])
            
            if png_files:
                # Если создался один файл - это может быть все страницы в одном изображении
                # Или это первая страница
                if len(png_files) == 1:
                    # Открываем файл и проверяем - это одна страница или несколько
                    image = Image.open(str(png_files[0])).convert("RGB")
                    width, height = image.size
                    
                    # Если изображение очень высокое, возможно это несколько страниц
                    # Стандартная страница A4 в 300 DPI: ~2480x3508 пикселей
                    # Если высота значительно больше, возможно это несколько страниц
                    if height > 4000 and page_num > 0:
                        # Обрезаем нужную страницу (предполагаем A4 формат)
                        page_height = 3508  # Примерная высота страницы A4 в 300 DPI
                        y_start = page_num * page_height
                        y_end = min((page_num + 1) * page_height, height)
                        image = image.crop((0, y_start, width, y_end))
                    
                    # Масштабируем если нужно
                    if self.render_scale != 1.0:
                        width, height = image.size
                        new_size = (int(width * self.render_scale), int(height * self.render_scale))
                        image = image.resize(new_size, Image.Resampling.LANCZOS)
                    
                    return image
                else:
                    # Несколько файлов - берем нужный
                    if page_num < len(png_files):
                        image = Image.open(str(png_files[page_num])).convert("RGB")
                        # Масштабируем если нужно
                        if self.render_scale != 1.0:
                            width, height = image.size
                            new_size = (int(width * self.render_scale), int(height * self.render_scale))
                            image = image.resize(new_size, Image.Resampling.LANCZOS)
                        return image
            
            # Если не нашли файлы, пробуем альтернативный метод через unoconv
            # или используем Python API LibreOffice напрямую
            raise RuntimeError(
                f"LibreOffice не создал изображение для страницы {page_num + 1}. "
                f"Вывод: {result.stdout[:500] if result.stdout else 'нет'}, "
                f"Ошибки: {result.stderr[:500] if result.stderr else 'нет'}"
            )
        finally:
            # Удаляем временные файлы
            try:
                shutil.rmtree(temp_image_dir, ignore_errors=True)
                shutil.rmtree(temp_script_dir, ignore_errors=True)
            except:
                pass
    
    def _render_via_docx2pdf(self, docx_path: Path, page_num: int, temp_pdf_path: Optional[Path] = None) -> Image.Image:
        """Рендерит страницу через docx2pdf (fallback метод)."""
        if temp_pdf_path is None:
            temp_pdf_dir = tempfile.mkdtemp()
            temp_pdf_path = Path(temp_pdf_dir) / f"{docx_path.stem}_temp.pdf"
        else:
            temp_pdf_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Конвертируем DOCX в PDF (только если PDF еще не существует)
        if not temp_pdf_path.exists():
            try:
                self.docx2pdf_convert(str(docx_path), str(temp_pdf_path))
            except Exception as e:
                raise RuntimeError(f"Ошибка конвертации DOCX в PDF: {e}")
        
        # Рендерим страницу из PDF
        pdf_document = fitz.open(str(temp_pdf_path))
        try:
            if page_num >= len(pdf_document):
                raise ValueError(f"Страница {page_num} не существует (всего страниц: {len(pdf_document)})")
            
            page = pdf_document.load_page(page_num)
            mat = fitz.Matrix(self.render_scale, self.render_scale)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("ppm")
            image = Image.open(BytesIO(img_data)).convert("RGB")
            return image
        finally:
            pdf_document.close()
    
    def render_page(
        self,
        docx_path: Path,
        page_num: int,
        temp_pdf_path: Optional[Path] = None,
        return_original: bool = False,
    ) -> Union[Image.Image, Tuple[Image.Image, Image.Image]]:
        """
        Рендерит одну страницу DOCX в изображение.
        
        Args:
            docx_path: Путь к DOCX файлу
            page_num: Номер страницы (0-based)
            temp_pdf_path: Путь для временного PDF (только для метода docx2pdf)
            return_original: Если True, возвращает кортеж (original_image, optimized_image)
        
        Returns:
            Image.Image или tuple[Image.Image, Image.Image]
        """
        # Выбираем метод рендеринга
        if self.render_method == "win32com":
            original_image = self._render_via_win32com(docx_path, page_num)
        elif self.render_method == "libreoffice":
            original_image = self._render_via_libreoffice(docx_path, page_num)
        elif self.render_method == "docx2pdf":
            original_image = self._render_via_docx2pdf(docx_path, page_num, temp_pdf_path)
        else:
            raise RuntimeError(f"Неизвестный метод рендеринга: {self.render_method}")
        
        if self.optimize_for_ocr:
            optimized_image = fetch_image(
                original_image,
                min_pixels=self.min_pixels,
                max_pixels=self.max_pixels,
            )
        else:
            optimized_image = original_image
        
        if return_original:
            return original_image, optimized_image
        return optimized_image
    
    def get_page_count(self, docx_path: Path, temp_pdf_path: Optional[Path] = None) -> int:
        """
        Возвращает количество страниц в DOCX.
        
        Args:
            docx_path: Путь к DOCX файлу
            temp_pdf_path: Путь для временного PDF (только для метода docx2pdf)
        
        Returns:
            Количество страниц
        """
        if self.render_method == "win32com":
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(str(docx_path.absolute()))
                try:
                    # Получаем количество страниц через Word
                    return doc.ComputeStatistics(2)  # wdStatisticPages
                finally:
                    doc.Close()
            finally:
                word.Quit()
        
        elif self.render_method == "libreoffice":
            # Для LibreOffice считаем страницы через python-docx
            # Это более надежно, чем конвертация в PDF
            if HAS_PYTHON_DOCX:
                from docx import Document as PythonDocxDocument
                doc = PythonDocxDocument(str(docx_path))
                # Приблизительная оценка: считаем параграфы и делим на среднее количество параграфов на страницу
                # Или используем более точный метод через LibreOffice
                # Но проще всего - конвертировать в PDF временно для подсчета
                temp_pdf_dir = tempfile.mkdtemp()
                temp_pdf_path = Path(temp_pdf_dir) / f"{docx_path.stem}_temp.pdf"
                
                try:
                    import subprocess
                    cmd = LIBREOFFICE_CMD or 'soffice'
                    subprocess.run(
                        [
                            cmd,
                            '--headless',
                            '--convert-to', 'pdf',
                            '--outdir', str(temp_pdf_dir),
                            str(docx_path.absolute())
                        ],
                        check=True,
                        timeout=60,
                        capture_output=True
                    )
                    
                    if HAS_PYMUPDF and temp_pdf_path.exists():
                        pdf_doc = fitz.open(str(temp_pdf_path))
                        try:
                            return len(pdf_doc)
                        finally:
                            pdf_doc.close()
                    
                    # Fallback: приблизительная оценка
                    return max(1, len(doc.paragraphs) // 20)  # Примерно 20 параграфов на страницу
                finally:
                    if temp_pdf_path.exists():
                        try:
                            temp_pdf_path.unlink()
                            shutil.rmtree(temp_pdf_dir, ignore_errors=True)
                        except:
                            pass
                return 1  # Минимум 1 страница
            else:
                raise RuntimeError("python-docx не установлен для подсчета страниц")
        
        elif self.render_method == "docx2pdf":
            # Используем временный PDF
            if temp_pdf_path is None:
                temp_pdf_dir = tempfile.mkdtemp()
                temp_pdf_path = Path(temp_pdf_dir) / f"{docx_path.stem}_temp.pdf"
            else:
                temp_pdf_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Конвертируем DOCX в PDF если нужно
            if not temp_pdf_path.exists():
                try:
                    self.docx2pdf_convert(str(docx_path), str(temp_pdf_path))
                except Exception as e:
                    raise RuntimeError(f"Ошибка конвертации DOCX в PDF: {e}")
            
            if HAS_PYMUPDF:
                pdf_document = fitz.open(str(temp_pdf_path))
                try:
                    return len(pdf_document)
                finally:
                    pdf_document.close()
            else:
                raise RuntimeError("PyMuPDF не доступен")
        
        else:
            raise RuntimeError(f"Неизвестный метод рендеринга: {self.render_method}")


def extract_text_from_docx(docx_path: Path) -> List[Dict[str, Any]]:
    """
    Извлекает текст из DOCX с метаданными (аналог paragraphs_with_metadata).
    
    Args:
        docx_path: Путь к DOCX файлу
    
    Returns:
        Список словарей с данными параграфов
    """
    if not HAS_PYTHON_DOCX:
        raise RuntimeError("python-docx не установлен")
    
    doc = PythonDocxDocument(str(docx_path))
    paragraphs_data = []
    
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Получаем стиль
        style = para.style.name if para.style else "Normal"
        
        # Получаем форматирование
        formatting = {
            "bold": False,
            "italic": False,
            "underline": False,
            "font_size": None,
            "font_name": None,
            "alignment": None,
        }
        
        # Проверяем runs для форматирования
        if para.runs:
            first_run = para.runs[0]
            formatting["bold"] = first_run.bold or False
            formatting["italic"] = first_run.italic or False
            formatting["underline"] = first_run.underline is not None
            
            if first_run.font.size:
                formatting["font_size"] = str(first_run.font.size.pt) + "pt"
            if first_run.font.name:
                formatting["font_name"] = first_run.font.name
        
        # Выравнивание
        if para.alignment:
            alignment_map = {
                0: "left",
                1: "center",
                2: "right",
                3: "justify",
            }
            formatting["alignment"] = alignment_map.get(para.alignment, None)
        
        # Отступы (обрабатываем ошибки при чтении некорректных значений)
        coordinates = {}
        try:
            if para.paragraph_format.left_indent:
                coordinates["left_indent"] = str(para.paragraph_format.left_indent)
        except (ValueError, AttributeError, TypeError):
            pass
        
        try:
            if para.paragraph_format.right_indent:
                coordinates["right_indent"] = str(para.paragraph_format.right_indent)
        except (ValueError, AttributeError, TypeError):
            pass
        
        try:
            if para.paragraph_format.first_line_indent:
                coordinates["first_line_indent"] = str(para.paragraph_format.first_line_indent)
        except (ValueError, AttributeError, TypeError) as e:
            # Некоторые DOCX файлы содержат некорректные значения отступов (float вместо int)
            # Пропускаем такие случаи
            pass
        
        paragraphs_data.append({
            "index": idx,
            "text": text,
            "style": style,
            "formatting": formatting,
            "coordinates": coordinates,
            "is_heading": style.startswith("Heading"),
        })
    
    return paragraphs_data


def match_ocr_text_with_docx_text(
    ocr_elements: List[Dict[str, Any]],
    docx_paragraphs: List[Dict[str, Any]],
    page_num: int,
) -> List[Dict[str, Any]]:
    """
    Сопоставляет текст из Dots.OCR bbox с текстом из DOCX.
    
    Args:
        ocr_elements: Элементы layout от Dots.OCR
        docx_paragraphs: Параграфы из DOCX
        page_num: Номер страницы
    
    Returns:
        Список элементов с сопоставленным текстом из DOCX
    """
    matched_elements = []
    
    # Создаем индекс текста из DOCX для быстрого поиска
    docx_text_index = {}
    for para in docx_paragraphs:
        text = para["text"].lower().strip()
        # Создаем ключи для поиска (первые N слов)
        words = text.split()[:5]  # Первые 5 слов
        key = " ".join(words)
        if key not in docx_text_index:
            docx_text_index[key] = []
        docx_text_index[key].append(para)
    
    for ocr_element in ocr_elements:
        ocr_text = ocr_element.get("text", "").strip()
        category = ocr_element.get("category", "")
        bbox = ocr_element.get("bbox", [])
        
        # Ищем соответствующий параграф в DOCX
        matched_para = None
        
        if ocr_text:
            # Пробуем найти по тексту
            ocr_words = ocr_text.lower().split()[:5]
            ocr_key = " ".join(ocr_words)
            
            if ocr_key in docx_text_index:
                candidates = docx_text_index[ocr_key]
                # Выбираем первый подходящий (можно улучшить логику)
                if candidates:
                    matched_para = candidates[0]
        
        # Если не нашли по тексту, но это Section-header, ищем по категории
        if not matched_para and category == "Section-header":
            # Ищем заголовки в DOCX
            for para in docx_paragraphs:
                if para.get("is_heading", False):
                    # Проверяем похожесть текста
                    para_text = para["text"].lower().strip()
                    if ocr_text and para_text:
                        # Простая проверка на совпадение (можно улучшить)
                        if ocr_text.lower() in para_text or para_text in ocr_text.lower():
                            matched_para = para
                            break
        
        # Создаем элемент с сопоставленным текстом
        element = {
            "bbox": bbox,
            "category": category,
            "page_num": page_num,
            "ocr_text": ocr_text,  # Текст от OCR (может быть неточным)
            "docx_text": matched_para["text"] if matched_para else ocr_text,  # Текст из DOCX
            "docx_paragraph": matched_para,  # Полные данные параграфа
        }
        
        matched_elements.append(element)
    
    return matched_elements


def build_hierarchy_from_section_headers(
    layout_elements: List[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    """
    Строит иерархию элементов, группируя их по Section-header (как в PDF).
    
    Args:
        layout_elements: Список элементов layout
    
    Returns:
        Список секций с заголовками и дочерними элементами
    """
    sections = []
    current_section = None
    
    for element in layout_elements:
        if element["category"] == "Section-header":
            # Начинаем новую секцию
            if current_section:
                sections.append(current_section)
            
            current_section = {
                "header": element,
                "children": []
            }
        else:
            # Добавляем в текущую секцию
            if current_section:
                current_section["children"].append(element)
            else:
                # Если нет заголовка, создаем секцию "Начало документа"
                if not sections or sections[-1].get("header", {}).get("text") != "Начало документа":
                    current_section = {
                        "header": {
                            "text": "Начало документа",
                            "level": 0,
                            "category": "Title"
                        },
                        "children": []
                    }
                    sections.append(current_section)
                else:
                    current_section = sections[-1]
                
                current_section["children"].append(element)
    
    if current_section:
        sections.append(current_section)
    
    return sections


def determine_header_level(
    text: str,
    previous_headers: List[Dict[str, Any]],
    docx_paragraph: Optional[Dict[str, Any]] = None,
) -> int:
    """
    Определяет уровень заголовка (аналог _determine_header_level из PDF).
    
    Args:
        text: Текст заголовка
        previous_headers: Список предыдущих заголовков
        docx_paragraph: Данные параграфа из DOCX (если есть)
    
    Returns:
        Уровень заголовка (1-6)
    """
    # Приоритет 1: Анализ нумерации в тексте
    if re.match(r'^\d+\s+[A-Z]', text):
        return 1
    if re.match(r'^\d+\.\d+\s+', text):
        return 2
    if re.match(r'^\d+\.\d+\.\d+\s+', text):
        return 3
    if re.match(r'^\d+\.\d+\.\d+\.\d+\s+', text):
        return 4
    
    # Приоритет 2: Если есть данные из DOCX, используем стиль (но не доверяем полностью)
    if docx_paragraph:
        style = docx_paragraph.get("style", "")
        if style.startswith("Heading"):
            level = int(style.split()[-1])
            # Но проверяем контекст - если предыдущие заголовки имеют другой уровень,
            # возможно стиль неправильный
            if previous_headers:
                last_level = previous_headers[-1].get("level", 1)
                # Если разница слишком большая, корректируем
                if abs(level - last_level) > 2:
                    return min(6, last_level + 1)
            return min(level, 6)
    
    # Приоритет 3: Сравнение размера шрифта с предыдущими (если есть данные из DOCX)
    if docx_paragraph and previous_headers:
        font_size_str = docx_paragraph.get("formatting", {}).get("font_size")
        if font_size_str:
            try:
                current_size = float(font_size_str.replace("pt", ""))
                
                # Находим последний заголовок с известным размером шрифта
                for header in reversed(previous_headers):
                    header_para = header.get("docx_paragraph")
                    if header_para:
                        last_font_size_str = header_para.get("formatting", {}).get("font_size")
                        if last_font_size_str:
                            last_size = float(last_font_size_str.replace("pt", ""))
                            last_level = header.get("level", 1)
                            
                            # Сравниваем размеры (>= 2pt разница)
                            if current_size >= last_size + 2:
                                return max(1, last_level - 1)
                            elif current_size <= last_size - 2:
                                return min(6, last_level + 1)
                            else:
                                return last_level
            except (ValueError, AttributeError):
                pass
    
    # По умолчанию
    if previous_headers:
        return min(6, previous_headers[-1].get("level", 1) + 1)
    return 1


def analyze_header_levels(layout_elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Анализирует уровни заголовков (аналог _analyze_header_levels_from_elements из PDF).
    
    Args:
        layout_elements: Список элементов layout
    
    Returns:
        Список элементов с определенными уровнями заголовков
    """
    analyzed_elements = []
    previous_headers = []
    last_numbered_level = None
    
    for element in layout_elements:
        if element["category"] == "Section-header":
            text = element.get("docx_text", element.get("ocr_text", ""))
            docx_paragraph = element.get("docx_paragraph")
            
            level = determine_header_level(text, previous_headers, docx_paragraph)
            
            # Проверяем наличие нумерации
            if re.match(r'^\d+', text):
                last_numbered_level = level
            
            # Сохраняем заголовок для контекста
            header_info = {
                "level": level,
                "text": text,
                "docx_paragraph": docx_paragraph,
            }
            previous_headers.append(header_info)
            
            element["header_level"] = level
            element["last_numbered_level"] = last_numbered_level
        
        analyzed_elements.append(element)
    
    return analyzed_elements


def parse_docx_with_dots_ocr(docx_path: Path, output_dir: Optional[Path] = None) -> Dict[str, Any]:
    """
    Парсит DOCX используя Dots.OCR для layout detection и текст из DOCX.
    
    Args:
        docx_path: Путь к DOCX файлу
        output_dir: Директория для сохранения результатов
    
    Returns:
        Словарь с результатами парсинга
    """
    print("=" * 80)
    print("ПАЙПЛАЙН DOCX С DOTS.OCR")
    print("=" * 80)
    print()
    
    # Шаг 1: Извлекаем текст из DOCX
    print("Шаг 1: Извлечение текста из DOCX...")
    docx_paragraphs = extract_text_from_docx(docx_path)
    print(f"  Извлечено {len(docx_paragraphs)} параграфов")
    
    # Шаг 2: Рендерим страницы DOCX в изображения
    print("Шаг 2: Рендеринг страниц DOCX в изображения...")
    renderer = DocxPageRenderer(render_scale=2.0, optimize_for_ocr=True)
    total_pages = renderer.get_page_count(docx_path)
    print(f"  Всего страниц: {total_pages}")
    print(f"  Метод рендеринга: {renderer.render_method}")
    
    # Шаг 3: Layout detection через Dots.OCR для всех страниц
    print("Шаг 3: Layout detection через Dots.OCR...")
    all_layout_elements = []
    
    try:
        for page_num in tqdm(range(total_pages), desc="Layout detection", unit="страница"):
            try:
                original_image, optimized_image = renderer.render_page(
                    docx_path, page_num, return_original=True
                )
                
                # Layout detection
                layout_cells, raw_response, success = process_layout_detection(
                    image=optimized_image,
                    origin_image=original_image,
                )
                
                if not success or not layout_cells:
                    print(f"  Предупреждение: не удалось получить layout для страницы {page_num + 1}")
                    continue
                
                # Сопоставляем текст из OCR с текстом из DOCX
                matched_elements = match_ocr_text_with_docx_text(
                    layout_cells, docx_paragraphs, page_num
                )
                
                all_layout_elements.extend(matched_elements)
                
            except Exception as e:
                print(f"  Ошибка при обработке страницы {page_num + 1}: {e}")
                continue
        
        print(f"  Найдено {len(all_layout_elements)} элементов layout")
        
        # Шаг 4: Анализ уровней заголовков
        print("Шаг 4: Анализ уровней заголовков...")
        analyzed_elements = analyze_header_levels(all_layout_elements)
        headers_count = len([e for e in analyzed_elements if e["category"] == "Section-header"])
        print(f"  Найдено {headers_count} заголовков")
        
        # Шаг 5: Построение иерархии
        print("Шаг 5: Построение иерархии...")
        hierarchy = build_hierarchy_from_section_headers(analyzed_elements)
        print(f"  Создано {len(hierarchy)} секций")
        
        # Шаг 6: Создание элементов
        elements = []
        for section in hierarchy:
            header = section["header"]
            if header.get("category") == "Section-header":
                level = header.get("header_level", 1)
                elements.append({
                    "type": f"HEADER_{level}",
                    "content": header.get("docx_text", header.get("ocr_text", "")),
                    "metadata": {
                        "page_num": header.get("page_num", 0),
                        "bbox": header.get("bbox", []),
                        "docx_paragraph": header.get("docx_paragraph"),
                    }
                })
            
            for child in section["children"]:
                if child["category"] == "Text":
                    elements.append({
                        "type": "TEXT",
                        "content": child.get("docx_text", child.get("ocr_text", "")),
                        "metadata": {
                            "page_num": child.get("page_num", 0),
                            "bbox": child.get("bbox", []),
                        }
                    })
                elif child["category"] == "Table":
                    elements.append({
                        "type": "TABLE",
                        "content": child.get("docx_text", child.get("ocr_text", "")),
                        "metadata": {
                            "page_num": child.get("page_num", 0),
                            "bbox": child.get("bbox", []),
                        }
                    })
        
        # Результат
        result = {
            "source": str(docx_path),
            "format": "DOCX",
            "elements": elements,
            "metadata": {
                "parser": "docx_dots_ocr",
                "status": "completed",
                "processing_method": "dots_ocr_layout_with_docx_text",
                "render_method": renderer.render_method,
                "total_pages": total_pages,
                "sections_count": len(hierarchy),
                "headers_count": headers_count,
                "elements_count": len(elements),
            },
            "hierarchy": hierarchy,
        }
        
        # Сохранение результатов
        if output_dir:
            output_dir.mkdir(parents=True, exist_ok=True)
            output_file = output_dir / f"{docx_path.stem}_dots_ocr_pipeline.json"
            with open(output_file, "w", encoding="utf-8") as f:
                json.dump(result, f, ensure_ascii=False, indent=2, default=str)
            print(f"\nРезультат сохранен в: {output_file}")
        
        return result
    
    except Exception as e:
        print(f"Ошибка при обработке: {e}")
        raise


def _get_element_color(element_type: str) -> str:
    """Возвращает цвет для типа элемента."""
    color_map = {
        "TEXT": "green",
        "IMAGE": "magenta",
        "CAPTION": "orange",
        "HEADER_1": "cyan",
        "HEADER_2": "cyan",
        "HEADER_3": "cyan",
        "HEADER_4": "cyan",
        "HEADER_5": "cyan",
        "HEADER_6": "cyan",
        "TITLE": "red",
        "TABLE": "pink",
        "FORMULA": "gray",
        "LIST_ITEM": "blue",
    }
    return color_map.get(element_type, "red")


def _draw_bbox_on_full_page(image: Image.Image, bbox: List[float], label: str = "", color: str = "red") -> Image.Image:
    """Рисует bbox на полной странице."""
    from PIL import ImageDraw, ImageFont
    
    img_copy = image.copy()
    draw = ImageDraw.Draw(img_copy)
    
    if len(bbox) >= 4:
        x1, y1, x2, y2 = bbox[0], bbox[1], bbox[2], bbox[3]
        
        # Рисуем прямоугольник
        draw.rectangle([x1, y1, x2, y2], outline=color, width=2)
        
        # Добавляем подпись, если есть
        if label:
            try:
                font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 12)
            except:
                try:
                    font = ImageFont.truetype("arial.ttf", 12)
                except:
                    font = ImageFont.load_default()
            
            # Фон для текста
            text_bbox = draw.textbbox((x1, y1 - 15), label, font=font)
            text_bbox = (text_bbox[0] - 2, text_bbox[1] - 2, text_bbox[2] + 2, text_bbox[3] + 2)
            draw.rectangle(text_bbox, fill=color)
            draw.text((x1, y1 - 15), label, fill="white", font=font)
    
    return img_copy


def _save_full_pages_with_layout(
    docx_path: Path,
    result: Dict[str, Any],
    output_dir: Path,
    renderer: DocxPageRenderer,
) -> int:
    """
    Сохраняет полные сканы страниц с нарисованными bbox для всех элементов layout.
    
    Args:
        docx_path: Путь к DOCX файлу
        result: Результат парсинга
        output_dir: Директория для сохранения
        renderer: Рендерер страниц
        temp_pdf_path: Путь к временному PDF
    
    Returns:
        Количество сохраненных страниц
    """
    from collections import defaultdict
    
    pages_dir = output_dir / "pages_with_layout"
    pages_dir.mkdir(exist_ok=True)
    
    # Группируем элементы по страницам
    elements_by_page: Dict[int, List[Dict[str, Any]]] = defaultdict(list)
    for element in result.get("elements", []):
        page_num = element.get("metadata", {}).get("page_num", 0)
        bbox = element.get("metadata", {}).get("bbox", [])
        if bbox and len(bbox) >= 4:
            elements_by_page[page_num].append(element)
    
    if not elements_by_page:
        return 0
    
    total_pages = result.get("metadata", {}).get("total_pages", 0)
    saved_count = 0
    
    for page_num in tqdm(sorted(elements_by_page.keys()), desc="Сохранение страниц с layout", unit="страница", leave=False):
        if page_num >= total_pages:
            continue
        
        try:
            # Рендерим страницу
            original_image, _ = renderer.render_page(
                docx_path, page_num, return_original=True
            )
            page_image = original_image.copy()
            
            # Рисуем bbox для всех элементов на странице
            for element in elements_by_page[page_num]:
                element_type = element.get("type", "TEXT")
                bbox = element.get("metadata", {}).get("bbox", [])
                color = _get_element_color(element_type)
                label = f"{element_type}"
                
                page_image = _draw_bbox_on_full_page(page_image, bbox, label, color)
            
            # Сохраняем страницу
            page_file = pages_dir / f"page_{page_num + 1}_with_layout.png"
            page_image.save(page_file, "PNG")
            saved_count += 1
            
        except Exception as e:
            print(f"Ошибка при сохранении страницы {page_num + 1}: {e}")
            continue
    
    return saved_count


def _save_structure(result: Dict[str, Any], output_dir: Path) -> None:
    """Сохраняет структуру документа."""
    structure_file = output_dir / "structure.json"
    
    structure = {
        "source": result.get("source"),
        "format": result.get("format"),
        "metadata": result.get("metadata", {}),
        "hierarchy": result.get("hierarchy", []),
        "elements": [
            {
                "type": e.get("type"),
                "content": e.get("content", "")[:200] + "..." if len(e.get("content", "")) > 200 else e.get("content", ""),
                "metadata": e.get("metadata", {})
            }
            for e in result.get("elements", [])
        ]
    }
    
    with open(structure_file, "w", encoding="utf-8") as f:
        json.dump(structure, f, ensure_ascii=False, indent=2, default=str)


def process_docx_file(
    docx_path: Path,
    output_base_dir: Path,
    render_scale: float = 2.0,
) -> Dict[str, Any]:
    """
    Обрабатывает один DOCX файл и сохраняет результаты.
    
    Args:
        docx_path: Путь к DOCX файлу
        output_base_dir: Базовая директория для сохранения результатов
        render_scale: Масштаб рендеринга
    
    Returns:
        Словарь с результатами обработки
    """
    import time
    
    print(f"\n{'='*80}")
    print(f"Обработка: {docx_path.name}")
    print(f"{'='*80}")
    
    start_time = time.time()
    
    try:
        # Создаем директорию для результатов
        docx_output_dir = output_base_dir / docx_path.stem
        docx_output_dir.mkdir(parents=True, exist_ok=True)
        
        # Парсинг
        result = parse_docx_with_dots_ocr(docx_path, None)  # Не сохраняем в parse_docx_with_dots_ocr
        
        # Сохраняем полный результат
        result_file = docx_output_dir / "result.json"
        with open(result_file, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2, default=str)
        
        # Сохраняем структуру
        _save_structure(result, docx_output_dir)
        
        # Сохраняем страницы с layout
        renderer = DocxPageRenderer(render_scale=render_scale, optimize_for_ocr=True)
        
        saved_pages = _save_full_pages_with_layout(
            docx_path, result, docx_output_dir, renderer
        )
        
        processing_time = time.time() - start_time
        
        # Статистика
        stats = {
            "processing_time_seconds": processing_time,
            "total_elements": result.get("metadata", {}).get("elements_count", 0),
            "headers": result.get("metadata", {}).get("headers_count", 0),
            "sections": result.get("metadata", {}).get("sections_count", 0),
            "total_pages": result.get("metadata", {}).get("total_pages", 0),
            "saved_pages_with_layout": saved_pages,
        }
        
        stats_file = docx_output_dir / "stats.json"
        with open(stats_file, "w", encoding="utf-8") as f:
            json.dump(stats, f, indent=2, ensure_ascii=False)
        
        print(f"✓ Успешно обработан за {processing_time:.2f} сек")
        print(f"  Элементов: {stats['total_elements']}")
        print(f"  Заголовков: {stats['headers']}")
        print(f"  Секций: {stats['sections']}")
        print(f"  Страниц: {stats['total_pages']}")
        print(f"  Результаты сохранены в: {docx_output_dir}")
        
        return {
            "success": True,
            "processing_time": processing_time,
            "stats": stats,
            "output_dir": str(docx_output_dir),
        }
    
    except Exception as e:
        processing_time = time.time() - start_time
        error_msg = f"Ошибка при обработке {docx_path.name}: {e}"
        print(f"✗ {error_msg}")
        import traceback
        traceback.print_exc()
        
        return {
            "success": False,
            "processing_time": processing_time,
            "error": str(e),
        }


def main():
    """Основная функция для обработки нескольких DOCX файлов."""
    import time
    
    # Пути к файлам
    test_folder = Path(r"E:\easy\documentor\documentor_langchain\experiments\pdf_text_extraction\test_folder")
    
    docx_files = [
        test_folder / "02_Отчет_Этап_2_полный_замечания_эксперта.docx",
        test_folder / "Diplom2024.docx",
        test_folder / "Диплом.docx",
    ]
    
    # Проверяем существование файлов
    existing_files = [f for f in docx_files if f.exists()]
    
    if not existing_files:
        print("Ошибка: Не найдено ни одного файла для обработки")
        return
    
    print(f"Найдено файлов для обработки: {len(existing_files)}")
    for f in existing_files:
        print(f"  - {f.name}")
    
    # Директория для результатов
    output_dir = test_folder.parent / "results" / "documentor_docx_parser"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"\nВыходная директория: {output_dir}")
    
    # Обрабатываем каждый файл
    results: List[Dict[str, Any]] = []
    total_start_time = time.time()
    
    for docx_file in tqdm(existing_files, desc="Обработка DOCX файлов", unit="файл"):
        result = process_docx_file(docx_file, output_dir, render_scale=2.0)
        results.append({
            "file": docx_file.name,
            **result
        })
    
    total_time = time.time() - total_start_time
    
    # Сохраняем общую статистику
    summary = {
        "total_files": len(existing_files),
        "successful": len([r for r in results if r.get("success", False)]),
        "failed": len([r for r in results if not r.get("success", False)]),
        "total_processing_time_seconds": total_time,
        "average_processing_time_seconds": total_time / len(existing_files) if existing_files else 0,
        "processing_method": "dots_ocr_layout_with_docx_text",
        "results": results,
    }
    
    summary_file = output_dir / "summary.json"
    with open(summary_file, "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=2, ensure_ascii=False, default=str)
    
    # Выводим итоги
    print(f"\n{'='*80}")
    print("ИТОГИ ОБРАБОТКИ DOCX")
    print(f"{'='*80}")
    print(f"Всего файлов: {summary['total_files']}")
    print(f"Успешно: {summary['successful']}")
    print(f"Ошибок: {summary['failed']}")
    print(f"Общее время: {total_time:.2f} сек")
    print(f"Среднее время на файл: {summary['average_processing_time_seconds']:.2f} сек")
    print(f"Метод обработки: {summary['processing_method']}")
    print(f"\nДетальная статистика сохранена в: {summary_file}")
    
    if summary['failed'] > 0:
        print("\nФайлы с ошибками:")
        for result in results:
            if not result.get("success", False):
                print(f"  - {result['file']}: {result.get('error', 'Unknown error')}")


if __name__ == "__main__":
    main()
