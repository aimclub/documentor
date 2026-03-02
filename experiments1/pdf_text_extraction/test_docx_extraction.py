"""
Скрипт для обработки DOCX файлов с разными выводами:
- С метаданными
- Сплошной текст
- С координатами (если возможно)
- Постраничный вывод (по параграфам/секциям)
- Использование пайплайна documentor
"""

import json
from pathlib import Path
from typing import Dict, List, Any, Optional
import sys

# Добавляем корень проекта в путь
project_root = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(project_root))

try:
    from docx import Document as PythonDocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    print("Предупреждение: python-docx не установлен. Установите: pip install python-docx")

try:
    from langchain_core.documents import Document
    from documentor import Pipeline
    HAS_PIPELINE = True
except ImportError as e:
    HAS_PIPELINE = False
    print(f"Предупреждение: documentor Pipeline недоступен: {e}")
    print("Будет использоваться только python-docx для извлечения данных")


def extract_text_with_metadata(docx_path: Path) -> List[Dict[str, Any]]:
    """
    Извлекает текст с метаданными (стили, форматирование).
    
    Returns:
        Список словарей с текстом и метаданными для каждого параграфа.
    """
    if not HAS_PYTHON_DOCX:
        return []
    
    doc = PythonDocxDocument(str(docx_path))
    paragraphs_data = []
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.strip():
            continue
        
        # Извлечение стиля
        style_name = paragraph.style.name if paragraph.style else "Normal"
        
        # Извлечение форматирования
        formatting = {
            "bold": False,
            "italic": False,
            "underline": False,
            "font_size": None,
            "font_name": None,
            "alignment": None,
        }
        
        # Проверка выравнивания
        if paragraph.alignment:
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: "left",
                WD_ALIGN_PARAGRAPH.CENTER: "center",
                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
            }
            formatting["alignment"] = alignment_map.get(paragraph.alignment, "left")
        
        # Извлечение форматирования из runs
        runs_info = []
        for run in paragraph.runs:
            run_info = {
                "text": run.text,
                "bold": run.bold if run.bold is not None else False,
                "italic": run.italic if run.italic is not None else False,
                "underline": run.underline is not None,
            }
            
            if run.font.size:
                formatting["font_size"] = f"{run.font.size.pt}pt"
            if run.font.name:
                formatting["font_name"] = run.font.name
            
            if run.bold:
                formatting["bold"] = True
            if run.italic:
                formatting["italic"] = True
            if run.underline:
                formatting["underline"] = True
            
            runs_info.append(run_info)
        
        paragraphs_data.append({
            "index": para_idx,
            "text": paragraph.text,
            "style": style_name,
            "formatting": formatting,
            "runs": runs_info,
            "is_heading": style_name.startswith("Heading"),
        })
    
    return paragraphs_data


def extract_text_with_coordinates(docx_path: Path) -> List[Dict[str, Any]]:
    """
    Извлекает текст с координатами через XML разметку DOCX.
    
    Returns:
        Список словарей с текстом и координатами (если доступны).
    """
    if not HAS_PYTHON_DOCX:
        return []
    
    doc = PythonDocxDocument(str(docx_path))
    paragraphs_data = []
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.strip():
            continue
        
        # Попытка извлечь координаты из XML
        coordinates = None
        try:
            # Получаем XML элемент параграфа
            para_xml = paragraph._element
            
            # Ищем информацию о позиции в XML
            # В DOCX координаты обычно не хранятся напрямую, но можно попробовать
            # извлечь информацию о позиции через свойства параграфа
            pPr = para_xml.find(qn('w:pPr'))
            
            # Информация о позиции (отступы)
            position_info = {}
            if pPr is not None:
                # Отступ слева
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    left_indent = ind.get(qn('w:left'))
                    if left_indent:
                        position_info["left_indent"] = left_indent
                
                # Отступ справа
                if ind is not None:
                    right_indent = ind.get(qn('w:right'))
                    if right_indent:
                        position_info["right_indent"] = right_indent
                
                # Отступ первой строки
                if ind is not None:
                    first_line = ind.get(qn('w:firstLine'))
                    if first_line:
                        position_info["first_line_indent"] = first_line
            
            if position_info:
                coordinates = position_info
        except Exception as e:
            # Если не удалось извлечь координаты, оставляем None
            pass
        
        paragraphs_data.append({
            "index": para_idx,
            "text": paragraph.text,
            "coordinates": coordinates,
        })
    
    return paragraphs_data


def extract_text_by_sections(docx_path: Path) -> List[Dict[str, Any]]:
    """
    Извлекает текст по секциям (аналог постраничного вывода для DOCX).
    
    Returns:
        Список словарей с текстом по секциям.
    """
    if not HAS_PYTHON_DOCX:
        return []
    
    doc = PythonDocxDocument(str(docx_path))
    sections_data = []
    
    # Группируем параграфы по секциям (разделяем по заголовкам)
    current_section = {
        "section_index": 0,
        "title": "Начало документа",
        "paragraphs": [],
    }
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.strip():
            continue
        
        style_name = paragraph.style.name if paragraph.style else "Normal"
        
        # Если это заголовок, начинаем новую секцию
        if style_name.startswith("Heading"):
            # Сохраняем предыдущую секцию
            if current_section["paragraphs"]:
                sections_data.append(current_section.copy())
            
            # Создаем новую секцию
            current_section = {
                "section_index": len(sections_data),
                "title": paragraph.text,
                "style": style_name,
                "paragraphs": [],
            }
        else:
            # Добавляем параграф в текущую секцию
            current_section["paragraphs"].append({
                "index": para_idx,
                "text": paragraph.text,
                "style": style_name,
            })
    
    # Добавляем последнюю секцию
    if current_section["paragraphs"]:
        sections_data.append(current_section)
    
    return sections_data


def extract_plain_text(docx_path: Path) -> str:
    """
    Извлекает сплошной текст без метаданных.
    
    Returns:
        Строка с полным текстом документа.
    """
    if not HAS_PYTHON_DOCX:
        return ""
    
    doc = PythonDocxDocument(str(docx_path))
    full_text = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)
    
    return "\n".join(full_text)


def extract_tables_info(docx_path: Path) -> List[Dict[str, Any]]:
    """
    Извлекает информацию о таблицах.
    
    Returns:
        Список словарей с информацией о таблицах.
    """
    if not HAS_PYTHON_DOCX:
        return []
    
    doc = PythonDocxDocument(str(docx_path))
    tables_data = []
    
    for table_idx, table in enumerate(doc.tables):
        table_info = {
            "index": table_idx,
            "rows": len(table.rows),
            "columns": len(table.columns) if table.rows else 0,
            "data": [],
        }
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_info["data"].append(row_data)
        
        tables_data.append(table_info)
    
    return tables_data


def format_output_metadata(paragraphs_data: List[Dict[str, Any]]) -> str:
    """Форматирует вывод с метаданными."""
    output = []
    output.append("=" * 80)
    output.append("ТЕКСТ С МЕТАДАННЫМИ")
    output.append("=" * 80)
    output.append("")
    
    for para in paragraphs_data:
        output.append(f"[Параграф {para['index']}]")
        output.append(f"Стиль: {para['style']}")
        output.append(f"Заголовок: {'Да' if para['is_heading'] else 'Нет'}")
        
        formatting = para['formatting']
        output.append(f"Форматирование:")
        output.append(f"  - Выравнивание: {formatting['alignment'] or 'left'}")
        output.append(f"  - Жирный: {'Да' if formatting['bold'] else 'Нет'}")
        output.append(f"  - Курсив: {'Да' if formatting['italic'] else 'Нет'}")
        output.append(f"  - Подчеркивание: {'Да' if formatting['underline'] else 'Нет'}")
        if formatting['font_size']:
            output.append(f"  - Размер шрифта: {formatting['font_size']}")
        if formatting['font_name']:
            output.append(f"  - Шрифт: {formatting['font_name']}")
        
        output.append(f"Текст: {para['text']}")
        output.append("")
        output.append("-" * 80)
        output.append("")
    
    return "\n".join(output)


def format_output_coordinates(paragraphs_data: List[Dict[str, Any]]) -> str:
    """Форматирует вывод с координатами."""
    output = []
    output.append("=" * 80)
    output.append("ТЕКСТ С КООРДИНАТАМИ")
    output.append("=" * 80)
    output.append("")
    
    for para in paragraphs_data:
        output.append(f"[Параграф {para['index']}]")
        output.append(f"Текст: {para['text']}")
        
        if para['coordinates']:
            output.append("Координаты/Позиция:")
            for key, value in para['coordinates'].items():
                output.append(f"  - {key}: {value}")
        else:
            output.append("Координаты: Недоступны")
        
        output.append("")
        output.append("-" * 80)
        output.append("")
    
    return "\n".join(output)


def format_output_sections(sections_data: List[Dict[str, Any]]) -> str:
    """Форматирует вывод по секциям."""
    output = []
    output.append("=" * 80)
    output.append("ТЕКСТ ПО СЕКЦИЯМ (АНАЛОГ ПОСТРАНИЧНОГО ВЫВОДА)")
    output.append("=" * 80)
    output.append("")
    
    for section in sections_data:
        output.append(f"СЕКЦИЯ {section['section_index']}: {section['title']}")
        if 'style' in section:
            output.append(f"Стиль заголовка: {section['style']}")
        output.append("")
        output.append("Параграфы:")
        output.append("-" * 80)
        
        for para in section['paragraphs']:
            output.append(f"[Параграф {para['index']}] Стиль: {para['style']}")
            output.append(para['text'])
            output.append("")
        
        output.append("=" * 80)
        output.append("")
    
    return "\n".join(output)


def format_output_pipeline(parsed_doc) -> str:
    """Форматирует вывод из пайплайна."""
    output = []
    output.append("=" * 80)
    output.append("ВЫВОД ИЗ ПАЙПЛАЙНА DOCUMENTOR")
    output.append("=" * 80)
    output.append("")
    
    output.append(f"Источник: {parsed_doc.source}")
    output.append(f"Формат: {parsed_doc.format.value}")
    output.append(f"Количество элементов: {len(parsed_doc.elements)}")
    output.append("")
    
    if parsed_doc.metadata:
        output.append("Метаданные:")
        for key, value in parsed_doc.metadata.items():
            if key == "pipeline_metrics":
                output.append(f"  {key}:")
                for metric_key, metric_value in value.items():
                    output.append(f"    - {metric_key}: {metric_value}")
            else:
                output.append(f"  - {key}: {value}")
        output.append("")
    
    output.append("Элементы:")
    output.append("-" * 80)
    
    for idx, element in enumerate(parsed_doc.elements):
        output.append(f"[Элемент {idx}]")
        output.append(f"Тип: {element.type.value}")
        output.append(f"Контент: {element.content[:200]}..." if len(element.content) > 200 else f"Контент: {element.content}")
        
        if element.metadata:
            output.append("Метаданные элемента:")
            for key, value in element.metadata.items():
                output.append(f"  - {key}: {value}")
        
        output.append("")
        output.append("-" * 80)
        output.append("")
    
    return "\n".join(output)


def main():
    """Основная функция."""
    # Путь к документу
    docx_path = Path(r"E:\easy\documentor\documentor_langchain\experiments\pdf_text_extraction\test_folder\Диплом.docx")
    
    if not docx_path.exists():
        print(f"Ошибка: Файл не найден: {docx_path}")
        return
    
    # Создаем директорию для результатов
    output_dir = docx_path.parent / "docx_extraction_results"
    output_dir.mkdir(exist_ok=True)
    
    print(f"Обработка документа: {docx_path.name}")
    print(f"Выходная директория: {output_dir}")
    print()
    
    # 1. Сплошной текст
    print("1. Извлечение сплошного текста...")
    plain_text = extract_plain_text(docx_path)
    plain_text_file = output_dir / f"{docx_path.stem}_plain_text.txt"
    with open(plain_text_file, "w", encoding="utf-8") as f:
        f.write(plain_text)
    print(f"   Сохранено: {plain_text_file}")
    print()
    
    # 2. Текст с метаданными
    print("2. Извлечение текста с метаданными...")
    paragraphs_metadata = extract_text_with_metadata(docx_path)
    metadata_output = format_output_metadata(paragraphs_metadata)
    metadata_file = output_dir / f"{docx_path.stem}_with_metadata.txt"
    with open(metadata_file, "w", encoding="utf-8") as f:
        f.write(metadata_output)
    print(f"   Сохранено: {metadata_file}")
    print()
    
    # 3. Текст с координатами
    print("3. Извлечение текста с координатами...")
    paragraphs_coords = extract_text_with_coordinates(docx_path)
    coords_output = format_output_coordinates(paragraphs_coords)
    coords_file = output_dir / f"{docx_path.stem}_with_coordinates.txt"
    with open(coords_file, "w", encoding="utf-8") as f:
        f.write(coords_output)
    print(f"   Сохранено: {coords_file}")
    print()
    
    # 4. Постраничный вывод (по секциям)
    print("4. Извлечение текста по секциям...")
    sections_data = extract_text_by_sections(docx_path)
    sections_output = format_output_sections(sections_data)
    sections_file = output_dir / f"{docx_path.stem}_by_sections.txt"
    with open(sections_file, "w", encoding="utf-8") as f:
        f.write(sections_output)
    print(f"   Сохранено: {sections_file}")
    print()
    
    # 5. Информация о таблицах
    print("5. Извлечение информации о таблицах...")
    tables_info = extract_tables_info(docx_path)
    tables_file = output_dir / f"{docx_path.stem}_tables.txt"
    with open(tables_file, "w", encoding="utf-8") as f:
        f.write("=" * 80 + "\n")
        f.write("ИНФОРМАЦИЯ О ТАБЛИЦАХ\n")
        f.write("=" * 80 + "\n\n")
        for table in tables_info:
            f.write(f"Таблица {table['index']}:\n")
            f.write(f"  Строк: {table['rows']}\n")
            f.write(f"  Столбцов: {table['columns']}\n")
            f.write("  Данные:\n")
            for row_idx, row_data in enumerate(table['data']):
                f.write(f"    Строка {row_idx}: {' | '.join(row_data)}\n")
            f.write("\n" + "-" * 80 + "\n\n")
    print(f"   Сохранено: {tables_file}")
    print()
    
    # 6. Использование пайплайна documentor (опционально)
    if HAS_PIPELINE:
        print("6. Обработка через пайплайн documentor...")
        try:
            # Загружаем документ как LangChain Document
            # Для DOCX нужно сначала извлечь текст
            if HAS_PYTHON_DOCX:
                doc = PythonDocxDocument(str(docx_path))
                full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            else:
                full_text = plain_text
            
            langchain_doc = Document(
                page_content=full_text,
                metadata={"source": str(docx_path)}
            )
            
            # Обрабатываем через пайплайн
            pipeline = Pipeline()
            parsed_doc = pipeline.parse(langchain_doc)
            
            # Форматируем вывод
            pipeline_output = format_output_pipeline(parsed_doc)
            pipeline_file = output_dir / f"{docx_path.stem}_pipeline_output.txt"
            with open(pipeline_file, "w", encoding="utf-8") as f:
                f.write(pipeline_output)
            print(f"   Сохранено: {pipeline_file}")
        except Exception as e:
            print(f"   Ошибка при обработке через пайплайн: {e}")
            import traceback
            traceback.print_exc()
        print()
    else:
        print("6. Пропущено: пайплайн documentor недоступен")
        print()
    
    # 7. JSON с полными данными
    print("7. Сохранение полных данных в JSON...")
    full_data = {
        "source": str(docx_path),
        "plain_text": plain_text,
        "paragraphs_with_metadata": paragraphs_metadata,
        "paragraphs_with_coordinates": paragraphs_coords,
        "sections": sections_data,
        "tables": tables_info,
    }
    json_file = output_dir / f"{docx_path.stem}_full_data.json"
    with open(json_file, "w", encoding="utf-8") as f:
        json.dump(full_data, f, ensure_ascii=False, indent=2)
    print(f"   Сохранено: {json_file}")
    print()
    
    print("=" * 80)
    print("ОБРАБОТКА ЗАВЕРШЕНА")
    print("=" * 80)
    print(f"Все файлы сохранены в: {output_dir}")


if __name__ == "__main__":
    main()
