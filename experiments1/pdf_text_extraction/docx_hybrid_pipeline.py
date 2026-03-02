"""
Комбинированный пайплайн для DOCX: PDF для координат + DOCX для точного текста и метаданных.

Подход:
1. DOCX → PDF (всегда, для получения координат и изображений страниц)
2. Используем PDF для:
   - Рендеринга страниц в изображения для Dots.OCR
   - Извлечения текста по координатам (как в PDF пайплайне)
   - Получения координат элементов
3. Используем DOCX для:
   - Сравнения текста из PDF с текстом из DOCX (берем точный текст из DOCX)
   - Получения метаданных (стили, форматирование)
   - Извлечения таблиц (готовые структурированные данные)
   - Извлечения изображений (оригинальные размеры)
4. Каскадная обработка:
   - Dots.OCR находит элементы по координатам
   - Для текста: извлекаем из PDF по координатам, сравниваем с DOCX, берем точный текст из DOCX
   - Для таблиц: если Dots.OCR нашел таблицу, ищем её в DOCX и берем готовую структуру
   - Для изображений: если Dots.OCR нашел изображение, ищем похожее в DOCX и сохраняем оригинал

ПАЙПЛАЙН СРАВНЕНИЯ ИЗОБРАЖЕНИЙ:
================================

Для точного сопоставления изображений из OCR (низкокачественных/отсканированных) 
и из DOCX (оригинальных) при разных размерах и произвольном порядке используется
комбинированный метод на основе анализа ключевых точек с геометрической верификацией.

Этапы пайплайна сравнения изображений:

1. МЕТОД 1: ORB + RANSAC (приоритет, наиболее точный)
   --------------------------------------------
   Алгоритм:
   a) Детекция ключевых точек через ORB (Oriented FAST and Rotated BRIEF)
      - Быстрый и патентно-свободный алгоритм
      - Находит характерные точки на изображении
      - Создает дескрипторы для каждой точки
   
   b) Матчинг дескрипторов через Brute Force Matcher
      - Сравнивает дескрипторы между двумя изображениями
      - Находит потенциальные совпадения точек
   
   c) Геометрическая верификация через RANSAC
      - Ищет гомографию (геометрическое преобразование) между изображениями
      - Отфильтровывает ложные совпадения (outliers)
      - Оставляет только геометрически согласованные совпадения (inliers)
   
   d) Критерии точного совпадения:
      - Минимум inliers >= 15 (настраиваемо)
      - Доля inliers >= 70% от общего числа матчей (настраиваемо)
      - Гомография не вырождена (determinant > 0.01)
   
   Преимущества:
   - Устойчивость к изменению размера (гомография моделирует масштабирование)
   - Устойчивость к поворотам и перспективным искажениям
   - Защита от ложных совпадений (геометрическая согласованность)
   - Точность >99.9% при корректной настройке
   
   Недостатки:
   - Требует OpenCV
   - Медленнее, чем perceptual hash (но все еще быстро - миллисекунды)

2. МЕТОД 2: Perceptual Hash (fallback, быстрый)
   --------------------------------------------
   Алгоритм:
   a) Вычисление perceptual hash (pHash, aHash, dHash)
      - pHash: воспринимает визуальное сходство
      - aHash: средний hash (быстрый)
      - dHash: difference hash (устойчив к яркости)
   
   b) Сравнение через Hamming distance
      - Чем меньше расстояние, тем более похожи изображения
      - Порог: distance <= 15 (настраиваемо)
   
   Преимущества:
   - Очень быстрый
   - Не требует OpenCV
   - Хорошо работает для визуально похожих изображений
   
   Недостатки:
   - Менее точный, чем ORB+RANSAC
   - Может давать ложные совпадения для похожих, но разных изображений

3. МЕТОД 3: Размер (только если другие не сработали)
   --------------------------------------------
   - Сравнение по размеру изображения
   - Используется как последний fallback
   - Низкая точность, но может помочь в простых случаях

ПРИОРИТЕТЫ СРАВНЕНИЯ:
---------------------
1. ORB + RANSAC match → match_type = "orb_exact", confidence = "high"
2. MD5 exact match → match_type = "exact", confidence = "high"
3. Perceptual hash match → match_type = "visual", confidence = "medium"
4. Size similarity → match_type = "size", confidence = "low"

ДОПОЛНИТЕЛЬНЫЕ УЛУЧШЕНИЯ:
-------------------------
- Отслеживание порядка изображений в документе (order_in_document)
- Группировка по страницам для лучшего сопоставления
- Флаг "matched" для предотвращения повторного использования
- Бонус за порядок: если порядок изображения в DOCX близок к порядку на странице

ПАРАМЕТРЫ НАСТРОЙКИ:
-------------------
- orb_min_inliers: 15 (минимум совпадений для ORB)
- orb_inlier_ratio: 0.7 (минимальная доля inliers)
- perceptual_threshold: 15 (порог для perceptual hash)
- max_reproj_error: 3.0 (максимальная ошибка репроекции для RANSAC)

Для мелких изображений (<100px): уменьшить orb_min_inliers до 8-10
Для сканов с шумом: увеличить max_reproj_error до 5-6
"""

import json
import re
import sys
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple, Union
from io import BytesIO
import tempfile
import shutil
import hashlib
import base64
import zipfile
from xml.etree import ElementTree as ET

from PIL import Image
from tqdm import tqdm
from collections import defaultdict

# Для perceptual hashing (визуальное сравнение изображений)
try:
    import imagehash
    HAS_IMAGEHASH = True
except ImportError:
    HAS_IMAGEHASH = False
    print("Предупреждение: imagehash не установлен. Установите: pip install imagehash")

# Для сравнения изображений через ключевые точки (ORB + RANSAC)
try:
    import cv2
    import numpy as np
    HAS_OPENCV = True
except ImportError:
    HAS_OPENCV = False
    print("Предупреждение: opencv-python не установлен. Установите: pip install opencv-python")

# Добавляем корень проекта в путь
project_root = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(project_root))

try:
    from docx import Document as PythonDocxDocument
    from docx.oxml.ns import qn
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    print("Предупреждение: python-docx не установлен. Установите: pip install python-docx")

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    print("Предупреждение: PyMuPDF не установлен")

# Методы конвертации DOCX → PDF
HAS_DOCX2PDF = False
try:
    from docx2pdf import convert as docx2pdf_convert
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False

HAS_WIN32COM = False
try:
    import win32com.client
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False

HAS_LIBREOFFICE = False
LIBREOFFICE_CMD = None
try:
    import subprocess
    for cmd in ['soffice', 'libreoffice', '/Applications/LibreOffice.app/Contents/MacOS/soffice']:
        try:
            result = subprocess.run([cmd, '--version'], capture_output=True, timeout=5)
            if result.returncode == 0:
                HAS_LIBREOFFICE = True
                LIBREOFFICE_CMD = cmd
                break
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
except:
    pass

# Импорты из documentor
from documentor.processing.parsers.pdf.ocr.dots_ocr_client import process_layout_detection
from documentor.processing.parsers.pdf.ocr.page_renderer import PdfPageRenderer
from documentor.utils.ocr_image_utils import fetch_image
from documentor.utils.ocr_consts import MIN_PIXELS, MAX_PIXELS

# Константы
RENDER_SCALE = 3.0  # Масштаб рендеринга для всех операций (layout detection, извлечение изображений, сохранение страниц)


def convert_docx_to_pdf(docx_path: Path, temp_pdf_path: Path) -> None:
    """
    Конвертирует DOCX в PDF с сохранением изображений.
    
    Использует приоритет:
    1. win32com (Word) - лучшая поддержка изображений на Windows
    2. LibreOffice - кроссплатформенный вариант
    3. docx2pdf - fallback
    
    Args:
        docx_path: Путь к DOCX файлу
        temp_pdf_path: Путь для сохранения PDF
    """
    # Приоритет 1: win32com (Word) - лучшая поддержка изображений
    if HAS_WIN32COM:
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False  # Отключаем предупреждения
            
            try:
                doc = word.Documents.Open(str(docx_path.absolute()))
                try:
                    # Экспортируем в PDF с сохранением всех элементов, включая изображения
                    # wdExportFormatPDF = 17
                    # wdExportOptimizeForPrint = 0 (для лучшего качества изображений)
                    doc.ExportAsFixedFormat(
                        OutputFileName=str(temp_pdf_path),
                        ExportFormat=17,  # wdExportFormatPDF
                        OpenAfterExport=False,
                        OptimizeFor=0,  # wdExportOptimizeForPrint (лучшее качество)
                        BitmapMissingFonts=True,  # Важно для изображений
                        UseISO19005_1=False,
                        IncludeDocProps=False,
                        KeepIRM=False,
                        CreateBookmarks=0,
                        DocStructureTags=False,
                    )
                    
                    if temp_pdf_path.exists():
                        print(f"  ✓ PDF создан через Word COM: {temp_pdf_path}")
                        return
                finally:
                    doc.Close(SaveChanges=False)
            finally:
                word.Quit()
        except Exception as e:
            print(f"  Предупреждение: не удалось конвертировать через Word COM: {e}")
            print(f"  Пробуем альтернативный метод...")
    
    # Приоритет 2: LibreOffice - хорошая поддержка изображений
    if HAS_LIBREOFFICE:
        try:
            cmd = LIBREOFFICE_CMD or 'soffice'
            result = subprocess.run(
                [
                    cmd,
                    '--headless',
                    '--nodefault',
                    '--nolockcheck',
                    '--invisible',
                    '--convert-to', 'pdf',
                    '--outdir', str(temp_pdf_path.parent),
                    str(docx_path.absolute())
                ],
                check=True,
                timeout=120,
                capture_output=True,
                text=True
            )
            
            # LibreOffice создает PDF с именем файла, нужно переименовать
            expected_pdf = temp_pdf_path.parent / f"{docx_path.stem}.pdf"
            if expected_pdf.exists() and expected_pdf != temp_pdf_path:
                expected_pdf.rename(temp_pdf_path)
            
            if temp_pdf_path.exists():
                print(f"  ✓ PDF создан через LibreOffice: {temp_pdf_path}")
                return
        except Exception as e:
            print(f"  Предупреждение: не удалось конвертировать через LibreOffice: {e}")
            print(f"  Пробуем fallback метод...")
    
    # Приоритет 3: docx2pdf (fallback)
    if HAS_DOCX2PDF:
        try:
            docx2pdf_convert(str(docx_path), str(temp_pdf_path))
            if temp_pdf_path.exists():
                print(f"  ✓ PDF создан через docx2pdf: {temp_pdf_path}")
                return
        except Exception as e:
            print(f"  Предупреждение: не удалось конвертировать через docx2pdf: {e}")
    
    # Если ничего не сработало
    raise RuntimeError(
        "Не удалось конвертировать DOCX в PDF. Установите один из:\n"
        "- Microsoft Word (Windows) + pywin32\n"
        "- LibreOffice (кроссплатформенный)\n"
        "- docx2pdf (pip install docx2pdf)"
    )


def extract_text_from_pdf_by_bbox(pdf_path: Path, bbox: List[float], page_num: int, render_scale: float = None) -> str:
    """
    Извлекает текст из PDF по координатам bbox (как в PDF пайплайне).
    
    Args:
        pdf_path: Путь к PDF файлу
        bbox: Координаты [x1, y1, x2, y2]
        page_num: Номер страницы (0-based)
        render_scale: Масштаб рендеринга (если None, используется RENDER_SCALE)
    
    Returns:
        Извлеченный текст
    """
    if render_scale is None:
        render_scale = RENDER_SCALE
    
    if not HAS_PYMUPDF:
        return ""
    
    pdf_document = fitz.open(str(pdf_path))
    try:
        if page_num >= len(pdf_document):
            return ""
        
        page = pdf_document.load_page(page_num)
        # Приводим координаты к масштабу оригинального PDF
        x1, y1, x2, y2 = (
            bbox[0] / render_scale,
            bbox[1] / render_scale,
            bbox[2] / render_scale,
            bbox[3] / render_scale,
        )
        rect = fitz.Rect(x1, y1, x2, y2)
        text = page.get_text("text", clip=rect).strip()
        return text
    finally:
        pdf_document.close()


def extract_all_text_from_docx(docx_path: Path) -> List[Dict[str, Any]]:
    """
    Извлекает весь текст из DOCX с метаданными.
    
    Args:
        docx_path: Путь к DOCX файлу
    
    Returns:
        Список параграфов с текстом и метаданными
    """
    if not HAS_PYTHON_DOCX:
        raise RuntimeError("python-docx не установлен")
    
    doc = PythonDocxDocument(str(docx_path))
    paragraphs_data = []
    
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        style = para.style.name if para.style else "Normal"
        
        # Форматирование
        formatting = {}
        if para.runs:
            first_run = para.runs[0]
            formatting["bold"] = first_run.bold or False
            formatting["italic"] = first_run.italic or False
            if first_run.font.size:
                formatting["font_size"] = str(first_run.font.size.pt) + "pt"
            if first_run.font.name:
                formatting["font_name"] = first_run.font.name
        
        paragraphs_data.append({
            "index": idx,
            "text": text,
            "style": style,
            "formatting": formatting,
            "is_heading": style.startswith("Heading"),
        })
    
    return paragraphs_data


def extract_tables_from_docx_xml(docx_path: Path) -> List[Dict[str, Any]]:
    """
    Извлекает все таблицы из DOCX через XML парсинг.
    Это дает больше информации о структуре, включая объединенные ячейки.
    Также пытается определить приблизительный номер страницы на основе разрывов страниц.
    
    Args:
        docx_path: Путь к DOCX файлу
    
    Returns:
        Список таблиц с данными и метаданными из XML, включая estimated_page
    """
    tables_data = []
    
    try:
        # DOCX - это ZIP архив, открываем его
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # Читаем document.xml
            document_xml = docx_zip.read('word/document.xml')
            
        # Парсим XML
        root = ET.fromstring(document_xml)
        
        # Определяем namespace для Word
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # Находим body документа
        body = root.find('w:body', ns)
        if body is None:
            return []
        
        # Собираем все элементы body в порядке появления
        all_elements = list(body)
        
        # Находим все разрывы страниц для приблизительного определения страниц
        page_breaks = []
        current_page = 1  # Начинаем с первой страницы
        
        # Проходим по всем элементам и находим разрывы страниц
        for elem_idx, elem in enumerate(all_elements):
            # Ищем разрывы страниц в параграфах
            if elem.tag.endswith('}p'):  # Параграф
                # Ищем w:br с type="page"
                for br in elem.findall('.//w:br', ns):
                    br_type = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') or br.get('type')
                    if br_type == 'page':
                        page_breaks.append({
                            'element_index': elem_idx,
                            'page': current_page + 1,
                        })
                        current_page += 1
                
                # Ищем w:lastRenderedPageBreak (разрыв страницы при последнем рендеринге)
                for page_break in elem.findall('.//w:lastRenderedPageBreak', ns):
                    page_breaks.append({
                        'element_index': elem_idx,
                        'page': current_page + 1,
                    })
                    current_page += 1
        
        # Функция для определения приблизительной страницы элемента
        def get_estimated_page(element_index: int) -> int:
            """Определяет приблизительный номер страницы на основе разрывов страниц."""
            if not page_breaks:
                # Если нет разрывов страниц, используем эвристику: ~50 элементов на страницу
                return max(1, (element_index // 50) + 1)
            
            # Находим последний разрыв страницы перед этим элементом
            estimated_page = 1
            for pb in page_breaks:
                if pb['element_index'] <= element_index:
                    estimated_page = pb['page']
                else:
                    break
            
            return estimated_page
        
        # Находим все таблицы
        tables = root.findall('.//w:tbl', ns)
        
        for table_idx, table_elem in enumerate(tables):
            # Находим позицию таблицы в списке всех элементов
            table_xml_position = None
            try:
                table_xml_position = all_elements.index(table_elem)
            except ValueError:
                # Если таблица не найдена напрямую, ищем её родительский элемент
                parent = table_elem.getparent()
                if parent is not None:
                    try:
                        table_xml_position = all_elements.index(parent)
                    except ValueError:
                        pass
            
            # Определяем приблизительный номер страницы
            estimated_page = 1
            if table_xml_position is not None:
                estimated_page = get_estimated_page(table_xml_position)
            
            table_info = {
                "index": table_idx,
                "rows": [],
                "rows_count": 0,
                "cols_count": 0,
                "style": None,
                "merged_cells": [],  # Информация об объединенных ячейках
                "xml_position": table_xml_position,  # Позиция в XML (для сопоставления)
                "estimated_page": estimated_page,  # Приблизительный номер страницы
                "page_breaks_before": len([pb for pb in page_breaks if pb['element_index'] < (table_xml_position or 0)]),
            }
            
            # Получаем стиль таблицы
            tbl_pr = table_elem.find('w:tblPr', ns)
            if tbl_pr is not None:
                tbl_style = tbl_pr.find('w:tblStyle', ns)
                if tbl_style is not None:
                    # Атрибут val - просто имя без namespace
                    table_info["style"] = tbl_style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') or tbl_style.get('val', '')
            
            # Находим все строки
            rows = table_elem.findall('.//w:tr', ns)
            table_info["rows_count"] = len(rows)
            
            max_cols = 0
            row_idx = 0
            
            for row_elem in rows:
                row_data = {
                    "row_index": row_idx,
                    "cells": [],
                    "cells_count": 0,
                }
                
                # Находим все ячейки в строке
                cells = row_elem.findall('.//w:tc', ns)
                
                col_idx = 0
                for cell_elem in cells:
                    # Извлекаем текст из ячейки
                    cell_text_parts = []
                    for para in cell_elem.findall('.//w:p', ns):
                        for run in para.findall('.//w:r', ns):
                            for text_elem in run.findall('.//w:t', ns):
                                cell_text_parts.append(text_elem.text or '')
                    
                    cell_text = ''.join(cell_text_parts).strip()
                    
                    # Проверяем объединенные ячейки
                    cell_props = cell_elem.find('w:tcPr', ns)
                    is_merged = False
                    rowspan = None
                    colspan = None
                    vmerge = None
                    
                    if cell_props is not None:
                        # Проверяем colspan (объединение по горизонтали)
                        grid_span = cell_props.find('w:gridSpan', ns)
                        if grid_span is not None:
                            # Атрибут val не имеет namespace
                            val_attr = grid_span.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') or grid_span.get('val')
                            if val_attr:
                                colspan = int(val_attr)
                                is_merged = True
                        
                        # Проверяем vmerge (объединение по вертикали)
                        v_merge = cell_props.find('w:vMerge', ns)
                        if v_merge is not None:
                            # Атрибут val не имеет namespace
                            vmerge_attr = v_merge.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') or v_merge.get('val')
                            if vmerge_attr:
                                vmerge = vmerge_attr
                                if vmerge != 'restart':
                                    is_merged = True
                        
                        # Проверяем rowspan (объединение строк)
                        # В Word это обычно через vMerge
                    
                    cell_info = {
                        "cell_index": col_idx,
                        "text": cell_text,
                        "text_length": len(cell_text),
                        "is_merged": is_merged,
                        "colspan": colspan or 1,
                        "rowspan": rowspan or 1,
                        "vmerge": vmerge,
                    }
                    
                    row_data["cells"].append(cell_info)
                    
                    # Учитываем colspan при подсчете столбцов
                    actual_cols = colspan or 1
                    col_idx += actual_cols
                    max_cols = max(max_cols, col_idx)
                    
                    if is_merged:
                        table_info["merged_cells"].append({
                            "row": row_idx,
                            "col": col_idx - actual_cols,
                            "colspan": colspan,
                            "vmerge": vmerge,
                        })
                
                row_data["cells_count"] = len(row_data["cells"])
                table_info["rows"].append(row_data)
                row_idx += 1
            
            table_info["cols_count"] = max_cols
            
            # Создаем упрощенную структуру данных (для обратной совместимости)
            table_data = []
            for row_info in table_info["rows"]:
                row_cells = [cell["text"] for cell in row_info["cells"]]
                table_data.append(row_cells)
            
            table_info["data"] = table_data
            
            # Добавляем обратную совместимость: rows и cols для старого кода
            table_info["rows"] = table_info["rows_count"]
            table_info["cols"] = table_info["cols_count"]
            
            # xml_position и estimated_page уже установлены выше
            
            tables_data.append(table_info)
    
    except Exception as e:
        print(f"Ошибка при парсинге XML DOCX: {e}")
        import traceback
        traceback.print_exc()
        # Fallback на python-docx
        if HAS_PYTHON_DOCX:
            return extract_tables_from_docx_fallback(docx_path)
        return []
    
    return tables_data


def extract_tables_from_docx_fallback(docx_path: Path) -> List[Dict[str, Any]]:
    """
    Fallback метод через python-docx (если XML парсинг не сработал).
    """
    if not HAS_PYTHON_DOCX:
        return []
    
    doc = PythonDocxDocument(str(docx_path))
    tables_data = []
    
    for table_idx, table in enumerate(doc.tables):
        table_rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(row_cells)
        
        table_info = {
            "index": table_idx,
            "rows_count": len(table.rows),
            "cols_count": len(table.columns) if table.rows else 0,
            "data": table_rows,
        }
        # Обратная совместимость
        table_info["rows"] = table_info["rows_count"]
        table_info["cols"] = table_info["cols_count"]
        tables_data.append(table_info)
    
    return tables_data


def extract_tables_from_docx(docx_path: Path) -> List[Dict[str, Any]]:
    """
    Извлекает все таблицы из DOCX (использует XML парсинг).
    
    Args:
        docx_path: Путь к DOCX файлу
    
    Returns:
        Список таблиц с данными
    """
    return extract_tables_from_docx_xml(docx_path)


def extract_images_from_docx(docx_path: Path) -> List[Dict[str, Any]]:
    """
    Извлекает все изображения из DOCX с оригинальными размерами и порядком.
    
    Args:
        docx_path: Путь к DOCX файлу
    
    Returns:
        Список изображений с данными (в порядке появления в документе)
    """
    if not HAS_PYTHON_DOCX:
        return []
    
    doc = PythonDocxDocument(str(docx_path))
    images_data = []
    image_counter = 0  # Счетчик для отслеживания порядка
    
    # Извлекаем изображения из параграфов
    for para_idx, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if run._element.xpath('.//a:blip'):
                # Найдено изображение
                blip = run._element.xpath('.//a:blip')[0]
                rId = blip.get(qn('r:embed'))
                
                if rId:
                    try:
                        image_part = doc.part.related_parts[rId]
                        image_bytes = image_part.blob
                        
                        # Проверяем, что image_bytes не пустой
                        if not image_bytes or len(image_bytes) == 0:
                            print(f"Предупреждение: пустой image_bytes для изображения в параграфе {para_idx}")
                            continue
                        
                        # Получаем размеры изображения
                        image = Image.open(BytesIO(image_bytes))
                        
                        # Конвертируем в RGB если нужно
                        if image.mode != 'RGB':
                            image = image.convert('RGB')
                        
                        width, height = image.size
                        
                        # Проверяем, что размеры валидные
                        if width == 0 or height == 0:
                            print(f"Предупреждение: нулевые размеры изображения в параграфе {para_idx}")
                            continue
                        
                        # Сохраняем image_bytes заново из конвертированного изображения
                        img_bytes_io = BytesIO()
                        image.save(img_bytes_io, format='PNG')
                        image_bytes = img_bytes_io.getvalue()
                        
                        # Создаем хеш для сравнения
                        image_hash = hashlib.md5(image_bytes).hexdigest()
                        
                        images_data.append({
                            "index": image_counter,
                            "order_in_document": image_counter,  # Порядок в документе
                            "paragraph_index": para_idx,
                            "width": width,
                            "height": height,
                            "image_bytes": image_bytes,
                            "image_hash": image_hash,
                            "format": "PNG",  # Всегда сохраняем как PNG
                            "matched": False,  # Флаг для отслеживания сопоставления
                        })
                        image_counter += 1
                    except Exception as e:
                        print(f"Ошибка при извлечении изображения из параграфа {para_idx}: {e}")
                        import traceback
                        traceback.print_exc()
                        continue
    
    # Извлекаем изображения из таблиц
    for table_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run._element.xpath('.//a:blip'):
                            blip = run._element.xpath('.//a:blip')[0]
                            rId = blip.get(qn('r:embed'))
                            
                            if rId:
                                try:
                                    image_part = doc.part.related_parts[rId]
                                    image_bytes = image_part.blob
                                    
                                    # Проверяем, что image_bytes не пустой
                                    if not image_bytes or len(image_bytes) == 0:
                                        continue
                                    
                                    image = Image.open(BytesIO(image_bytes))
                                    
                                    # Конвертируем в RGB если нужно
                                    if image.mode != 'RGB':
                                        image = image.convert('RGB')
                                    
                                    width, height = image.size
                                    
                                    # Проверяем, что размеры валидные
                                    if width == 0 or height == 0:
                                        continue
                                    
                                    # Сохраняем image_bytes заново из конвертированного изображения
                                    img_bytes_io = BytesIO()
                                    image.save(img_bytes_io, format='PNG')
                                    image_bytes = img_bytes_io.getvalue()
                                    
                                    image_hash = hashlib.md5(image_bytes).hexdigest()
                                    
                                    images_data.append({
                                        "index": image_counter,
                                        "order_in_document": image_counter,  # Порядок в документе
                                        "table_index": table_idx,
                                        "width": width,
                                        "height": height,
                                        "image_bytes": image_bytes,
                                        "image_hash": image_hash,
                                        "format": "PNG",  # Всегда сохраняем как PNG
                                        "matched": False,  # Флаг для отслеживания сопоставления
                                    })
                                    image_counter += 1
                                except Exception as e:
                                    continue
    
    return images_data


def match_pdf_text_with_docx_text(
    pdf_text: str,
    docx_paragraphs: List[Dict[str, Any]],
    threshold: float = 0.7
) -> Optional[Dict[str, Any]]:
    """
    Сопоставляет текст из PDF с текстом из DOCX.
    
    Args:
        pdf_text: Текст из PDF
        docx_paragraphs: Список параграфов из DOCX
        threshold: Порог схожести (0-1)
    
    Returns:
        Соответствующий параграф из DOCX или None
    """
    if not pdf_text:
        return None
    
    pdf_text_lower = pdf_text.lower().strip()
    pdf_words = set(pdf_text_lower.split())
    
    best_match = None
    best_score = 0.0
    
    for para in docx_paragraphs:
        docx_text = para["text"].lower().strip()
        docx_words = set(docx_text.split())
        
        if not docx_words:
            continue
        
        # Вычисляем Jaccard similarity
        intersection = len(pdf_words & docx_words)
        union = len(pdf_words | docx_words)
        similarity = intersection / union if union > 0 else 0.0
        
        if similarity > best_score and similarity >= threshold:
            best_score = similarity
            best_match = para
    
    return best_match


def extract_table_text_from_pdf(
    pdf_path: Path,
    table_bbox: List[float],
    page_num: int,
    render_scale: float = None
) -> str:
    """
    Извлекает весь текст из области таблицы в PDF.
    
    Args:
        pdf_path: Путь к PDF файлу
        table_bbox: Координаты таблицы [x1, y1, x2, y2]
        page_num: Номер страницы (0-based)
        render_scale: Масштаб рендеринга
    
    Returns:
        Текст из таблицы (все ячейки, разделенные пробелами/переносами)
    """
    if render_scale is None:
        render_scale = RENDER_SCALE
    
    if not HAS_PYMUPDF:
        return ""
    
    try:
        pdf_document = fitz.open(str(pdf_path))
        if page_num >= len(pdf_document):
            return ""
        
        page = pdf_document.load_page(page_num)
        
        # Приводим координаты к масштабу оригинального PDF
        x1, y1, x2, y2 = (
            table_bbox[0] / render_scale,
            table_bbox[1] / render_scale,
            table_bbox[2] / render_scale,
            table_bbox[3] / render_scale,
        )
        rect = fitz.Rect(x1, y1, x2, y2)
        text = page.get_text("text", clip=rect).strip()
        pdf_document.close()
        return text
    except Exception as e:
        print(f"Ошибка при извлечении текста таблицы из PDF: {e}")
        return ""


def normalize_table_text(text: str) -> str:
    """
    Нормализует текст таблицы для сравнения.
    Удаляет лишние пробелы, приводит к нижнему регистру, удаляет спецсимволы.
    """
    if not text:
        return ""
    # Приводим к нижнему регистру
    text = text.lower()
    # Удаляем лишние пробелы и переносы строк
    text = " ".join(text.split())
    # Удаляем спецсимволы (оставляем только буквы, цифры и пробелы)
    text = re.sub(r'[^\w\s]', '', text)
    return text.strip()


def calculate_table_content_similarity(
    docx_table: Dict[str, Any],
    ocr_table_text: str,
    pdf_path: Optional[Path] = None,
    table_bbox: Optional[List[float]] = None,
    page_num: Optional[int] = None
) -> float:
    """
    Вычисляет схожесть содержимого таблицы из DOCX и текста из OCR.
    
    Args:
        docx_table: Таблица из DOCX
        ocr_table_text: Текст, извлеченный из области таблицы в PDF/OCR
        pdf_path: Путь к PDF (для дополнительного извлечения текста, если ocr_table_text пустой)
        table_bbox: Координаты таблицы (если нужно извлечь текст)
        page_num: Номер страницы (если нужно извлечь текст)
    
    Returns:
        Коэффициент схожести от 0.0 до 1.0
    """
    # Извлекаем текст из DOCX таблицы
    docx_text_parts = []
    if "data" in docx_table:
        for row in docx_table["data"]:
            if isinstance(row, list):
                for cell in row:
                    if cell:
                        docx_text_parts.append(str(cell))
            elif isinstance(row, dict) and "cells" in row:
                for cell in row["cells"]:
                    if isinstance(cell, dict) and "text" in cell:
                        docx_text_parts.append(str(cell["text"]))
    
    docx_text = " ".join(docx_text_parts)
    docx_text_normalized = normalize_table_text(docx_text)
    
    # Если текст из OCR не предоставлен, пытаемся извлечь
    if not ocr_table_text and pdf_path and table_bbox and page_num is not None:
        ocr_table_text = extract_table_text_from_pdf(pdf_path, table_bbox, page_num)
    
    ocr_text_normalized = normalize_table_text(ocr_table_text)
    
    if not docx_text_normalized and not ocr_text_normalized:
        return 1.0  # Обе таблицы пустые - считаем совпадением
    
    if not docx_text_normalized or not ocr_text_normalized:
        return 0.0  # Одна пустая, другая нет - не совпадают
    
    # Простое сравнение: доля общих слов
    docx_words = set(docx_text_normalized.split())
    ocr_words = set(ocr_text_normalized.split())
    
    if not docx_words or not ocr_words:
        return 0.0
    
    # Вычисляем Jaccard similarity (пересечение / объединение)
    intersection = len(docx_words & ocr_words)
    union = len(docx_words | ocr_words)
    
    if union == 0:
        return 0.0
    
    similarity = intersection / union
    
    # Дополнительно: проверяем, что хотя бы 50% слов совпадают
    docx_coverage = intersection / len(docx_words) if docx_words else 0.0
    ocr_coverage = intersection / len(ocr_words) if ocr_words else 0.0
    
    # Используем минимальное покрытие (обе таблицы должны иметь достаточно общих слов)
    min_coverage = min(docx_coverage, ocr_coverage)
    
    # Комбинируем Jaccard similarity и минимальное покрытие
    final_similarity = (similarity * 0.6 + min_coverage * 0.4)
    
    return final_similarity


def find_ocr_table_for_docx_table(
    docx_table: Dict[str, Any],
    ocr_table_elements: List[Dict[str, Any]],
    used_ocr_tables: set,
    docx_paragraphs: List[Dict[str, Any]],
    pdf_path: Optional[Path] = None
) -> Optional[Dict[str, Any]]:
    """
    Находит таблицу в результатах OCR для таблицы из DOCX.
    Использует комбинированный подход:
    1. Сравнение содержимого (текст из ячеек)
    2. estimated_page и порядок таблиц
    3. Позиция на странице
    
    Args:
        docx_table: Таблица из DOCX
        ocr_table_elements: Список таблиц из Dots.OCR
        used_ocr_tables: Множество ID уже использованных таблиц из OCR
        docx_paragraphs: Список параграфов из DOCX (для контекста)
        pdf_path: Путь к PDF файлу (для извлечения текста таблиц)
    
    Returns:
        Найденная таблица из OCR или None
    """
    if not ocr_table_elements:
        return None
    
    # Фильтруем неиспользованные таблицы
    available_ocr_tables = [
        t for t in ocr_table_elements 
        if id(t) not in used_ocr_tables
    ]
    
    if not available_ocr_tables:
        return None
    
    # Приоритет 1: Сравнение содержимого таблиц (если есть PDF)
    if pdf_path:
        best_match = None
        best_similarity = 0.0
        
        estimated_page = docx_table.get("estimated_page")
        estimated_page_0based = (estimated_page - 1) if estimated_page else None
        
        for ocr_table in available_ocr_tables:
            # Фильтруем по странице (если есть estimated_page)
            if estimated_page_0based is not None:
                ocr_page = ocr_table.get("page_num", 9999)
                if abs(ocr_page - estimated_page_0based) > 3:  # Слишком далеко
                    continue
            
            table_bbox = ocr_table.get("bbox", [])
            page_num = ocr_table.get("page_num", 0)
            
            if not table_bbox or len(table_bbox) < 4:
                continue
            
            # Извлекаем текст из области таблицы
            ocr_table_text = extract_table_text_from_pdf(pdf_path, table_bbox, page_num)
            
            # Вычисляем схожесть
            similarity = calculate_table_content_similarity(
                docx_table, ocr_table_text, pdf_path, table_bbox, page_num
            )
            
            if similarity > best_similarity:
                best_similarity = similarity
                best_match = ocr_table
        
        # Если нашли хорошее совпадение (similarity > 0.3), используем его
        if best_match and best_similarity > 0.3:
            return best_match
    
    # Приоритет 2: Используем estimated_page из DOCX для поиска в OCR
    estimated_page = docx_table.get("estimated_page")
    if estimated_page is not None:
        # Конвертируем в 0-based для сравнения с page_num из OCR
        estimated_page_0based = estimated_page - 1
        
        # Ищем таблицы на той же странице или близкой
        tables_on_similar_page = [
            t for t in available_ocr_tables
            if abs(t.get("page_num", 9999) - estimated_page_0based) <= 2
        ]
        
        if tables_on_similar_page:
            # Сортируем по близости к estimated_page, затем по позиции Y
            tables_on_similar_page.sort(
                key=lambda t: (
                    abs(t.get("page_num", 9999) - estimated_page_0based),
                    t.get("bbox", [0])[1] if t.get("bbox") else 0
                )
            )
            return tables_on_similar_page[0]
    
    # Приоритет 3: Используем порядок таблиц (xml_position)
    xml_position = docx_table.get("xml_position")
    if xml_position is not None:
        # Сортируем таблицы из OCR по порядку появления (по page_num и позиции на странице)
        available_ocr_tables.sort(key=lambda t: (t.get("page_num", 9999), t.get("bbox", [0])[1] if t.get("bbox") else 0))
        
        # Используем индекс таблицы в DOCX как приблизительный порядок
        docx_table_index = docx_table.get("index", 0)
        
        if docx_table_index < len(available_ocr_tables):
            return available_ocr_tables[docx_table_index]
    
    # Fallback: возвращаем первую доступную таблицу
    return available_ocr_tables[0] if available_ocr_tables else None


def find_table_in_docx_by_position(
    table_bbox: List[float],
    page_num: int,
    docx_tables: List[Dict[str, Any]],
    docx_paragraphs: List[Dict[str, Any]]
) -> Optional[Dict[str, Any]]:
    """
    Находит таблицу в DOCX по позиции и контексту.
    Использует XML позицию, estimated_page и порядок таблиц для более точного сопоставления.
    
    Args:
        table_bbox: Координаты таблицы из Dots.OCR
        page_num: Номер страницы (0-based, из Dots.OCR)
        docx_tables: Список таблиц из DOCX (с XML метаданными)
        docx_paragraphs: Список параграфов из DOCX
    
    Returns:
        Найденная таблица или None
    """
    if not docx_tables:
        return None
    
    # Конвертируем page_num из 0-based в 1-based для сравнения с estimated_page
    page_num_1based = page_num + 1
    
    # Приоритет 1: Ищем таблицы с estimated_page, близким к page_num
    tables_with_estimated_page = [
        t for t in docx_tables 
        if t.get("estimated_page") is not None
    ]
    
    if tables_with_estimated_page:
        # Сортируем по близости estimated_page к page_num
        def page_distance_score(table):
            est_page = table.get("estimated_page", 9999)
            distance = abs(est_page - page_num_1based)
            # Бонус за точное совпадение
            if est_page == page_num_1based:
                return (distance, -1000)  # Приоритет точному совпадению
            return (distance, table.get("xml_position", 9999))
        
        tables_with_estimated_page.sort(key=page_distance_score)
        
        # Берем таблицу с ближайшим estimated_page
        best_match = tables_with_estimated_page[0]
        est_page = best_match.get("estimated_page")
        
        # Если estimated_page близок к page_num (разница <= 2), используем его
        if abs(est_page - page_num_1based) <= 2:
            return best_match
    
    # Приоритет 2: Используем порядок таблиц по XML позиции
    tables_with_position = [t for t in docx_tables if t.get("xml_position") is not None]
    
    if tables_with_position:
        # Сортируем по XML позиции
        tables_with_position.sort(key=lambda x: x.get("xml_position", 9999))
        
        # Если есть estimated_page, используем его для фильтрации
        # Иначе просто возвращаем первую таблицу
        for table in tables_with_position:
            est_page = table.get("estimated_page")
            if est_page is None or abs(est_page - page_num_1based) <= 3:
                return table
        
        # Если ничего не подошло, возвращаем первую
        return tables_with_position[0]
    
    # Fallback: возвращаем первую таблицу
    return docx_tables[0] if docx_tables else None


def extract_image_from_pdf_by_bbox(
    pdf_path: Path,
    image_bbox: List[float],
    page_num: int,
    render_scale: float = None,
    rendered_page_image: Optional[Image.Image] = None
) -> Optional[Image.Image]:
    """
    Извлекает изображение из PDF по координатам bbox.
    
    Args:
        pdf_path: Путь к PDF файлу
        image_bbox: Координаты [x1, y1, x2, y2] из Dots.OCR (в масштабе отрендеренного изображения)
        page_num: Номер страницы (0-based)
        render_scale: Масштаб рендеринга (если None, используется RENDER_SCALE)
        rendered_page_image: Опционально - уже отрендеренное изображение страницы (для ускорения)
    
    Returns:
        PIL Image или None
    """
    if render_scale is None:
        render_scale = RENDER_SCALE
    
    if not HAS_PYMUPDF:
        return None
    
    # Если есть уже отрендеренное изображение страницы, используем его
    if rendered_page_image is not None:
        try:
            x1, y1, x2, y2 = float(image_bbox[0]), float(image_bbox[1]), float(image_bbox[2]), float(image_bbox[3])
            
            # Координаты bbox из Dots.OCR могут относиться к optimized_image
            # Нужно проверить размеры и при необходимости масштабировать
            width, height = rendered_page_image.size
            
            # Проверяем, не нужно ли масштабировать координаты
            # Если bbox выходит за границы изображения, возможно координаты в другом масштабе
            if x2 > width or y2 > height:
                # Пробуем найти масштаб
                # Обычно optimized_image может быть меньше original_image
                # Но в нашем случае мы используем original_image, так что координаты должны совпадать
                pass
            
            # Обрезаем изображение по координатам bbox
            x1 = max(0, int(x1))
            y1 = max(0, int(y1))
            x2 = min(width, int(x2))
            y2 = min(height, int(y2))
            
            # Проверяем валидность координат
            if x2 > x1 and y2 > y1 and (x2 - x1) > 10 and (y2 - y1) > 10:
                cropped_image = rendered_page_image.crop((x1, y1, x2, y2))
                
                # Проверяем, что изображение не пустое (не все пиксели черные)
                # Это поможет избежать черных квадратов
                import numpy as np
                img_array = np.array(cropped_image)
                if img_array.size > 0:
                    # Проверяем, что не все пиксели черные (0,0,0)
                    non_black_pixels = np.sum(np.any(img_array != [0, 0, 0], axis=2))
                    total_pixels = img_array.shape[0] * img_array.shape[1]
                    
                    if non_black_pixels > total_pixels * 0.1:  # Хотя бы 10% пикселей не черные
                        return cropped_image
                    else:
                        print(f"  Предупреждение: извлеченное изображение слишком темное (возможно неправильные координаты)")
                        # Пробуем fallback метод
        except Exception as e:
            print(f"Ошибка при обрезке изображения: {e}, пробуем fallback метод")
    
    # Fallback: извлекаем из PDF напрямую
    pdf_document = fitz.open(str(pdf_path))
    try:
        if page_num >= len(pdf_document):
            return None
        
        page = pdf_document.load_page(page_num)
        
        # Приводим координаты к масштабу оригинального PDF
        x1, y1, x2, y2 = (
            image_bbox[0] / render_scale,
            image_bbox[1] / render_scale,
            image_bbox[2] / render_scale,
            image_bbox[3] / render_scale,
        )
        rect = fitz.Rect(x1, y1, x2, y2)
        
        # Извлекаем изображение из области bbox с тем же масштабом
        mat = fitz.Matrix(render_scale, render_scale)
        pix = page.get_pixmap(matrix=mat, clip=rect)
        img_data = pix.tobytes("png")
        pix = None
        
        image = Image.open(BytesIO(img_data)).convert("RGB")
        return image
    except Exception as e:
        print(f"Ошибка при извлечении изображения из PDF: {e}")
        return None
    finally:
        pdf_document.close()


def calculate_image_hash(image: Image.Image) -> Dict[str, Any]:
    """
    Вычисляет различные хеши для изображения.
    
    Args:
        image: PIL Image
    
    Returns:
        Словарь с хешами и метаданными
    """
    result = {
        "md5_hash": None,
        "perceptual_hash": None,
        "average_hash": None,
        "p_hash": None,
        "d_hash": None,
        "width": image.width,
        "height": image.height,
        "size": image.width * image.height,
    }
    
    # MD5 хеш (для идентичных изображений)
    img_bytes = BytesIO()
    image.save(img_bytes, format="PNG")
    img_bytes.seek(0)
    result["md5_hash"] = hashlib.md5(img_bytes.getvalue()).hexdigest()
    
    # Perceptual hashes (для визуально похожих изображений)
    if HAS_IMAGEHASH:
        try:
            result["perceptual_hash"] = str(imagehash.phash(image))
            result["average_hash"] = str(imagehash.average_hash(image))
            result["p_hash"] = str(imagehash.phash(image))  # Perceptual hash
            result["d_hash"] = str(imagehash.dhash(image))  # Difference hash
        except Exception as e:
            print(f"Ошибка при вычислении perceptual hash: {e}")
    
    return result


def compare_images(
    image1: Image.Image,
    image2: Image.Image,
    use_perceptual_hash: bool = True
) -> Dict[str, Any]:
    """
    Сравнивает два изображения и возвращает метрики схожести.
    
    Args:
        image1: Первое изображение
        image2: Второе изображение
        use_perceptual_hash: Использовать perceptual hash для сравнения
    
    Returns:
        Словарь с метриками сравнения
    """
    result = {
        "md5_match": False,
        "perceptual_hash_distance": None,
        "average_hash_distance": None,
        "p_hash_distance": None,
        "d_hash_distance": None,
        "size_similarity": 0.0,
        "is_exact_match": False,
        "is_visual_match": False,
    }
    
    # Вычисляем хеши для обоих изображений
    hash1 = calculate_image_hash(image1)
    hash2 = calculate_image_hash(image2)
    
    # MD5 сравнение (идентичные изображения)
    result["md5_match"] = hash1["md5_hash"] == hash2["md5_hash"]
    result["is_exact_match"] = result["md5_match"]
    
    # Сравнение размеров
    size1 = hash1["size"]
    size2 = hash2["size"]
    if size1 > 0 and size2 > 0:
        result["size_similarity"] = min(size1, size2) / max(size1, size2)
    
    # Perceptual hash сравнение (визуально похожие изображения)
    if use_perceptual_hash and HAS_IMAGEHASH:
        try:
            if hash1["perceptual_hash"] and hash2["perceptual_hash"]:
                phash1 = imagehash.hex_to_hash(hash1["perceptual_hash"])
                phash2 = imagehash.hex_to_hash(hash2["perceptual_hash"])
                result["perceptual_hash_distance"] = phash1 - phash2
            
            if hash1["average_hash"] and hash2["average_hash"]:
                ahash1 = imagehash.hex_to_hash(hash1["average_hash"])
                ahash2 = imagehash.hex_to_hash(hash2["average_hash"])
                result["average_hash_distance"] = ahash1 - ahash2
            
            if hash1["p_hash"] and hash2["p_hash"]:
                phash1 = imagehash.hex_to_hash(hash1["p_hash"])
                phash2 = imagehash.hex_to_hash(hash2["p_hash"])
                result["p_hash_distance"] = phash1 - phash2
            
            if hash1["d_hash"] and hash2["d_hash"]:
                dhash1 = imagehash.hex_to_hash(hash1["d_hash"])
                dhash2 = imagehash.hex_to_hash(hash2["d_hash"])
                result["d_hash_distance"] = dhash1 - dhash2
            
            # Считаем визуально похожими, если perceptual hash distance <= 5
            # (чем меньше расстояние, тем более похожи изображения)
            if result["perceptual_hash_distance"] is not None:
                result["is_visual_match"] = result["perceptual_hash_distance"] <= 5
            elif result["p_hash_distance"] is not None:
                result["is_visual_match"] = result["p_hash_distance"] <= 5
        except Exception as e:
            print(f"Ошибка при сравнении perceptual hash: {e}")
    
    return result


def compare_images_orb_ransac(
    image1: Image.Image,
    image2: Image.Image,
    min_inliers: int = 15,
    max_reproj_error: float = 3.0,
    inlier_ratio: float = 0.7
) -> Tuple[bool, Dict[str, Any]]:
    """
    Точное сравнение двух изображений разного размера через ключевые точки (ORB + RANSAC).
    
    Этот метод использует детекцию ключевых точек и геометрическую верификацию через RANSAC
    для точного сопоставления изображений, которые могут отличаться размером, качеством,
    но визуально идентичны.
    
    Алгоритм:
    1. Детекция и описание ключевых точек (ORB)
    2. Матчинг дескрипторов (Brute Force Matcher)
    3. Геометрическая верификация через RANSAC (поиск гомографии)
    4. Проверка критериев точного совпадения
    
    Args:
        image1: Первое изображение (PIL Image)
        image2: Второе изображение (PIL Image)
        min_inliers: Минимум совпадений после RANSAC для признания совпадения
        max_reproj_error: Максимальная ошибка репроекции в пикселях
        inlier_ratio: Минимальная доля inliers от общего числа матчей
    
    Returns:
        Tuple[bool, Dict]: (совпадают ли изображения, статистика сравнения)
    """
    if not HAS_OPENCV:
        return False, {"error": "OpenCV не установлен"}
    
    try:
        # 1. Конвертируем PIL Image в numpy array (grayscale)
        img1_array = np.array(image1.convert("L"))
        img2_array = np.array(image2.convert("L"))
        
        # 2. Детекция и описание ключевых точек (ORB — быстрый и патентно-свободный)
        orb = cv2.ORB_create(nfeatures=1000, scaleFactor=1.2, edgeThreshold=10)
        kp1, des1 = orb.detectAndCompute(img1_array, None)
        kp2, des2 = orb.detectAndCompute(img2_array, None)
        
        if des1 is None or des2 is None or len(kp1) < 10 or len(kp2) < 10:
            return False, {
                "error": "Недостаточно ключевых точек",
                "kp1_count": len(kp1) if kp1 else 0,
                "kp2_count": len(kp2) if kp2 else 0
            }
        
        # 3. Матчинг дескрипторов
        bf = cv2.BFMatcher(cv2.NORM_HAMMING, crossCheck=True)
        matches = bf.match(des1, des2)
        
        if len(matches) < min_inliers:
            return False, {
                "matches_total": len(matches),
                "min_required": min_inliers
            }
        
        # 4. Подготовка координат для геометрической верификации
        src_pts = np.float32([kp1[m.queryIdx].pt for m in matches]).reshape(-1, 1, 2)
        dst_pts = np.float32([kp2[m.trainIdx].pt for m in matches]).reshape(-1, 1, 2)
        
        # 5. Поиск гомографии с RANSAC (проверка геометрической согласованности)
        try:
            M, mask = cv2.findHomography(
                src_pts, dst_pts,
                cv2.RANSAC,
                ransacReprojThreshold=max_reproj_error
            )
            
            if mask is None:
                return False, {
                    "matches_total": len(matches),
                    "inliers": 0,
                    "error": "Не удалось вычислить гомографию"
                }
            
            inliers = int(mask.sum())
            inlier_ratio_actual = inliers / len(matches) if len(matches) > 0 else 0.0
            
            # 6. Критерии точного совпадения
            is_match = (
                inliers >= min_inliers and
                inlier_ratio_actual >= inlier_ratio
            )
            
            # Проверяем, что гомография не вырождена
            if M is not None:
                det = cv2.determinant(M[:2, :2])
                if det <= 0.01:
                    is_match = False
            
            stats = {
                "matches_total": len(matches),
                "inliers": inliers,
                "inlier_ratio": round(inlier_ratio_actual, 3),
                "scale_estimate": round(np.sqrt(abs(cv2.determinant(M[:2, :2]))), 3) if M is not None else None,
                "is_match": is_match
            }
            
            return is_match, stats
            
        except cv2.error as e:
            return False, {
                "error": f"Ошибка при вычислении гомографии: {str(e)}",
                "matches_total": len(matches)
            }
    
    except Exception as e:
        return False, {"error": f"Ошибка при сравнении изображений: {str(e)}"}


def compare_images_combined(
    image1: Image.Image,
    image2: Image.Image,
    use_orb: bool = True,
    use_perceptual_hash: bool = True,
    orb_min_inliers: int = 15,
    orb_inlier_ratio: float = 0.7,
    perceptual_threshold: int = 15
) -> Dict[str, Any]:
    """
    Комбинированное сравнение изображений: ORB+RANSAC + Perceptual Hash.
    
    Использует два метода:
    1. ORB + RANSAC (точное геометрическое сопоставление) - приоритет
    2. Perceptual Hash (быстрое визуальное сравнение) - fallback
    
    Args:
        image1: Первое изображение
        image2: Второе изображение
        use_orb: Использовать ORB+RANSAC метод
        use_perceptual_hash: Использовать perceptual hash
        orb_min_inliers: Минимум inliers для ORB
        orb_inlier_ratio: Минимальная доля inliers
        perceptual_threshold: Порог для perceptual hash
    
    Returns:
        Словарь с результатами сравнения обоими методами
    """
    result = {
        "method": "combined",
        "orb_match": False,
        "orb_stats": None,
        "perceptual_match": False,
        "perceptual_stats": None,
        "is_match": False,
        "match_confidence": "none"  # "high" (ORB), "medium" (perceptual), "low" (size only), "none"
    }
    
    # Метод 1: ORB + RANSAC (наиболее точный)
    if use_orb and HAS_OPENCV:
        orb_match, orb_stats = compare_images_orb_ransac(
            image1, image2,
            min_inliers=orb_min_inliers,
            inlier_ratio=orb_inlier_ratio
        )
        result["orb_match"] = orb_match
        result["orb_stats"] = orb_stats
        
        if orb_match:
            result["is_match"] = True
            result["match_confidence"] = "high"
            return result
    
    # Метод 2: Perceptual Hash (fallback, если ORB не сработал)
    if use_perceptual_hash and HAS_IMAGEHASH:
        perceptual_result = compare_images(image1, image2, use_perceptual_hash=True)
        result["perceptual_stats"] = perceptual_result
        
        # Проверяем визуальное совпадение
        visual_match = False
        if perceptual_result.get("perceptual_hash_distance") is not None:
            visual_match = perceptual_result["perceptual_hash_distance"] <= perceptual_threshold
        elif perceptual_result.get("p_hash_distance") is not None:
            visual_match = perceptual_result["p_hash_distance"] <= perceptual_threshold
        elif perceptual_result.get("average_hash_distance") is not None:
            visual_match = perceptual_result["average_hash_distance"] <= perceptual_threshold
        
        result["perceptual_match"] = visual_match
        
        if visual_match:
            result["is_match"] = True
            if result["match_confidence"] == "none":
                result["match_confidence"] = "medium"
    
    # Метод 3: Размер (только если другие методы не сработали)
    if not result["is_match"]:
        perceptual_result = compare_images(image1, image2, use_perceptual_hash=False)
        size_similarity = perceptual_result.get("size_similarity", 0.0)
        if size_similarity > 0.9:
            result["is_match"] = True
            result["match_confidence"] = "low"
            result["perceptual_stats"] = perceptual_result
    
    return result


def find_exact_image_in_docx(
    image_bbox: List[float],
    page_num: int,
    docx_images: List[Dict[str, Any]],
    pdf_path: Path,
    render_scale: float = None,
    perceptual_threshold: int = 10,  # Увеличен порог для более мягкого сравнения
    rendered_page_image: Optional[Image.Image] = None
) -> Optional[Dict[str, Any]]:
    """
    Улучшенное точное сопоставление изображения из Dots.OCR с изображением из DOCX.
    
    Алгоритм:
    1. Извлекаем изображение из PDF по координатам bbox
    2. Вычисляем хеши (MD5, perceptual hash) для изображения из PDF
    3. Сравниваем с каждым изображением из DOCX используя множественные метрики
    4. Находим наиболее похожее изображение (приоритет: точное совпадение > визуальное > размер)
    
    Args:
        image_bbox: Координаты изображения из Dots.OCR [x1, y1, x2, y2]
        page_num: Номер страницы (0-based)
        docx_images: Список изображений из DOCX (с предвычисленными хешами)
        pdf_path: Путь к PDF файлу
        render_scale: Масштаб рендеринга (если None, используется RENDER_SCALE)
        perceptual_threshold: Порог для perceptual hash (чем больше, тем мягче)
        rendered_page_image: Опционально - уже отрендеренное изображение страницы
    
    Returns:
        Найденное изображение из DOCX или None
    """
    if render_scale is None:
        render_scale = RENDER_SCALE
    
    if not docx_images:
        return None
    
    # Шаг 1: Извлекаем изображение из PDF по координатам
    pdf_image = extract_image_from_pdf_by_bbox(pdf_path, image_bbox, page_num, render_scale, rendered_page_image)
    if pdf_image is None:
        return None
    
    # Проверяем, что изображение не пустое
    if pdf_image.size[0] == 0 or pdf_image.size[1] == 0:
        return None
    
    # Шаг 2: Вычисляем хеши для изображения из PDF
    pdf_hash = calculate_image_hash(pdf_image)
    pdf_size = pdf_image.size[0] * pdf_image.size[1]
    
    # Шаг 3: Сравниваем с изображениями из DOCX
    best_match = None
    best_score = -1.0
    best_match_type = None  # "exact", "visual", "size"
    best_comparison_details = None
    
    for docx_image in docx_images:
        # Загружаем изображение из DOCX
        try:
            if not docx_image.get("image_bytes"):
                continue
            
            docx_img = Image.open(BytesIO(docx_image["image_bytes"])).convert("RGB")
            
            # Проверяем валидность изображения
            if docx_img.size[0] == 0 or docx_img.size[1] == 0:
                continue
        except Exception as e:
            continue
        
        # Вычисляем хеши для изображения из DOCX (если еще не вычислены)
        if "perceptual_hash" not in docx_image or docx_image.get("perceptual_hash") is None:
            docx_hash = calculate_image_hash(docx_img)
            docx_image["perceptual_hash"] = docx_hash.get("perceptual_hash")
            docx_image["average_hash"] = docx_hash.get("average_hash")
            docx_image["p_hash"] = docx_hash.get("p_hash")
            docx_image["d_hash"] = docx_hash.get("d_hash")
        
        # Сравнение (комбинированный метод: ORB + Perceptual Hash)
        comparison = compare_images_combined(
            pdf_image, docx_img,
            use_orb=True,
            use_perceptual_hash=True,
            orb_min_inliers=15,
            orb_inlier_ratio=0.7,
            perceptual_threshold=perceptual_threshold
        )
        
        # Приоритет 1: Точное совпадение через ORB (наиболее надежно)
        if comparison.get("orb_match", False):
            docx_image["match_type"] = "orb_exact"
            docx_image["match_score"] = 100.0
            docx_image["comparison_details"] = comparison
            return docx_image
        
        # Приоритет 2: Точное совпадение (MD5)
        perceptual_result = compare_images(pdf_image, docx_img, use_perceptual_hash=True)
        if perceptual_result.get("is_exact_match", False):
            docx_image["match_type"] = "exact"
            docx_image["match_score"] = 100.0
            docx_image["comparison_details"] = comparison
            return docx_image  # Нашли точное совпадение, возвращаем сразу
        
        # Приоритет 3: Визуальное совпадение через Perceptual Hash
        visual_match = comparison.get("perceptual_match", False)
        visual_score = 0.0
        
        if visual_match:
            # Используем статистику из perceptual hash для score
            perceptual_stats = comparison.get("perceptual_stats", {})
            if perceptual_stats.get("perceptual_hash_distance") is not None:
                phash_dist = perceptual_stats["perceptual_hash_distance"]
                visual_score = 100.0 - (phash_dist / perceptual_threshold * 50.0)
            elif perceptual_stats.get("p_hash_distance") is not None:
                phash_dist = perceptual_stats["p_hash_distance"]
                visual_score = 100.0 - (phash_dist / perceptual_threshold * 50.0)
            elif perceptual_stats.get("average_hash_distance") is not None:
                phash_dist = perceptual_stats["average_hash_distance"]
                visual_score = 100.0 - (phash_dist / perceptual_threshold * 50.0)
            else:
                visual_score = 80.0  # Базовый score для визуального совпадения
        
        if visual_match:
            # Бонус за похожесть размеров
            size_bonus = 0.0
            if comparison["size_similarity"] > 0.7:
                size_bonus = comparison["size_similarity"] * 20.0
            
            total_score = visual_score + size_bonus
            
            if total_score > best_score:
                best_score = total_score
                best_match = docx_image
                best_match_type = "visual"
                best_comparison_details = comparison
        
        # Приоритет 3: Похожесть по размеру (если визуальное не найдено)
        elif best_match_type != "visual":
            size_score = comparison["size_similarity"] * 50.0  # Максимум 50 баллов за размер
            
            # Небольшой бонус, если размеры очень похожи
            if comparison["size_similarity"] > 0.9:
                size_score += 20.0
            
            if size_score > best_score:
                best_score = size_score
                best_match = docx_image
                best_match_type = "size"
                best_comparison_details = comparison
    
    # Возвращаем лучшее совпадение
    if best_match is not None:
        best_match["match_type"] = best_match_type
        best_match["match_score"] = best_score
        best_match["comparison_details"] = best_comparison_details
        return best_match
    
    return None


def extract_context_near_image(
    image_bbox: List[float],
    page_num: int,
    all_layout_elements: List[Dict[str, Any]],
    max_distance: float = 100.0
) -> Dict[str, Any]:
    """
    Извлекает контекст вокруг изображения (текст, подписи, заголовки).
    
    Args:
        image_bbox: Координаты изображения [x1, y1, x2, y2]
        page_num: Номер страницы
        all_layout_elements: Все элементы layout
        max_distance: Максимальное расстояние для поиска контекста
    
    Returns:
        Словарь с контекстом (текст выше, ниже, рядом)
    """
    context = {
        "text_above": [],
        "text_below": [],
        "text_nearby": [],
        "captions": [],
        "figure_numbers": []
    }
    
    x1, y1, x2, y2 = image_bbox[0], image_bbox[1], image_bbox[2], image_bbox[3]
    image_center_y = (y1 + y2) / 2
    image_center_x = (x1 + x2) / 2
    
    # Ищем элементы на той же странице
    page_elements = [e for e in all_layout_elements if e.get("page_num") == page_num]
    
    for element in page_elements:
        elem_bbox = element.get("bbox", [])
        if len(elem_bbox) < 4:
            continue
        
        ex1, ey1, ex2, ey2 = elem_bbox[0], elem_bbox[1], elem_bbox[2], elem_bbox[3]
        elem_center_y = (ey1 + ey2) / 2
        elem_center_x = (ex1 + ex2) / 2
        
        # Расстояние по Y
        y_distance = abs(elem_center_y - image_center_y)
        # Расстояние по X
        x_distance = abs(elem_center_x - image_center_x)
        
        # Текст выше изображения
        if ey2 < y1 and x_distance < max_distance:
            text = element.get("text", "").strip()
            if text:
                context["text_above"].append({
                    "text": text,
                    "distance": y_distance,
                    "category": element.get("category", "")
                })
        
        # Текст ниже изображения
        elif ey1 > y2 and x_distance < max_distance:
            text = element.get("text", "").strip()
            if text:
                context["text_below"].append({
                    "text": text,
                    "distance": y_distance,
                    "category": element.get("category", "")
                })
                # Проверяем на подпись (Caption)
                if element.get("category") == "Caption":
                    context["captions"].append(text)
                    # Ищем номера рисунков
                    figure_match = re.search(r'(?:рис|рисунок|figure|fig)[\s\.]*(\d+)', text.lower())
                    if figure_match:
                        context["figure_numbers"].append(int(figure_match.group(1)))
        
        # Текст рядом (по горизонтали)
        elif abs(ey1 - image_center_y) < max_distance and x_distance < max_distance * 2:
            text = element.get("text", "").strip()
            if text:
                context["text_nearby"].append({
                    "text": text,
                    "distance": x_distance,
                    "category": element.get("category", "")
                })
    
    # Сортируем по расстоянию
    context["text_above"].sort(key=lambda x: x["distance"])
    context["text_below"].sort(key=lambda x: x["distance"])
    context["text_nearby"].sort(key=lambda x: x["distance"])
    
    return context


def find_image_by_context(
    context: Dict[str, Any],
    docx_paragraphs: List[Dict[str, Any]],
    docx_images: List[Dict[str, Any]]
) -> Optional[Dict[str, Any]]:
    """
    Ищет изображение в DOCX по контексту (подписям, номерам рисунков).
    
    Args:
        context: Контекст вокруг изображения из OCR
        docx_paragraphs: Параграфы из DOCX
        docx_images: Изображения из DOCX
    
    Returns:
        Найденное изображение или None
    """
    # Ищем по номерам рисунков
    if context.get("figure_numbers"):
        figure_num = context["figure_numbers"][0]
        
        # Ищем параграф с упоминанием этого номера
        for para in docx_paragraphs:
            text_lower = para.get("text", "").lower()
            # Ищем паттерны типа "рисунок 1", "рис. 1", "figure 1"
            pattern = rf'(?:рис|рисунок|figure|fig)[\s\.]*{figure_num}\b'
            if re.search(pattern, text_lower):
                # Нашли текст с номером, ищем ближайшее изображение
                para_idx = para.get("index", -1)
                
                # Ищем изображения в близких параграфах
                for docx_img in docx_images:
                    img_para_idx = docx_img.get("paragraph_index", -1)
                    if img_para_idx != -1 and abs(img_para_idx - para_idx) <= 3:
                        return docx_img
    
    # Ищем по подписям (captions)
    if context.get("captions"):
        caption_text = context["captions"][0].lower()
        
        # Ищем похожий текст в DOCX
        for para in docx_paragraphs:
            para_text_lower = para.get("text", "").lower().strip()
            # Простое сравнение подписи
            if len(caption_text) > 10:  # Только для достаточно длинных подписей
                # Используем частичное совпадение
                words_caption = set(caption_text.split())
                words_para = set(para_text_lower.split())
                common_words = words_caption & words_para
                if len(common_words) >= min(3, len(words_caption) * 0.5):
                    # Нашли похожий текст, ищем ближайшее изображение
                    para_idx = para.get("index", -1)
                    for docx_img in docx_images:
                        img_para_idx = docx_img.get("paragraph_index", -1)
                        if img_para_idx != -1 and abs(img_para_idx - para_idx) <= 2:
                            return docx_img
    
    return None


def calculate_image_position_score(
    image_bbox: List[float],
    page_width: float,
    page_height: float,
    docx_image: Dict[str, Any],
    docx_paragraphs: List[Dict[str, Any]]
) -> float:
    """
    Вычисляет score на основе позиции изображения относительно страницы.
    
    Args:
        image_bbox: Координаты изображения из OCR
        page_width: Ширина страницы
        page_height: Высота страницы
        docx_image: Изображение из DOCX
        docx_paragraphs: Параграфы из DOCX
    
    Returns:
        Score от 0 до 1 (1 = идеальное совпадение позиции)
    """
    if page_width == 0 or page_height == 0:
        return 0.0
    
    # Относительная позиция изображения из OCR
    ocr_x_center = (image_bbox[0] + image_bbox[2]) / 2 / page_width
    ocr_y_center = (image_bbox[1] + image_bbox[3]) / 2 / page_height
    
    # Пытаемся определить позицию изображения из DOCX
    # Если есть paragraph_index, можем оценить позицию
    para_idx = docx_image.get("paragraph_index", -1)
    if para_idx != -1 and para_idx < len(docx_paragraphs):
        # Примерная оценка: изображения обычно идут в порядке параграфов
        # Это грубая оценка, но может помочь
        total_paras = len(docx_paragraphs)
        estimated_y = para_idx / max(total_paras, 1)  # Относительная позиция по документу
        
        # Сравниваем относительные позиции
        y_diff = abs(ocr_y_center - estimated_y)
        score = max(0.0, 1.0 - y_diff * 2)  # Чем ближе, тем выше score
        
        return score
    
    return 0.0


def compare_images_histogram(
    image1: Image.Image,
    image2: Image.Image
) -> Dict[str, float]:
    """
    Сравнивает изображения через гистограммы цветов.
    
    Args:
        image1: Первое изображение
        image2: Второе изображение
    
    Returns:
        Словарь с метриками сравнения гистограмм
    """
    if not HAS_OPENCV:
        return {"histogram_similarity": 0.0}
    
    try:
        # Конвертируем в numpy arrays
        img1_array = np.array(image1.convert("RGB"))
        img2_array = np.array(image2.convert("RGB"))
        
        # Вычисляем гистограммы для каждого канала
        hist1_b = cv2.calcHist([img1_array], [0], None, [256], [0, 256])
        hist1_g = cv2.calcHist([img1_array], [1], None, [256], [0, 256])
        hist1_r = cv2.calcHist([img1_array], [2], None, [256], [0, 256])
        
        hist2_b = cv2.calcHist([img2_array], [0], None, [256], [0, 256])
        hist2_g = cv2.calcHist([img2_array], [1], None, [256], [0, 256])
        hist2_r = cv2.calcHist([img2_array], [2], None, [256], [0, 256])
        
        # Сравниваем гистограммы (корреляция)
        corr_b = cv2.compareHist(hist1_b, hist2_b, cv2.HISTCMP_CORREL)
        corr_g = cv2.compareHist(hist1_g, hist2_g, cv2.HISTCMP_CORREL)
        corr_r = cv2.compareHist(hist1_r, hist2_r, cv2.HISTCMP_CORREL)
        
        # Средняя корреляция
        avg_corr = (corr_b + corr_g + corr_r) / 3.0
        
        return {
            "histogram_similarity": float(avg_corr),
            "histogram_corr_b": float(corr_b),
            "histogram_corr_g": float(corr_g),
            "histogram_corr_r": float(corr_r)
        }
    except Exception as e:
        return {"histogram_similarity": 0.0, "error": str(e)}


def find_exact_image_in_docx_improved(
    image_bbox: List[float],
    page_num: int,
    docx_images: List[Dict[str, Any]],
    pdf_path: Path,
    render_scale: float = None,
    perceptual_threshold: int = 15,  # Еще более мягкий порог
    rendered_page_image: Optional[Image.Image] = None,
    image_order_in_page: int = 0,  # Порядок изображения на странице
    all_layout_elements: Optional[List[Dict[str, Any]]] = None,  # Все элементы для контекста
    docx_paragraphs: Optional[List[Dict[str, Any]]] = None,  # Параграфы для контекста
    page_width: Optional[float] = None,  # Ширина страницы
    page_height: Optional[float] = None  # Высота страницы
) -> Optional[Dict[str, Any]]:
    """
    Улучшенное сопоставление изображений с учетом порядка и позиции.
    
    Args:
        image_bbox: Координаты изображения из Dots.OCR [x1, y1, x2, y2]
        page_num: Номер страницы (0-based)
        docx_images: Список изображений из DOCX
        pdf_path: Путь к PDF файлу
        render_scale: Масштаб рендеринга
        perceptual_threshold: Порог для perceptual hash
        rendered_page_image: Отрендеренное изображение страницы
        image_order_in_page: Порядок изображения на странице (0-based)
    
    Returns:
        Найденное изображение из DOCX или None
    """
    if render_scale is None:
        render_scale = RENDER_SCALE
    
    if not docx_images:
        return None
    
    # НОВЫЙ ПОДХОД: Используем контекст для первоначальной фильтрации
    context_match = None
    if all_layout_elements and docx_paragraphs:
        context = extract_context_near_image(image_bbox, page_num, all_layout_elements)
        context_match = find_image_by_context(context, docx_paragraphs, docx_images)
        
        # Если нашли по контексту, проверяем визуально
        if context_match:
            pdf_image = extract_image_from_pdf_by_bbox(pdf_path, image_bbox, page_num, render_scale, rendered_page_image)
            if pdf_image and pdf_image.size[0] > 0 and pdf_image.size[1] > 0:
                try:
                    docx_img = Image.open(BytesIO(context_match["image_bytes"])).convert("RGB")
                    # Быстрая проверка через гистограммы
                    hist_result = compare_images_histogram(pdf_image, docx_img)
                    if hist_result.get("histogram_similarity", 0.0) > 0.7:
                        # Проверяем через комбинированный метод
                        comparison = compare_images_combined(
                            pdf_image, docx_img,
                            use_orb=True,
                            use_perceptual_hash=True,
                            orb_min_inliers=10,  # Более мягкий порог
                            orb_inlier_ratio=0.6,
                            perceptual_threshold=20
                        )
                        if comparison.get("is_match", False):
                            context_match["match_type"] = "context_" + comparison.get("match_confidence", "medium")
                            context_match["match_score"] = 95.0
                            context_match["comparison_details"] = comparison
                            return context_match
                except Exception:
                    pass
    
    # Извлекаем изображение из PDF
    pdf_image = extract_image_from_pdf_by_bbox(pdf_path, image_bbox, page_num, render_scale, rendered_page_image)
    if pdf_image is None or pdf_image.size[0] == 0 or pdf_image.size[1] == 0:
        return None
    
    # Вычисляем хеши для PDF изображения
    pdf_hash = calculate_image_hash(pdf_image)
    pdf_size = pdf_image.size[0] * pdf_image.size[1]
    
    # Сортируем изображения из DOCX по порядку в документе
    sorted_docx_images = sorted(docx_images, key=lambda x: x.get("order_in_document", x.get("index", 9999)))
    
    best_match = None
    best_score = -1.0
    best_match_type = None
    best_comparison_details = None
    
    for docx_image in sorted_docx_images:
        # Пропускаем уже сопоставленные (если есть флаг)
        if docx_image.get("matched", False):
            continue
        
        try:
            if not docx_image.get("image_bytes"):
                continue
            
            docx_img = Image.open(BytesIO(docx_image["image_bytes"])).convert("RGB")
            
            if docx_img.size[0] == 0 or docx_img.size[1] == 0:
                continue
        except Exception:
            continue
        
        # Вычисляем хеши
        if "perceptual_hash" not in docx_image or docx_image.get("perceptual_hash") is None:
            docx_hash = calculate_image_hash(docx_img)
            docx_image["perceptual_hash"] = docx_hash.get("perceptual_hash")
            docx_image["average_hash"] = docx_hash.get("average_hash")
            docx_image["p_hash"] = docx_hash.get("p_hash")
            docx_image["d_hash"] = docx_hash.get("d_hash")
        
        # Сравнение (комбинированный метод: ORB + Perceptual Hash + Histogram)
        comparison = compare_images_combined(
            pdf_image, docx_img,
            use_orb=True,
            use_perceptual_hash=True,
            orb_min_inliers=10,  # Более мягкий порог
            orb_inlier_ratio=0.6,  # Более мягкий порог
            perceptual_threshold=20  # Более мягкий порог
        )
        
        # Добавляем сравнение через гистограммы
        hist_result = compare_images_histogram(pdf_image, docx_img)
        comparison["histogram_similarity"] = hist_result.get("histogram_similarity", 0.0)
        
        # Добавляем score позиции
        if page_width and page_height:
            position_score = calculate_image_position_score(
                image_bbox, page_width, page_height, docx_image, docx_paragraphs or []
            )
            comparison["position_score"] = position_score
        
        # Приоритет 1: Точное совпадение через ORB (наиболее надежно)
        if comparison.get("orb_match", False):
            docx_image["match_type"] = "orb_exact"
            docx_image["match_score"] = 100.0
            docx_image["comparison_details"] = comparison
            return docx_image
        
        # Приоритет 2: Точное совпадение (MD5)
        perceptual_result = compare_images(pdf_image, docx_img, use_perceptual_hash=True)
        if perceptual_result.get("is_exact_match", False):
            docx_image["match_type"] = "exact"
            docx_image["match_score"] = 100.0
            docx_image["comparison_details"] = comparison
            return docx_image
        
        # Приоритет 3: Визуальное совпадение через Perceptual Hash
        visual_match = comparison.get("perceptual_match", False)
        visual_score = 0.0
        
        if visual_match:
            # Используем статистику из perceptual hash для score
            perceptual_stats = comparison.get("perceptual_stats", {})
            if perceptual_stats.get("perceptual_hash_distance") is not None:
                phash_dist = perceptual_stats["perceptual_hash_distance"]
                visual_score = 100.0 - (phash_dist / 20.0 * 50.0)  # Используем 20 как порог
            elif perceptual_stats.get("p_hash_distance") is not None:
                phash_dist = perceptual_stats["p_hash_distance"]
                visual_score = 100.0 - (phash_dist / 20.0 * 50.0)
            elif perceptual_stats.get("average_hash_distance") is not None:
                phash_dist = perceptual_stats["average_hash_distance"]
                visual_score = 100.0 - (phash_dist / 20.0 * 50.0)
            else:
                visual_score = 80.0  # Базовый score для визуального совпадения
        
        if visual_match:
            # Бонусы
            size_bonus = 0.0
            perceptual_stats = comparison.get("perceptual_stats", {})
            size_similarity = perceptual_stats.get("size_similarity", 0.0)
            if size_similarity > 0.7:
                size_bonus = size_similarity * 20.0
            
            # Бонус за гистограмму
            hist_bonus = 0.0
            hist_sim = comparison.get("histogram_similarity", 0.0)
            if hist_sim > 0.8:
                hist_bonus = hist_sim * 15.0
            
            # Бонус за позицию
            position_bonus = 0.0
            pos_score = comparison.get("position_score", 0.0)
            if pos_score > 0.7:
                position_bonus = pos_score * 10.0
            
            # Бонус за порядок (если порядок близок к ожидаемому)
            order_bonus = 0.0
            docx_order = docx_image.get("order_in_document", docx_image.get("index", 9999))
            # Предполагаем, что изображения на странице идут в порядке появления в документе
            # Если порядок близок к ожидаемому, даем бонус
            if abs(docx_order - image_order_in_page) <= 2:
                order_bonus = 10.0
            
            total_score = visual_score + size_bonus + hist_bonus + position_bonus + order_bonus
            
            if total_score > best_score:
                best_score = total_score
                best_match = docx_image
                best_match_type = "visual"
                best_comparison_details = comparison
        
        # Приоритет 4: Похожесть по размеру
        elif best_match_type != "visual":
            perceptual_stats = comparison.get("perceptual_stats", {})
            size_similarity = perceptual_stats.get("size_similarity", 0.0)
            size_score = size_similarity * 50.0
            
            # Бонус за порядок
            docx_order = docx_image.get("order_in_document", docx_image.get("index", 9999))
            if abs(docx_order - image_order_in_page) <= 2:
                size_score += 15.0
            
            if size_score > best_score:
                best_score = size_score
                best_match = docx_image
                best_match_type = "size"
                best_comparison_details = comparison
    
    if best_match is not None:
        best_match["match_type"] = best_match_type
        best_match["match_score"] = best_score
        best_match["comparison_details"] = best_comparison_details
        return best_match
    
    return None


def parse_docx_hybrid(docx_path: Path, output_dir: Optional[Path] = None, temp_pdf_path: Optional[Path] = None) -> Dict[str, Any]:
    """
    Парсит DOCX используя комбинированный подход: PDF для координат + DOCX для точного текста.
    
    Args:
        docx_path: Путь к DOCX файлу
        output_dir: Директория для сохранения результатов
        temp_pdf_path: Опциональный путь к временному PDF (если None - создается внутри)
    
    Returns:
        Словарь с результатами парсинга
    """
    print("=" * 80)
    print("КОМБИНИРОВАННЫЙ ПАЙПЛАЙН DOCX (PDF + DOCX)")
    print("=" * 80)
    print()
    
    # Шаг 1: Конвертируем DOCX в PDF (если не передан)
    temp_pdf_dir = None
    created_temp_pdf = False
    
    if temp_pdf_path is None:
        print("Шаг 1: Конвертация DOCX → PDF...")
        temp_pdf_dir = tempfile.mkdtemp()
        temp_pdf_path = Path(temp_pdf_dir) / f"{docx_path.stem}_temp.pdf"
        created_temp_pdf = True
        convert_docx_to_pdf(docx_path, temp_pdf_path)
        print(f"  ✓ PDF создан: {temp_pdf_path}")
    else:
        print(f"Шаг 1: Использование существующего PDF: {temp_pdf_path}")
    
    try:
        
        # Шаг 2: Извлекаем данные из DOCX
        print("Шаг 2: Извлечение данных из DOCX...")
        docx_paragraphs = extract_all_text_from_docx(docx_path)
        docx_tables = extract_tables_from_docx(docx_path)
        # Не извлекаем изображения из DOCX - используем только из Dots.OCR
        print(f"  Параграфов: {len(docx_paragraphs)}")
        print(f"  Таблиц: {len(docx_tables)}")
        
        # Шаг 3: Layout detection через Dots.OCR (используем PDF)
        print("Шаг 3: Layout detection через Dots.OCR (из PDF)...")
        renderer = PdfPageRenderer(render_scale=RENDER_SCALE, optimize_for_ocr=True)
        pdf_document = fitz.open(str(temp_pdf_path))
        total_pages = len(pdf_document)
        pdf_document.close()
        
        print(f"  Всего страниц: {total_pages}")
        
        all_layout_elements = []
        
        # Сохраняем отрендеренные изображения страниц для последующего извлечения изображений
        rendered_pages: Dict[int, Image.Image] = {}
        
        for page_num in tqdm(range(total_pages), desc="Layout detection", unit="страница"):
            try:
                original_image, optimized_image = renderer.render_page(
                    temp_pdf_path, page_num, return_original=True
                )
                
                # Сохраняем отрендеренное изображение для извлечения изображений
                rendered_pages[page_num] = original_image
                
                # Layout detection
                layout_cells, raw_response, success = process_layout_detection(
                    image=optimized_image,
                    origin_image=original_image,
                )
                
                if not success or not layout_cells:
                    continue
                
                # Добавляем номер страницы и ссылку на отрендеренное изображение
                for element in layout_cells:
                    element["page_num"] = page_num
                    element["_rendered_page_image"] = original_image  # Временная ссылка для извлечения изображений
                
                all_layout_elements.extend(layout_cells)
                
            except Exception as e:
                print(f"  Ошибка при обработке страницы {page_num + 1}: {e}")
                continue
        
        print(f"  Найдено {len(all_layout_elements)} элементов layout")
        
        # Шаг 4: Извлекаем текст из PDF по координатам и сопоставляем с DOCX
        print("Шаг 4: Извлечение текста из PDF и сопоставление с DOCX...")
        
        for element in tqdm(all_layout_elements, desc="Сопоставление текста", unit="элемент", leave=False):
            category = element.get("category", "")
            bbox = element.get("bbox", [])
            page_num = element.get("page_num", 0)
            
            if category in ["Text", "Section-header", "Title", "Caption"]:
                # Извлекаем текст из PDF по координатам
                pdf_text = extract_text_from_pdf_by_bbox(
                    temp_pdf_path, bbox, page_num, RENDER_SCALE
                )
                element["pdf_text"] = pdf_text
                
                # Сопоставляем с текстом из DOCX
                matched_para = match_pdf_text_with_docx_text(pdf_text, docx_paragraphs)
                if matched_para:
                    element["docx_text"] = matched_para["text"]
                    element["docx_paragraph"] = matched_para
                else:
                    element["docx_text"] = pdf_text  # Fallback на PDF текст
        
        # Шаг 5: Обработка таблиц (отталкиваемся от DOCX, а не от OCR)
        print("Шаг 5: Обработка таблиц (отталкиваемся от DOCX)...")
        
        # Получаем все таблицы из OCR для поиска
        ocr_table_elements = [e for e in all_layout_elements if e.get("category") == "Table"]
        print(f"  Найдено таблиц в Dots.OCR: {len(ocr_table_elements)}")
        print(f"  Найдено таблиц в DOCX: {len(docx_tables)}")
        
        # Создаем директорию для сохранения сравнений таблиц (если output_dir указан)
        tables_comparison_dir = None
        if output_dir:
            tables_comparison_dir = output_dir / "tables_comparison"
            tables_comparison_dir.mkdir(parents=True, exist_ok=True)
        
        # Список для хранения обработанных таблиц (только те, что есть в DOCX)
        processed_table_elements = []
        
        # Флаг для отслеживания, какие таблицы из OCR уже использованы
        used_ocr_tables = set()
        
        # Обрабатываем таблицы из DOCX
        for docx_table_idx, docx_table in enumerate(tqdm(docx_tables, desc="Обработка таблиц из DOCX", unit="таблица", leave=False)):
            # Ищем соответствующую таблицу в результатах OCR
            ocr_table_element = find_ocr_table_for_docx_table(
                docx_table, ocr_table_elements, used_ocr_tables, docx_paragraphs, temp_pdf_path
            )
            
            if ocr_table_element:
                # Нашли соответствие в OCR
                used_ocr_tables.add(id(ocr_table_element))
                
                table_bbox = ocr_table_element.get("bbox", [])
                page_num = ocr_table_element.get("page_num", 0)
                
                # Вычисляем схожесть содержимого для верификации
                ocr_table_text = extract_table_text_from_pdf(temp_pdf_path, table_bbox, page_num)
                content_similarity = calculate_table_content_similarity(
                    docx_table, ocr_table_text, temp_pdf_path, table_bbox, page_num
                )
                
                # Получаем отрендеренное изображение страницы
                rendered_page_image = ocr_table_element.get("_rendered_page_image")
                
                # Удаляем временную ссылку
                if "_rendered_page_image" in ocr_table_element:
                    del ocr_table_element["_rendered_page_image"]
                
                # Извлекаем изображение таблицы из PDF (из Dots.OCR)
                table_image = extract_image_from_pdf_by_bbox(
                    temp_pdf_path, table_bbox, page_num, RENDER_SCALE, rendered_page_image
                )
                
                # Форматируем таблицу в markdown
                table_md = []
                for row in docx_table["data"]:
                    table_md.append("| " + " | ".join(str(cell) for cell in row) + " |")
                
                # Создаем элемент таблицы с данными из DOCX
                table_element = {
                    "category": "Table",
                    "bbox": table_bbox,
                    "page_num": page_num,
                    "docx_table": docx_table,
                    "table_markdown": "\n".join(table_md),
                    "source": "docx_with_ocr_coords",
                    "content_similarity": round(content_similarity, 3),  # Схожесть содержимого
                    "match_method": "content_based" if content_similarity > 0.3 else "position_based",
                }
                
                # Сохраняем сравнение таблиц
                if tables_comparison_dir and table_image:
                    try:
                        # Сохраняем изображение таблицы из Dots.OCR
                        ocr_table_path = tables_comparison_dir / f"table_{docx_table_idx + 1}_from_dots_ocr_page_{page_num + 1}.png"
                        table_image.save(ocr_table_path, "PNG")
                        table_element["ocr_table_image_path"] = str(ocr_table_path.relative_to(tables_comparison_dir.parent))
                        
                        # Сохраняем структуру таблицы из DOCX
                        docx_table_path = tables_comparison_dir / f"table_{docx_table_idx + 1}_from_docx_page_{page_num + 1}.txt"
                        with open(docx_table_path, "w", encoding="utf-8") as f:
                            f.write(f"Таблица из DOCX (индекс: {docx_table.get('index', 'N/A')})\n")
                            rows_count = docx_table.get('rows_count', docx_table.get('rows', 0))
                            cols_count = docx_table.get('cols_count', docx_table.get('cols', 0))
                            f.write(f"Размеры: {rows_count} строк × {cols_count} столбцов\n")
                            if docx_table.get('estimated_page'):
                                f.write(f"Приблизительная страница (из XML): {docx_table['estimated_page']}\n")
                            if table_element.get('content_similarity') is not None:
                                f.write(f"Схожесть содержимого с OCR: {table_element['content_similarity']:.3f}\n")
                                f.write(f"Метод сопоставления: {table_element.get('match_method', 'unknown')}\n")
                            f.write("=" * 80 + "\n\n")
                            f.write("Структура таблицы (Markdown):\n")
                            f.write(table_element.get("table_markdown", ""))
                            f.write("\n\n")
                            f.write("=" * 80 + "\n\n")
                            f.write("Текст из OCR (для сравнения):\n")
                            f.write(ocr_table_text[:1000] + ("..." if len(ocr_table_text) > 1000 else ""))
                            f.write("\n\n")
                            f.write("=" * 80 + "\n\n")
                            f.write("Структура таблицы (JSON):\n")
                            f.write(json.dumps(docx_table, ensure_ascii=False, indent=2, default=str))
                        table_element["docx_table_path"] = str(docx_table_path.relative_to(tables_comparison_dir.parent))
                        
                        # Создаем визуальное сравнение
                        comparison_image = create_table_comparison(
                            table_image, docx_table, docx_table_idx + 1, page_num + 1
                        )
                        comparison_path = tables_comparison_dir / f"table_{docx_table_idx + 1}_comparison_page_{page_num + 1}.png"
                        comparison_image.save(comparison_path, "PNG")
                        table_element["table_comparison_path"] = str(comparison_path.relative_to(tables_comparison_dir.parent))
                    except Exception as e:
                        print(f"  Предупреждение: не удалось сохранить сравнение таблицы {docx_table_idx + 1}: {e}")
                
                processed_table_elements.append(table_element)
            else:
                # Таблица из DOCX не найдена в OCR - создаем элемент без координат
                print(f"  Предупреждение: таблица {docx_table_idx + 1} из DOCX не найдена в результатах OCR")
                table_md = []
                for row in docx_table["data"]:
                    table_md.append("| " + " | ".join(str(cell) for cell in row) + " |")
                
                table_element = {
                    "category": "Table",
                    "bbox": None,
                    "page_num": docx_table.get("estimated_page", 1) - 1 if docx_table.get("estimated_page") else None,
                    "docx_table": docx_table,
                    "table_markdown": "\n".join(table_md),
                    "source": "docx_only",
                    "match_status": "not_found_in_ocr",
                }
                processed_table_elements.append(table_element)
        
        # Обрабатываем таблицы из OCR, которые не были сопоставлены с DOCX
        unmatched_ocr_tables = [
            t for t in ocr_table_elements 
            if id(t) not in used_ocr_tables
        ]
        
        if unmatched_ocr_tables:
            print(f"  Предупреждение: {len(unmatched_ocr_tables)} таблиц найдено в OCR, но не найдено в DOCX")
            # Пытаемся найти их в DOCX по содержимому
            for ocr_table in unmatched_ocr_tables:
                table_bbox = ocr_table.get("bbox", [])
                page_num = ocr_table.get("page_num", 0)
                
                if not table_bbox or len(table_bbox) < 4:
                    continue
                
                # Извлекаем текст из области таблицы
                ocr_table_text = extract_table_text_from_pdf(temp_pdf_path, table_bbox, page_num)
                
                # Ищем похожую таблицу в DOCX по содержимому
                best_docx_match = None
                best_similarity = 0.0
                
                for docx_table in docx_tables:
                    # Пропускаем уже сопоставленные таблицы
                    if any(e.get("docx_table", {}).get("index") == docx_table.get("index") 
                           for e in processed_table_elements if e.get("source") == "docx_with_ocr_coords"):
                        continue
                    
                    similarity = calculate_table_content_similarity(
                        docx_table, ocr_table_text, temp_pdf_path, table_bbox, page_num
                    )
                    
                    if similarity > best_similarity:
                        best_similarity = similarity
                        best_docx_match = docx_table
                
                # Если нашли хорошее совпадение, используем его
                if best_docx_match and best_similarity > 0.3:
                    print(f"    Найдено совпадение для таблицы OCR (страница {page_num + 1}) с DOCX таблицей {best_docx_match.get('index')} (схожесть: {best_similarity:.3f})")
                    # Создаем элемент таблицы с данными из DOCX
                    table_md = []
                    for row in best_docx_match["data"]:
                        table_md.append("| " + " | ".join(str(cell) for cell in row) + " |")
                    
                    table_element = {
                        "category": "Table",
                        "bbox": table_bbox,
                        "page_num": page_num,
                        "docx_table": best_docx_match,
                        "table_markdown": "\n".join(table_md),
                        "source": "docx_with_ocr_coords",
                        "content_similarity": round(best_similarity, 3),
                        "match_method": "content_based_retry",
                    }
                    processed_table_elements.append(table_element)
                    used_ocr_tables.add(id(ocr_table))
        
        # Обновляем all_layout_elements: удаляем неиспользованные таблицы из OCR
        # и добавляем обработанные таблицы
        all_layout_elements = [
            e for e in all_layout_elements 
            if e.get("category") != "Table" or id(e) in used_ocr_tables
        ]
        
        # Добавляем обработанные таблицы
        all_layout_elements.extend(processed_table_elements)
        
        print(f"  Обработано таблиц из DOCX: {len(processed_table_elements)}")
        matched_in_ocr = len([e for e in processed_table_elements if e.get("source") == "docx_with_ocr_coords"])
        print(f"  Найдено в OCR: {matched_in_ocr}")
        if len(processed_table_elements) > matched_in_ocr:
            print(f"  Предупреждение: {len(processed_table_elements) - matched_in_ocr} таблиц из DOCX не найдено в OCR")
        if tables_comparison_dir:
            print(f"  Сравнения таблиц сохранены в: {tables_comparison_dir}")
        
        # Шаг 6: Обработка изображений (только из Dots.OCR, без сопоставления с DOCX)
        print("Шаг 6: Обработка изображений из Dots.OCR...")
        image_elements = [e for e in all_layout_elements if e.get("category") == "Picture"]
        
        # Создаем директорию для сохранения изображений (если output_dir указан)
        images_dir = None
        if output_dir:
            images_dir = output_dir / "images_from_ocr"
            images_dir.mkdir(parents=True, exist_ok=True)
        
        # Обрабатываем изображения из Dots.OCR
        for img_idx, image_element in enumerate(tqdm(image_elements, desc="Обработка изображений", unit="изображение", leave=False)):
            image_bbox = image_element.get("bbox", [])
            page_num = image_element.get("page_num", 0)
            
            # Получаем отрендеренное изображение страницы (если доступно)
            rendered_page_image = image_element.get("_rendered_page_image")
            
            # Удаляем временную ссылку
            if "_rendered_page_image" in image_element:
                del image_element["_rendered_page_image"]
            
            # Извлекаем изображение из PDF (из Dots.OCR)
            pdf_image = extract_image_from_pdf_by_bbox(
                temp_pdf_path, image_bbox, page_num, RENDER_SCALE, rendered_page_image
            )
            
            if pdf_image is None:
                continue
            
            # Сохраняем изображение в base64 для метаданных
            img_bytes_io = BytesIO()
            pdf_image.save(img_bytes_io, format="PNG")
            img_bytes = img_bytes_io.getvalue()
            image_element["image_base64"] = base64.b64encode(img_bytes).decode("utf-8")
            
            # Сохраняем метаданные изображения
            image_element["image_metadata"] = {
                "width": pdf_image.width,
                "height": pdf_image.height,
                "format": "PNG",
                "source": "dots_ocr",
                "page_num": page_num,
                "bbox": image_bbox,
            }
            
            # Сохраняем изображение в файл (если указана директория)
            if images_dir:
                try:
                    image_path = images_dir / f"image_{img_idx + 1}_page_{page_num + 1}.png"
                    pdf_image.save(image_path, "PNG")
                    image_element["image_path"] = str(image_path.relative_to(images_dir.parent))
                except Exception as e:
                    print(f"  Предупреждение: не удалось сохранить изображение {img_idx + 1}: {e}")
        
        print(f"  Найдено изображений в Dots.OCR: {len(image_elements)}")
        if images_dir:
            print(f"  Изображения сохранены в: {images_dir}")
        
        # Шаг 7: Анализ уровней заголовков (как в PDF)
        print("Шаг 7: Анализ уровней заголовков...")
        # Используем логику из PDF пайплайна
        analyzed_elements = analyze_header_levels_hybrid(all_layout_elements)
        headers_count = len([e for e in analyzed_elements if e.get("category") == "Section-header"])
        print(f"  Найдено {headers_count} заголовков")
        
        # Шаг 8: Построение иерархии
        print("Шаг 8: Построение иерархии...")
        hierarchy = build_hierarchy_from_section_headers(analyzed_elements)
        print(f"  Создано {len(hierarchy)} секций")
        
        # Шаг 9: Создание элементов
        elements = []
        for section in hierarchy:
            header = section["header"]
            if header.get("category") == "Section-header":
                level = header.get("header_level", 1)
                elements.append({
                    "type": f"HEADER_{level}",
                    "content": header.get("docx_text", header.get("pdf_text", "")),
                    "metadata": {
                        "page_num": header.get("page_num", 0),
                        "bbox": header.get("bbox", []),
                        "docx_paragraph": header.get("docx_paragraph"),
                    }
                })
            
            for child in section["children"]:
                if child["category"] == "Text":
                    elements.append({
                        "type": "TEXT",
                        "content": child.get("docx_text", child.get("pdf_text", "")),
                        "metadata": {
                            "page_num": child.get("page_num", 0),
                            "bbox": child.get("bbox", []),
                            "docx_paragraph": child.get("docx_paragraph"),
                        }
                    })
                elif child["category"] == "Table":
                    elements.append({
                        "type": "TABLE",
                        "content": child.get("table_markdown", ""),
                        "metadata": {
                            "page_num": child.get("page_num", 0),
                            "bbox": child.get("bbox", []),
                            "docx_table": child.get("docx_table"),
                        }
                    })
                elif child["category"] == "Picture":
                    elements.append({
                        "type": "IMAGE",
                        "content": "",  # Изображения хранятся в метаданных
                        "metadata": {
                            "page_num": child.get("page_num", 0),
                            "bbox": child.get("bbox", []),
                            "image_metadata": child.get("image_metadata"),
                            "image_base64": child.get("image_base64"),
                            "image_path": child.get("image_path"),
                        }
                    })
        
        # Результат
        result = {
            "source": str(docx_path),
            "format": "DOCX",
            "elements": elements,
            "metadata": {
                "parser": "docx_hybrid",
                "status": "completed",
                "processing_method": "hybrid_pdf_coords_docx_text",
                "total_pages": total_pages,
                "sections_count": len(hierarchy),
                "headers_count": headers_count,
                "elements_count": len(elements),
                "tables_count": len([e for e in elements if e["type"] == "TABLE"]),
                "images_count": len([e for e in elements if e["type"] == "IMAGE"]),
            },
            "hierarchy": hierarchy,
        }
        
        # Сохранение результатов
        if output_dir:
            output_dir.mkdir(parents=True, exist_ok=True)
            output_file = output_dir / f"{docx_path.stem}_hybrid_pipeline.json"
            with open(output_file, "w", encoding="utf-8") as f:
                json.dump(result, f, ensure_ascii=False, indent=2, default=str)
            print(f"\nРезультат сохранен в: {output_file}")
        
        return result
    
    finally:
        # Удаляем временный PDF только если мы его создавали
        if created_temp_pdf and temp_pdf_path and temp_pdf_path.exists():
            try:
                temp_pdf_path.unlink()
                if temp_pdf_dir:
                    shutil.rmtree(temp_pdf_dir, ignore_errors=True)
            except:
                pass


def determine_header_level_hybrid(
    text: str,
    previous_headers: List[Dict[str, Any]],
    docx_paragraph: Optional[Dict[str, Any]] = None,
) -> int:
    """
    Определяет уровень заголовка (аналог из PDF пайплайна).
    
    Args:
        text: Текст заголовка
        previous_headers: Список предыдущих заголовков
        docx_paragraph: Данные параграфа из DOCX
    
    Returns:
        Уровень заголовка (1-6)
    """
    # Приоритет 1: Анализ нумерации
    if re.match(r'^\d+\s+[A-Z]', text):
        return 1
    if re.match(r'^\d+\.\d+\s+', text):
        return 2
    if re.match(r'^\d+\.\d+\.\d+\s+', text):
        return 3
    if re.match(r'^\d+\.\d+\.\d+\.\d+\s+', text):
        return 4
    
    # Приоритет 2: Если есть данные из DOCX, используем стиль (но не доверяем полностью)
    if docx_paragraph:
        style = docx_paragraph.get("style", "")
        if style.startswith("Heading"):
            level = int(style.split()[-1])
            if previous_headers:
                last_level = previous_headers[-1].get("level", 1)
                if abs(level - last_level) > 2:
                    return min(6, last_level + 1)
            return min(level, 6)
    
    # Приоритет 3: Сравнение размера шрифта с предыдущими
    if docx_paragraph and previous_headers:
        font_size_str = docx_paragraph.get("formatting", {}).get("font_size")
        if font_size_str:
            try:
                current_size = float(font_size_str.replace("pt", ""))
                for header in reversed(previous_headers):
                    header_para = header.get("docx_paragraph")
                    if header_para:
                        last_font_size_str = header_para.get("formatting", {}).get("font_size")
                        if last_font_size_str:
                            last_size = float(last_font_size_str.replace("pt", ""))
                            last_level = header.get("level", 1)
                            if current_size >= last_size + 2:
                                return max(1, last_level - 1)
                            elif current_size <= last_size - 2:
                                return min(6, last_level + 1)
                            else:
                                return last_level
            except (ValueError, AttributeError):
                pass
    
    # По умолчанию
    if previous_headers:
        return min(6, previous_headers[-1].get("level", 1) + 1)
    return 1


def analyze_header_levels_hybrid(layout_elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Анализирует уровни заголовков (аналог из PDF пайплайна).
    
    Args:
        layout_elements: Список элементов layout
    
    Returns:
        Список элементов с определенными уровнями заголовков
    """
    analyzed_elements = []
    previous_headers = []
    last_numbered_level = None
    
    for element in layout_elements:
        if element.get("category") == "Section-header":
            text = element.get("docx_text", element.get("pdf_text", ""))
            docx_paragraph = element.get("docx_paragraph")
            
            level = determine_header_level_hybrid(text, previous_headers, docx_paragraph)
            
            # Проверяем наличие нумерации
            if re.match(r'^\d+', text):
                last_numbered_level = level
            
            header_info = {
                "level": level,
                "text": text,
                "docx_paragraph": docx_paragraph,
            }
            previous_headers.append(header_info)
            
            element["header_level"] = level
            element["last_numbered_level"] = last_numbered_level
        
        analyzed_elements.append(element)
    
    return analyzed_elements


def build_hierarchy_from_section_headers(layout_elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Строит иерархию элементов, группируя их по Section-header (как в PDF).
    
    Args:
        layout_elements: Список элементов layout
    
    Returns:
        Список секций с заголовками и дочерними элементами
    """
    sections = []
    current_section = None
    
    for element in layout_elements:
        if element.get("category") == "Section-header":
            if current_section:
                sections.append(current_section)
            
            current_section = {
                "header": element,
                "children": []
            }
        else:
            if current_section:
                current_section["children"].append(element)
            else:
                if not sections or sections[-1].get("header", {}).get("text") != "Начало документа":
                    current_section = {
                        "header": {
                            "text": "Начало документа",
                            "level": 0,
                            "category": "Title"
                        },
                        "children": []
                    }
                    sections.append(current_section)
                else:
                    current_section = sections[-1]
                
                current_section["children"].append(element)
    
    if current_section:
        sections.append(current_section)
    
    return sections


def process_docx_file_hybrid(
    docx_path: Path,
    output_base_dir: Path,
) -> Dict[str, Any]:
    """
    Обрабатывает один DOCX файл используя комбинированный подход.
    
    Args:
        docx_path: Путь к DOCX файлу
        output_base_dir: Базовая директория для сохранения результатов
        render_scale: Масштаб рендеринга
    
    Returns:
        Словарь с результатами обработки
    """
    import time
    
    print(f"\n{'='*80}")
    print(f"Обработка: {docx_path.name}")
    print(f"{'='*80}")
    
    start_time = time.time()
    
    try:
        # Создаем директорию для результатов
        docx_output_dir = output_base_dir / docx_path.stem
        docx_output_dir.mkdir(parents=True, exist_ok=True)
        
        # Создаем временный PDF для парсинга (будет использован и для layout изображений)
        temp_pdf_dir = tempfile.mkdtemp()
        temp_pdf_path = Path(temp_pdf_dir) / f"{docx_path.stem}_temp.pdf"
        convert_docx_to_pdf(docx_path, temp_pdf_path)
        
        if not temp_pdf_path.exists():
            print(f"Ошибка: не удалось создать временный PDF для обработки")
            return {"error": "Failed to create temporary PDF"}
        
        try:
            # Парсинг (передаем output_dir для сохранения изображений сравнения)
            result = parse_docx_hybrid(docx_path, docx_output_dir, temp_pdf_path=temp_pdf_path)
            
            # Сохраняем полный результат
            result_file = docx_output_dir / "result.json"
            with open(result_file, "w", encoding="utf-8") as f:
                json.dump(result, f, ensure_ascii=False, indent=2, default=str)
            
            # Сохраняем структуру
            _save_structure_hybrid(result, docx_output_dir)
            
            # Сохраняем страницы с layout (используем тот же PDF)
            saved_pages = _save_full_pages_with_layout_hybrid(
                docx_path, result, docx_output_dir, temp_pdf_path=temp_pdf_path
            )
        finally:
            # Удаляем временный PDF только после всех операций
            if temp_pdf_path.exists():
                try:
                    temp_pdf_path.unlink()
                    if temp_pdf_path.parent != Path(tempfile.gettempdir()):
                        shutil.rmtree(temp_pdf_path.parent, ignore_errors=True)
                except:
                    pass
        
        processing_time = time.time() - start_time
        
        # Статистика
        stats = {
            "processing_time_seconds": processing_time,
            "total_elements": result.get("metadata", {}).get("elements_count", 0),
            "headers": result.get("metadata", {}).get("headers_count", 0),
            "sections": result.get("metadata", {}).get("sections_count", 0),
            "tables": result.get("metadata", {}).get("tables_count", 0),
            "images": result.get("metadata", {}).get("images_count", 0),
            "total_pages": result.get("metadata", {}).get("total_pages", 0),
            "saved_pages_with_layout": saved_pages,
        }
        
        stats_file = docx_output_dir / "stats.json"
        with open(stats_file, "w", encoding="utf-8") as f:
            json.dump(stats, f, indent=2, ensure_ascii=False)
        
        print(f"✓ Успешно обработан за {processing_time:.2f} сек")
        print(f"  Элементов: {stats['total_elements']}")
        print(f"  Заголовков: {stats['headers']}")
        print(f"  Секций: {stats['sections']}")
        print(f"  Таблиц: {stats['tables']}")
        print(f"  Изображений: {stats['images']}")
        print(f"  Страниц: {stats['total_pages']}")
        print(f"  Результаты сохранены в: {docx_output_dir}")
        
        return {
            "success": True,
            "processing_time": processing_time,
            "stats": stats,
            "output_dir": str(docx_output_dir),
        }
    
    except Exception as e:
        processing_time = time.time() - start_time
        error_msg = f"Ошибка при обработке {docx_path.name}: {e}"
        print(f"✗ {error_msg}")
        import traceback
        traceback.print_exc()
        
        return {
            "success": False,
            "processing_time": processing_time,
            "error": str(e),
        }


def _save_structure_hybrid(result: Dict[str, Any], output_dir: Path) -> None:
    """Сохраняет структуру документа."""
    structure_file = output_dir / "structure.json"
    
    structure = {
        "source": result.get("source"),
        "format": result.get("format"),
        "metadata": result.get("metadata", {}),
        "hierarchy": result.get("hierarchy", []),
        "elements": [
            {
                "type": e.get("type"),
                "content": e.get("content", "")[:200] + "..." if len(e.get("content", "")) > 200 else e.get("content", ""),
                "metadata": {k: v for k, v in e.get("metadata", {}).items() if k != "image_base64"}  # Не сохраняем base64 в структуре
            }
            for e in result.get("elements", [])
        ]
    }
    
    with open(structure_file, "w", encoding="utf-8") as f:
        json.dump(structure, f, ensure_ascii=False, indent=2, default=str)


def _save_full_pages_with_layout_hybrid(
    docx_path: Path,
    result: Dict[str, Any],
    output_dir: Path,
    temp_pdf_path: Optional[Path] = None,
) -> int:
    """
    Сохраняет полные сканы страниц с нарисованными bbox для всех элементов layout.
    
    Args:
        docx_path: Путь к DOCX файлу
        result: Результат парсинга
        output_dir: Директория для сохранения
        temp_pdf_path: Путь к временному PDF (если None - создается)
    
    Returns:
        Количество сохраненных страниц
    """
    from PIL import ImageDraw, ImageFont
    
    pages_dir = output_dir / "pages_with_layout"
    pages_dir.mkdir(exist_ok=True)
    
    # Группируем элементы по страницам
    elements_by_page: Dict[int, List[Dict[str, Any]]] = defaultdict(list)
    elements = result.get("elements", [])
    
    if not elements:
        print("  Предупреждение: нет элементов для отображения на страницах")
        return 0
    
    for element in elements:
        page_num = element.get("metadata", {}).get("page_num", 0)
        bbox = element.get("metadata", {}).get("bbox", [])
        if bbox and len(bbox) >= 4:
            elements_by_page[page_num].append(element)
    
    if not elements_by_page:
        print("  Предупреждение: нет элементов с bbox для отображения")
        return 0
    
    total_pages = result.get("metadata", {}).get("total_pages", 0)
    if total_pages == 0:
        # Пытаемся определить количество страниц из максимального номера страницы
        if elements_by_page:
            total_pages = max(elements_by_page.keys()) + 1
        else:
            print("  Предупреждение: не удалось определить количество страниц")
            return 0
    
    saved_count = 0
    
    # Создаем временный PDF если нужно
    if temp_pdf_path is None:
        temp_pdf_dir = tempfile.mkdtemp()
        temp_pdf_path = Path(temp_pdf_dir) / f"{docx_path.stem}_temp.pdf"
        print(f"  Создание временного PDF для рендеринга страниц...")
        convert_docx_to_pdf(docx_path, temp_pdf_path)
        if not temp_pdf_path.exists():
            print(f"  Ошибка: не удалось создать временный PDF")
            return 0
    
    # Используем тот же render_scale, что и при layout detection
    # Координаты bbox относятся к изображению с этим масштабом
    renderer = PdfPageRenderer(render_scale=RENDER_SCALE, optimize_for_ocr=False)
    
    for page_num in tqdm(sorted(elements_by_page.keys()), desc="Сохранение страниц с layout", unit="страница", leave=False):
        if page_num >= total_pages:
            continue
        
        try:
            # Рендерим страницу из PDF с тем же масштабом, что использовался для layout detection
            original_image, _ = renderer.render_page(
                temp_pdf_path, page_num, return_original=True
            )
            
            if original_image is None:
                print(f"  Предупреждение: не удалось отрендерить страницу {page_num + 1}")
                continue
            
            page_image = original_image.copy()
            
            # Рисуем bbox для всех элементов на странице
            # Координаты bbox уже в правильном масштабе (RENDER_SCALE)
            for element in elements_by_page[page_num]:
                element_type = element.get("type", "TEXT")
                bbox = element.get("metadata", {}).get("bbox", [])
                
                if not bbox or len(bbox) < 4:
                    continue
                
                color = _get_element_color_hybrid(element_type)
                label = f"{element_type}"
                
                page_image = _draw_bbox_on_full_page_hybrid(page_image, bbox, label, color)
            
            # Сохраняем страницу
            page_file = pages_dir / f"page_{page_num + 1}_with_layout.png"
            page_image.save(page_file, "PNG")
            saved_count += 1
            
        except Exception as e:
            print(f"Ошибка при сохранении страницы {page_num + 1}: {e}")
            continue
    
    # Удаляем временный PDF (только если мы его создавали)
    if temp_pdf_path and temp_pdf_path.exists():
        try:
            temp_pdf_path.unlink()
            if temp_pdf_path.parent != Path(tempfile.gettempdir()):
                shutil.rmtree(temp_pdf_path.parent, ignore_errors=True)
        except:
            pass
    
    if saved_count > 0:
        print(f"  Сохранено страниц с layout: {saved_count} в {pages_dir}")
    else:
        print(f"  Предупреждение: не было сохранено ни одной страницы с layout")
    
    return saved_count


def _get_element_color_hybrid(element_type: str) -> str:
    """Возвращает цвет для типа элемента."""
    color_map = {
        "TEXT": "green",
        "IMAGE": "magenta",
        "CAPTION": "orange",
        "HEADER_1": "cyan",
        "HEADER_2": "cyan",
        "HEADER_3": "cyan",
        "HEADER_4": "cyan",
        "HEADER_5": "cyan",
        "HEADER_6": "cyan",
        "TITLE": "red",
        "TABLE": "pink",
        "FORMULA": "gray",
        "LIST_ITEM": "blue",
    }
    return color_map.get(element_type, "red")


def create_table_comparison(
    ocr_table_image: Image.Image,
    docx_table: Dict[str, Any],
    table_num: int,
    page_num: int
) -> Image.Image:
    """
    Создает визуальное сравнение таблицы из Dots.OCR и структуры из DOCX.
    
    Args:
        ocr_table_image: Изображение таблицы из Dots.OCR (PDF)
        docx_table: Структура таблицы из DOCX
        table_num: Номер таблицы
        page_num: Номер страницы
    
    Returns:
        PIL Image с side-by-side сравнением
    """
    from PIL import ImageDraw, ImageFont
    
    # Приводим изображение таблицы к разумному размеру для отображения
    max_display_width = 800
    max_display_height = 600
    
    if ocr_table_image.width > max_display_width or ocr_table_image.height > max_display_height:
        scale = min(max_display_width / ocr_table_image.width, max_display_height / ocr_table_image.height)
        new_width = int(ocr_table_image.width * scale)
        new_height = int(ocr_table_image.height * scale)
        ocr_table_image = ocr_table_image.resize((new_width, new_height), Image.Resampling.LANCZOS)
    
    # Создаем текстовое представление таблицы из DOCX
    table_text_lines = []
    table_text_lines.append(f"Таблица из DOCX (индекс: {docx_table.get('index', 'N/A')})")
    
    # Размеры (используем rows_count/cols_count из XML, если есть)
    rows_count = docx_table.get('rows_count', docx_table.get('rows', 0))
    cols_count = docx_table.get('cols_count', docx_table.get('cols', 0))
    table_text_lines.append(f"Размеры: {rows_count} строк × {cols_count} столбцов")
    
    # Добавляем информацию из XML
    if docx_table.get('style'):
        table_text_lines.append(f"Стиль: {docx_table['style']}")
    if docx_table.get('xml_position') is not None:
        table_text_lines.append(f"Позиция в XML: {docx_table['xml_position']}")
    if docx_table.get('estimated_page') is not None:
        table_text_lines.append(f"Приблизительная страница (из XML): {docx_table['estimated_page']}")
    merged_cells = docx_table.get('merged_cells', [])
    if merged_cells:
        table_text_lines.append(f"Объединенных ячеек: {len(merged_cells)}")
    
    table_text_lines.append("")
    table_text_lines.append("Структура:")
    table_text_lines.append("=" * 60)
    
    # Добавляем данные таблицы
    table_data = docx_table.get("data", [])
    for row_idx, row in enumerate(table_data[:20]):  # Ограничиваем 20 строками для отображения
        row_text = " | ".join(str(cell)[:30] for cell in row)  # Ограничиваем длину ячеек
        table_text_lines.append(row_text)
    
    if len(table_data) > 20:
        table_text_lines.append(f"... и еще {len(table_data) - 20} строк")
    
    # Добавляем информацию об объединенных ячейках
    if merged_cells:
        table_text_lines.append("")
        table_text_lines.append("Объединенные ячейки:")
        for merged in merged_cells[:5]:  # Показываем первые 5
            table_text_lines.append(f"  Строка {merged['row']}, Столбец {merged['col']}, Colspan: {merged.get('colspan', 1)}")
        if len(merged_cells) > 5:
            table_text_lines.append(f"  ... и еще {len(merged_cells) - 5}")
    
    # Вычисляем размеры для текстовой части
    try:
        font = ImageFont.truetype("arial.ttf", 14)
        font_small = ImageFont.truetype("arial.ttf", 12)
    except:
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 14)
            font_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 12)
        except:
            font = ImageFont.load_default()
            font_small = ImageFont.load_default()
    
    # Вычисляем размеры текста
    text_width = 600
    line_height = 20
    text_height = len(table_text_lines) * line_height + 40
    
    # Создаем изображение для текста
    text_image = Image.new("RGB", (text_width, text_height), "white")
    draw_text = ImageDraw.Draw(text_image)
    
    y_pos = 10
    for line in table_text_lines:
        if len(line) > 80:
            # Разбиваем длинные строки
            words = line.split()
            current_line = ""
            for word in words:
                if len(current_line + word) < 80:
                    current_line += word + " "
                else:
                    if current_line:
                        draw_text.text((10, y_pos), current_line.strip(), fill="black", font=font_small)
                        y_pos += line_height
                    current_line = word + " "
            if current_line:
                draw_text.text((10, y_pos), current_line.strip(), fill="black", font=font_small)
                y_pos += line_height
        else:
            draw_text.text((10, y_pos), line, fill="black", font=font_small if y_pos > 50 else font)
            y_pos += line_height
    
    # Создаем итоговое изображение (side-by-side)
    comparison_width = ocr_table_image.width + text_width + 40
    comparison_height = max(ocr_table_image.height, text_height) + 100
    
    comparison = Image.new("RGB", (comparison_width, comparison_height), "white")
    draw = ImageDraw.Draw(comparison)
    
    # Заголовок
    title = f"Сравнение таблицы {table_num} (страница {page_num})"
    draw.text((10, 10), title, fill="blue", font=font)
    
    # Подзаголовки
    draw.text((10, 35), "Изображение из Dots.OCR (PDF)", fill="gray", font=font_small)
    draw.text((ocr_table_image.width + 30, 35), "Структура из DOCX", fill="gray", font=font_small)
    
    # Вставляем изображения
    comparison.paste(ocr_table_image, (10, 60))
    comparison.paste(text_image, (ocr_table_image.width + 30, 60))
    
    return comparison


def create_image_comparison(
    ocr_image: Image.Image,
    docx_image: Image.Image,
    match_type: str = "unknown",
    match_score: float = 0.0,
    comparison_details: Optional[Dict[str, Any]] = None
) -> Image.Image:
    """
    Создает side-by-side сравнение изображений из Dots.OCR и DOCX.
    
    Args:
        ocr_image: Изображение из Dots.OCR (PDF)
        docx_image: Изображение из DOCX
        match_type: Тип совпадения ("exact", "visual", "size")
        match_score: Score совпадения
    
    Returns:
        PIL Image с side-by-side сравнением
    """
    from PIL import ImageDraw, ImageFont
    
    # Приводим изображения к одному размеру для сравнения
    max_height = max(ocr_image.height, docx_image.height)
    max_width = max(ocr_image.width, docx_image.width)
    
    # Создаем изображения одинакового размера
    ocr_resized = ocr_image.resize((max_width, max_height), Image.Resampling.LANCZOS)
    docx_resized = docx_image.resize((max_width, max_height), Image.Resampling.LANCZOS)
    
    # Определяем высоту заголовков (зависит от наличия comparison_details)
    header_height = 100
    if comparison_details:
        details_count = 0
        if comparison_details.get("perceptual_hash_distance") is not None:
            details_count += 1
        if comparison_details.get("size_similarity") is not None:
            details_count += 1
        header_height = 100 + (details_count * 20)
    
    # Создаем итоговое изображение (side-by-side)
    comparison_width = max_width * 2 + 40  # Два изображения + отступы
    comparison_height = max_height + header_height  # Изображения + заголовки
    
    comparison = Image.new("RGB", (comparison_width, comparison_height), "white")
    draw = ImageDraw.Draw(comparison)
    
    try:
        font_large = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 16)
        font_small = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 12)
    except:
        try:
            font_large = ImageFont.truetype("arial.ttf", 16)
            font_small = ImageFont.truetype("arial.ttf", 12)
        except:
            font_large = ImageFont.load_default()
            font_small = ImageFont.load_default()
    
    # Заголовки
    draw.text((10, 10), "Изображение из Dots.OCR (PDF)", fill="black", font=font_large)
    draw.text((max_width + 30, 10), "Изображение из DOCX", fill="black", font=font_large)
    
    # Информация о совпадении
    match_info = f"Match: {match_type} (score: {match_score:.2f})"
    draw.text((10, 35), match_info, fill="blue", font=font_small)
    
    # Дополнительная информация о сравнении
    if comparison_details:
        details_y = 55
        if comparison_details.get("perceptual_hash_distance") is not None:
            phash_info = f"pHash distance: {comparison_details['perceptual_hash_distance']}"
            draw.text((10, details_y), phash_info, fill="gray", font=font_small)
            details_y += 20
        if comparison_details.get("size_similarity") is not None:
            size_info = f"Size similarity: {comparison_details['size_similarity']:.2f}"
            draw.text((10, details_y), size_info, fill="gray", font=font_small)
    
    # Размеры изображений
    size_info_ocr = f"Size: {ocr_image.width}x{ocr_image.height}"
    size_info_docx = f"Size: {docx_image.width}x{docx_image.height}"
    
    # Определяем Y-позицию для размеров (зависит от наличия comparison_details)
    size_y = 55
    if comparison_details:
        # Если есть детали сравнения, размеры будут ниже
        details_count = 0
        if comparison_details.get("perceptual_hash_distance") is not None:
            details_count += 1
        if comparison_details.get("size_similarity") is not None:
            details_count += 1
        size_y = 55 + (details_count * 20)
    
    draw.text((10, size_y), size_info_ocr, fill="gray", font=font_small)
    draw.text((max_width + 30, size_y), size_info_docx, fill="gray", font=font_small)
    
    # Y-позиция для изображений
    image_y = size_y + 25
    
    # Вставляем изображения
    comparison.paste(ocr_resized, (10, image_y))
    comparison.paste(docx_resized, (max_width + 30, image_y))
    
    # Рисуем рамку
    draw.rectangle([5, 70, max_width + 15, max_height + 80], outline="black", width=2)
    draw.rectangle([max_width + 25, 70, comparison_width - 5, max_height + 80], outline="black", width=2)
    
    return comparison


def _draw_bbox_on_full_page_hybrid(image: Image.Image, bbox: List[float], label: str = "", color: str = "red") -> Image.Image:
    """Рисует bbox на полной странице."""
    from PIL import ImageDraw, ImageFont
    
    img_copy = image.copy()
    draw = ImageDraw.Draw(img_copy)
    
    if len(bbox) >= 4:
        x1, y1, x2, y2 = bbox[0], bbox[1], bbox[2], bbox[3]
        
        # Рисуем прямоугольник
        draw.rectangle([x1, y1, x2, y2], outline=color, width=2)
        
        # Добавляем подпись, если есть
        if label:
            try:
                font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 12)
            except:
                try:
                    font = ImageFont.truetype("arial.ttf", 12)
                except:
                    font = ImageFont.load_default()
            
            # Фон для текста
            text_bbox = draw.textbbox((x1, y1 - 15), label, font=font)
            text_bbox = (text_bbox[0] - 2, text_bbox[1] - 2, text_bbox[2] + 2, text_bbox[3] + 2)
            draw.rectangle(text_bbox, fill=color)
            draw.text((x1, y1 - 15), label, fill="white", font=font)
    
    return img_copy


def main():
    """Основная функция для обработки одного или нескольких DOCX файлов."""
    import time
    import sys
    
    # Пути к файлам
    test_folder = Path(r"E:\easy\documentor\documentor_langchain\experiments\pdf_text_extraction\test_folder")
    
    # Если передан аргумент командной строки - обрабатываем один файл
    if len(sys.argv) > 1:
        docx_file_path = Path(sys.argv[1])
        if not docx_file_path.exists():
            print(f"Ошибка: Файл не найден: {docx_file_path}")
            return
        
        # Директория для результатов
        output_dir = test_folder.parent / "results" / "documentor_docx_hybrid"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        print(f"Обработка одного файла: {docx_file_path.name}")
        print(f"Выходная директория: {output_dir}\n")
        
        result = process_docx_file_hybrid(docx_file_path, output_dir)
        
        if result.get("success", False):
            print(f"\n✓ Успешно обработан: {docx_file_path.name}")
            print(f"  Результаты сохранены в: {result.get('output_dir')}")
        else:
            print(f"\n✗ Ошибка при обработке: {result.get('error', 'Unknown error')}")
        
        return
    
    # Иначе обрабатываем все файлы
    docx_files = [
        test_folder / "Отчёт ГОСТ.docx",
        test_folder / "Отчёт НИР Хаухия АВ.docx",
        test_folder / "02_Отчет_Этап_2_полный_замечания_эксперта.docx",
        test_folder / "Диплом.docx",
    ]
    
    # Проверяем существование файлов
    existing_files = [f for f in docx_files if f.exists()]
    
    if not existing_files:
        print("Ошибка: Не найдено ни одного файла для обработки")
        return
    
    print(f"Найдено файлов для обработки: {len(existing_files)}")
    for f in existing_files:
        print(f"  - {f.name}")
    
    # Директория для результатов
    output_dir = test_folder.parent / "results" / "documentor_docx_hybrid"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"\nВыходная директория: {output_dir}")
    
    # Обрабатываем каждый файл
    results: List[Dict[str, Any]] = []
    total_start_time = time.time()
    
    for docx_file in tqdm(existing_files, desc="Обработка DOCX файлов", unit="файл"):
        result = process_docx_file_hybrid(docx_file, output_dir)
        results.append({
            "file": docx_file.name,
            **result
        })
    
    total_time = time.time() - total_start_time
    
    # Сохраняем общую статистику
    summary = {
        "total_files": len(existing_files),
        "successful": len([r for r in results if r.get("success", False)]),
        "failed": len([r for r in results if not r.get("success", False)]),
        "total_processing_time_seconds": total_time,
        "average_processing_time_seconds": total_time / len(existing_files) if existing_files else 0,
        "processing_method": "hybrid_pdf_coords_docx_text",
        "results": results,
    }
    
    summary_file = output_dir / "summary.json"
    with open(summary_file, "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=2, ensure_ascii=False, default=str)
    
    # Выводим итоги
    print(f"\n{'='*80}")
    print("ИТОГИ ОБРАБОТКИ DOCX (КОМБИНИРОВАННЫЙ ПОДХОД)")
    print(f"{'='*80}")
    print(f"Всего файлов: {summary['total_files']}")
    print(f"Успешно: {summary['successful']}")
    print(f"Ошибок: {summary['failed']}")
    print(f"Общее время: {total_time:.2f} сек")
    print(f"Среднее время на файл: {summary['average_processing_time_seconds']:.2f} сек")
    print(f"Метод обработки: {summary['processing_method']}")
    print(f"\nДетальная статистика сохранена в: {summary_file}")
    
    if summary['failed'] > 0:
        print("\nФайлы с ошибками:")
        for result in results:
            if not result.get("success", False):
                print(f"  - {result['file']}: {result.get('error', 'Unknown error')}")


if __name__ == "__main__":
    main()
