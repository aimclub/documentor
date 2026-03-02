"""
Скрипт для демонстрации структуры таблиц в DOCX.
Показывает, как хранятся таблицы и что можно извлечь.
"""

from pathlib import Path
import json
import sys

# Добавляем корень проекта в путь
project_root = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(project_root))

try:
    from docx import Document as PythonDocxDocument
    from docx.oxml.ns import qn
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    print("Предупреждение: python-docx не установлен. Установите: pip install python-docx")
    sys.exit(1)


def analyze_table_structure(docx_path: Path) -> None:
    """
    Анализирует структуру таблиц в DOCX и выводит подробную информацию.
    """
    doc = PythonDocxDocument(str(docx_path))
    
    print("=" * 80)
    print(f"АНАЛИЗ СТРУКТУРЫ ТАБЛИЦ В DOCX: {docx_path.name}")
    print("=" * 80)
    print()
    
    print(f"Всего таблиц в документе: {len(doc.tables)}")
    print()
    
    for table_idx, table in enumerate(doc.tables):
        print(f"{'='*80}")
        print(f"ТАБЛИЦА {table_idx + 1}")
        print(f"{'='*80}")
        
        # Основная информация
        print(f"Индекс: {table_idx}")
        print(f"Количество строк: {len(table.rows)}")
        print(f"Количество столбцов: {len(table.columns) if table.rows else 0}")
        
        # Стиль таблицы
        if table.style:
            print(f"Стиль таблицы: {table.style.name}")
        
        # Анализ структуры
        table_structure = {
            "index": table_idx,
            "rows_count": len(table.rows),
            "cols_count": len(table.columns) if table.rows else 0,
            "style": table.style.name if table.style else None,
            "rows": []
        }
        
        # Анализ каждой строки
        for row_idx, row in enumerate(table.rows):
            row_info = {
                "row_index": row_idx,
                "cells_count": len(row.cells),
                "cells": []
            }
            
            # Анализ каждой ячейки
            for cell_idx, cell in enumerate(row.cells):
                cell_info = {
                    "cell_index": cell_idx,
                    "text": cell.text.strip(),
                    "text_length": len(cell.text.strip()),
                    "paragraphs_count": len(cell.paragraphs),
                    "is_merged": False,  # python-docx не показывает напрямую, нужно проверять через XML
                }
                
                # Проверка на объединенные ячейки через XML
                try:
                    tc = cell._element
                    # Проверяем атрибуты rowspan и colspan
                    rowspan = tc.get(qn('w:rowSpan'))
                    colspan = tc.get(qn('w:gridSpan'))
                    vmerge = tc.get(qn('w:vMerge'))
                    
                    if rowspan or colspan or vmerge:
                        cell_info["is_merged"] = True
                        if rowspan:
                            cell_info["rowspan"] = rowspan
                        if colspan:
                            cell_info["colspan"] = colspan
                        if vmerge:
                            cell_info["vmerge"] = vmerge
                except:
                    pass
                
                # Информация о параграфах в ячейке
                cell_paragraphs = []
                for para in cell.paragraphs:
                    para_info = {
                        "text": para.text.strip(),
                        "style": para.style.name if para.style else None,
                        "runs_count": len(para.runs),
                    }
                    cell_paragraphs.append(para_info)
                
                cell_info["paragraphs"] = cell_paragraphs
                row_info["cells"].append(cell_info)
            
            table_structure["rows"].append(row_info)
        
        # Выводим структуру
        print("\nСтруктура таблицы:")
        print(json.dumps(table_structure, ensure_ascii=False, indent=2))
        
        # Выводим данные таблицы в текстовом виде
        print("\nДанные таблицы (текст):")
        print("-" * 80)
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            print(f"Строка {row_idx + 1}: {' | '.join(row_data)}")
        
        print()
    
    # Сохраняем полную структуру в JSON
    output_file = docx_path.parent / f"{docx_path.stem}_table_structure.json"
    all_tables_structure = []
    
    for table_idx, table in enumerate(doc.tables):
        table_data = {
            "index": table_idx,
            "rows": len(table.rows),
            "cols": len(table.columns) if table.rows else 0,
            "style": table.style.name if table.style else None,
            "data": []
        }
        
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data["data"].append(row_data)
        
        all_tables_structure.append(table_data)
    
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(all_tables_structure, f, ensure_ascii=False, indent=2)
    
    print(f"\nПолная структура сохранена в: {output_file}")


if __name__ == "__main__":
    # Пример использования
    test_folder = Path(r"E:\easy\documentor\documentor_langchain\experiments\pdf_text_extraction\test_folder")
    
    if len(sys.argv) > 1:
        docx_file = Path(sys.argv[1])
    else:
        # Используем первый доступный файл
        docx_files = list(test_folder.glob("*.docx"))
        if not docx_files:
            print("Не найдено DOCX файлов в test_folder")
            sys.exit(1)
        docx_file = docx_files[0]
    
    if not docx_file.exists():
        print(f"Файл не найден: {docx_file}")
        sys.exit(1)
    
    analyze_table_structure(docx_file)
