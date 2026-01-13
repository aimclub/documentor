#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Принудительный парсинг через python-docx для проблемных файлов
"""

from docx import Document
from pathlib import Path
import json


def parse_with_python_docx(docx_path: Path):
    """Парсинг через python-docx."""
    print(f"ПАРСИНГ ЧЕРЕЗ PYTHON-DOCX: {docx_path.name}")
    print("=" * 80)
    
    try:
        doc = Document(str(docx_path))
        
        headings = []
        paragraphs = []
        
        print(f"Всего параграфов в документе: {len(doc.paragraphs)}")
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            style_name = paragraph.style.name
            
            # Проверяем на заголовок
            if style_name.startswith('Heading') and text:
                level = style_name.replace('Heading ', '')
                headings.append({
                    'content': text,
                    'level': f'Heading {level}',
                    'style': style_name,
                    'paragraph_index': i
                })
                print(f"ЗАГОЛОВОК [{level}]: {text[:60]}...")
            elif text:  # Обычный параграф
                paragraphs.append({
                    'content': text,
                    'style': style_name,
                    'paragraph_index': i
                })
        
        print(f"\nНАЙДЕНО:")
        print(f"- Заголовков: {len(headings)}")
        print(f"- Параграфов: {len(paragraphs)}")
        
        if headings:
            print(f"\nЗАГОЛОВКИ ({len(headings)} шт.):")
            print("-" * 60)
            
            for i, h in enumerate(headings, 1):
                level = h['level']
                content = h['content']
                
                if "Heading 1" in level:
                    prefix = "1."
                elif "Heading 2" in level:
                    prefix = "2."
                elif "Heading 3" in level:
                    prefix = "3."
                else:
                    prefix = "•"
                    
                print(f"{i:2d}. {prefix} {content}")
        
        return headings
        
    except Exception as e:
        print(f"ОШИБКА: {e}")
        import traceback
        print(traceback.format_exc())
        return None


def process_problem_files():
    """Обработать проблемные файлы через python-docx."""
    problem_files = [
        "02_Отчет_Этап_2_полный_замечания_эксперта.docx",
        "Diplom2024.docx"
    ]
    
    test_folder = Path("test_folder")
    results = {}
    
    for filename in problem_files:
        file_path = test_folder / filename
        
        if not file_path.exists():
            print(f"Файл не найден: {filename}")
            continue
        
        print(f"\n{'='*80}")
        headings = parse_with_python_docx(file_path)
        
        if headings:
            results[filename] = {
                'success': True,
                'headings_count': len(headings),
                'headings': headings,
                'method': 'python-docx'
            }
            
            # Сохраняем результат
            output_file = file_path.with_suffix('.python_docx_headings.json')
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(headings, f, ensure_ascii=False, indent=2)
            print(f"\nРезультат сохранен в: {output_file}")
        else:
            results[filename] = {
                'success': False,
                'headings_count': 0,
                'headings': [],
                'method': 'python-docx'
            }
    
    # Итоговая статистика
    print(f"\n{'='*80}")
    print("ИТОГОВАЯ СТАТИСТИКА:")
    print("=" * 80)
    
    successful = sum(1 for r in results.values() if r['success'])
    total_headings = sum(r['headings_count'] for r in results.values())
    
    print(f"Успешно обработано: {successful}/{len(problem_files)}")
    print(f"Всего заголовков: {total_headings}")
    
    for filename, result in results.items():
        status = "УСПЕХ" if result['success'] else "ОШИБКА"
        print(f"{status}: {filename} - {result['headings_count']} заголовков")


if __name__ == "__main__":
    process_problem_files()




